"""Extract plain text from .docx files. Usage: python read-docx.py path/to/file.docx"""
import sys
from docx import Document

for path in sys.argv[1:]:
    print(f'\n{"="*80}\n{path}\n{"="*80}')
    doc = Document(path)
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        text = p.text.strip()
        if not text:
            print()
            continue
        if style.startswith('Heading') or style == 'Title':
            print(f'\n## [{style}] {text}\n')
        else:
            print(text)
    # Also dump tables
    for ti, t in enumerate(doc.tables):
        print(f'\n--- Table {ti+1} ({len(t.rows)} rows × {len(t.columns)} cols) ---')
        for row in t.rows:
            cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
            print(' | '.join(cells))
