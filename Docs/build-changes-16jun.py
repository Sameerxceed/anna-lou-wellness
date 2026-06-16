"""
Build a standalone doc summarising what shipped on 16 June 2026.

Anna can read this in 5 minutes without trawling through the full user
manual. Output: Docs/ALW_What_Shipped_16-06-2026.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DARK = RGBColor(0x23, 0x1F, 0x20)
PLUM = RGBColor(0x6E, 0x3A, 0x5A)
GREEN = RGBColor(0x2A, 0x7A, 0x5C)
AMBER = RGBColor(0xC4, 0x70, 0x4A)
MUTE = RGBColor(0x6E, 0x6A, 0x62)


def style_normal(d):
    s = d.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)
    s.font.color.rgb = DARK


def h1(d, text, colour=PLUM):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = colour


def h2(d, text, colour=AMBER):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = colour


def h3(d, text, colour=PLUM):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = colour


def para(d, text, italic=False, colour=DARK, size=11):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.italic = italic; r.font.size = Pt(size); r.font.color.rgb = colour


def bullet(d, text):
    p = d.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)


def status_done(d, text):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("SHIPPED: "); r.bold = True; r.font.color.rgb = GREEN; r.font.size = Pt(10.5)
    r2 = p.add_run(text); r2.font.size = Pt(10.5); r2.font.color.rgb = DARK


d = Document()
style_normal(d)

# Cover
h1(d, "What shipped on 16 June 2026")
para(d, "A summary of every change live on staging after today's deploys. The full how-to lives in ANNA_USER_MANUAL section 17 — this is the 5-minute version.", italic=True, colour=MUTE)

# Section: CMS improvements
h2(d, "Inside Strapi admin")

h3(d, "Bulk fill missing SEO (one button)")
para(d, "Sidebar -> SEO & AI Files -> Run bulk SEO backfill. Walks every Article / Programme / Experience / Coaching Session / Generic Page / Page Builder entry and writes the SEO title + description for any that are empty. Live progress shown. Never overwrites your edits. Takes about 1 minute per 60 entries. Run it again whenever you publish a batch of new articles.")

h3(d, "Site URLs lookup (new sidebar item)")
para(d, "Click Site URLs in the left sidebar to see every page on the site, grouped by section. Search box at the top, two Copy buttons per row: 'Copy path' for Navigation, 'Copy full' for emails and social. Use this instead of typing URLs by hand.")

h3(d, "Quick Photo Editor now works")
para(d, "Sidebar -> Quick Photos. Lists every hero / portrait image across the site as a thumbnail. One-tap Replace per row. Previously said 'Coming soon' — fixed today.")

h3(d, "Upsells field on every page entry")
para(d, "Almost every page singleton now has an Upsells field at the bottom (up to 3 cards). Each card: Label, Link, Eyebrow, Blurb, optional Image. Cards render at the bottom of the page as a 'Where next' section. Leave the field empty to hide the block.")

h3(d, "Searchable URL picker on Upsells")
para(d, "The Link field on each upsell card is no longer free text — it's a dropdown of every URL on the site. Click and type to filter; click to set. You can still paste an external URL like a Calendly link.")

h3(d, "FAQ extended to 6 more pages")
para(d, "FAQ collection now supports per-page Q&A on: Practitioners, REGULATED course, Reset Stories, Life, Love & Relationships, Work & Money. Open FAQ -> +Create -> pick the new page slug from the dropdown.")

h3(d, "Field help text added across the CMS")
para(d, "72 fields across 38 collections now have plain-English help text shown beneath the input. Look out for examples on: SEO title, SEO description, slug, hero image, booking URL, mailchimpTag, kicker, is_active, sort_order, etc.")

h3(d, "Experiences collections renamed for clarity")
para(d, "'Experiences . Event' -> 'Experiences . Event Bookings' (specific events like Align & Amplify, named retreats, workshops). 'Experiences . Sub-page' -> 'Experiences . Category Pages' (the 4 main category landings). Each description now cross-references the other so it's clear where to look for what.")

# Section: Front-end
h2(d, "On the public site (staging)")

h3(d, "Practitioner enquiry form on /practitioners")
para(d, "New 'Apply to be listed' CTA above the FAQ on the Practitioners page. Wellness professionals submit name, email, phone, optional practice, optional message. You get an email AND they get tagged in Mailchimp as 'Practitioner Enquiry' so any Customer Journey you wire to that tag fires.")

h3(d, "Event cards now lead with images")
para(d, "Upcoming events on /experiences/retreats, /experiences/workshops etc. now render as image-led cards (4:3 hero image, date pill overlay in your accent colour, title + meta + CTAs below). Upload hero_image at 1600x1000px+ for best results. Cards stack 3-2-1 across desktop / tablet / mobile.")

h3(d, "Per-page upsells rendering on 25+ pages")
para(d, "The 'Where next' block now renders on Homepage, About, Contact, Community, Reset Letters, Reset Stories, Life, Love & Relationships, Work & Money, Experiences landing, Work with Anna, Sessions hub, Membership, Decoder, Testimonials, Shop, the Returning Circle, every Article detail page, every Programme, every Experience event, every Page Builder page, Reset Room, Practitioners. Stays hidden until you fill the field.")

h3(d, "Decoder quiz popup on homepage (10s after load)")
para(d, "First-time visitors see a small popup 10 seconds in promoting the Nervous System Decoder quiz with a CTA. Dismissible via close / Escape / backdrop click. Won't reappear for 7 days once closed. Tell Sameer if you want the timing or copy moved into CMS.")

h3(d, "Site-wide font bump")
para(d, "Everything is roughly 10% larger after your 'I can't read the small text' feedback. Top nav links, body copy, footer, kickers, eyebrows — all proportionally bigger. No colour or layout changes, just size.")

h3(d, "Scroll-to-top arrow positioned above Ask Anna")
para(d, "The small back-to-top arrow on long pages no longer hides behind the floating Ask Anna pill. It now sits comfortably above it on every viewport.")

# Section: What's still pending
h2(d, "What's NOT live yet (waiting on you)")

bullet(d, "Reset Letters logo — needs the brand asset from you.")
bullet(d, "Founder Reset + Returning Circle sales pages — needs your scope choice.")
bullet(d, "Substack URL + section names — needs confirmation so we can verify the auto-pull.")
bullet(d, "Work & Money submenu structure — needs your clarification.")
bullet(d, "Chris Corsini design screenshot — needs the reference image.")
bullet(d, "Image upload bug — needs a collaborative retest with you.")
bullet(d, "Homepage 6-category magazine grid — needs your reply to the 5 questions.")
bullet(d, "Corporate Wellbeing artwork — needs the hero image uploaded by you.")
bullet(d, "Mailchimp Customer Journey activation — needs you to design each email body and turn the journey on.")
bullet(d, "Stripe live mode — needs your bank verification on Stripe Connect.")
bullet(d, "DNS go-live — needs your cutover date.")

# Section: How to test today's changes
h2(d, "Try it (15 minutes)")
para(d, "Walk through these in order to see everything today's work changed. You'll have first-hand confirmation that each piece works.")

para(d, "1. Open Strapi admin -> Site URLs (left sidebar). Type 'reset' in the search. Click Copy path on any row.", colour=DARK)
para(d, "2. Open Strapi admin -> SEO & AI Files. Scroll to the top 'Bulk fill missing SEO' panel. (Optional: click Run bulk SEO backfill and watch it process.)", colour=DARK)
para(d, "3. Open Strapi admin -> Homepage. Scroll to Upsells. Add an entry. Click the Link field — see the dropdown. Type 'reset' to filter. Pick one. Save + Publish.", colour=DARK)
para(d, "4. Open staging.annalouwellness.com in a fresh tab. Scroll to the bottom of the homepage — see your upsell card render.", colour=DARK)
para(d, "5. Open staging.annalouwellness.com/practitioners. Scroll to the bottom — click 'Apply to be listed'. Fill the form with your own test details. Submit. Check your inbox for the notification.", colour=DARK)
para(d, "6. Open staging.annalouwellness.com/experiences/retreats. See the upcoming-event cards leading with hero images.", colour=DARK)
para(d, "7. Wait 10 seconds on the homepage in an incognito tab — see the Decoder quiz popup.", colour=DARK)

# Footer
para(d, "", colour=DARK)
para(d, "Generated " + __import__('datetime').datetime.now().strftime('%d %b %Y'), italic=True, colour=MUTE, size=9)

out = "Docs/ALW_What_Shipped_16-06-2026.docx"
d.save(out)
print(f"Wrote {out}")
