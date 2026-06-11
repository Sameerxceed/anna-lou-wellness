"""
Build a docx answering every question/concern in Anna's 10 June notes.
Idempotent — re-run after editing.

Run from project root:  python Docs/build-answers-for-anna.py
Output: Docs/ALW_Answers_For_Anna_10-06-2026.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DARK = RGBColor(0x23, 0x1F, 0x20)
PLUM = RGBColor(0x6E, 0x3A, 0x5A)
GREEN = RGBColor(0x2A, 0x7A, 0x5C)
AMBER = RGBColor(0xC4, 0x70, 0x4A)
MUTE = RGBColor(0x6E, 0x6A, 0x62)


def style_normal(d):
    s = d.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)
    s.font.color.rgb = DARK


def h1(d, text, colour=PLUM):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(19); r.font.color.rgb = colour


def h2(d, text, colour=DARK):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = colour


def h3(d, text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = AMBER


def para(d, text, italic=False, colour=DARK, size=11):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.italic = italic; r.font.size = Pt(size); r.font.color.rgb = colour


def quote(d, text):
    """Anna's exact wording from her notes."""
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"“{text}”")
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTE


def bullet(d, text):
    p = d.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(11)


def numbered(d, text):
    p = d.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(11)


def callout(d, label, body, colour=GREEN):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(f"{label}  ")
    r1.bold = True; r1.font.color.rgb = colour; r1.font.size = Pt(11)
    r2 = p.add_run(body); r2.font.size = Pt(11); r2.font.color.rgb = DARK


# ────────────────────────────────────────────────
d = Document()
style_normal(d)

# Cover
title = d.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Anna Lou Wellness"); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = PLUM
sub = d.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Answers to your 10 June notes"); r.italic = True; r.font.size = Pt(16); r.font.color.rgb = DARK
ver = d.add_paragraph(); ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ver.add_run("From Sameer  ·  v3 updated Thursday 11 June 2026"); r.font.size = Pt(10); r.font.color.rgb = MUTE
d.add_paragraph()

para(d,
     "I've gone through every line of your notes. Some of these were already shipped a day or two ago and you might be testing on an older deploy — refresh your browser hard (Ctrl+Shift+R or pull down twice on iPhone) to see the latest. Others are genuinely open and I've grouped them by what I'll do this week.",
     italic=True, colour=MUTE)
para(d,
     "This document follows the order of your notes. Each section quotes you and then explains what's happening.",
     italic=True, colour=MUTE)

# ─── 1. Homepage ───
h1(d, "1. Homepage — Cup of Jo magazine layout")
quote(d, "Homepage needs to look like Cup of Jo almost exactly! So it's clear it is a magazine. Ours has very tiny text on mobile so you can barely read it.")
quote(d, "So can you make sure it only Image, category, preview text, add Read More… so people click read more to read the full article")
quote(d, "Should show a blog post from 6 different categories. The featured one at the top left for each category. Home page text to go elsewhere as it's too much text.")

h3(d, "Status")
para(d, "Partially done. I rebuilt the homepage on 8 Jun to use Cup-of-Jo style cards on mobile — bigger images, bigger text, less clutter (commit reference: e6f71cd). If you're still seeing tiny text, you may be on an older cached version. Hard-refresh and tell me if it's still small.")
para(d, "The 6-category magazine grid you describe is the bigger job. I have NOT started it because I want to be sure I build the right thing — there are a few design calls only you can make.")

h3(d, "5 quick questions I need from you before I build it")
para(d, "Sameer has sent these on WhatsApp. Reply with answers next to each number (1: …, 2: …) and I'll start the build the same day:")
numbered(d, "Cup-of-Jo layout on MOBILE only, or BOTH mobile and desktop? (Your note mentioned tiny mobile text, but also said the layout should be Cup-of-Jo from the start — want to confirm scope.)")
numbered(d, "Which 6 of your 22 sub-categories do you want on the homepage? Pick six, or say 'just pick' and I'll choose sensible defaults you can change later in Strapi.")
numbered(d, "Where does the existing homepage body copy go? (Move to an About page / move to footer / delete / other.)")
numbered(d, "Keep or drop each of: Featured Products strip · Testimonials strip · Press mentions strip · Certifications? (Cup-of-Jo is pure magazine; ALW is part-magazine, part-shop, part-service — your call.)")
numbered(d, "One BIG featured article above the 6 bands, or jump straight into the bands?")

h3(d, "Once you've answered — what I'll build")
bullet(d, "Strip the long body copy off the homepage (per your section 3 answer).")
bullet(d, "Build the category bands. Each: one featured article (large, top-left) + 3 smaller cards from that category.")
bullet(d, "Card structure: image, category label, headline, preview text, Read More — exactly the Cup-of-Jo pattern you described.")
bullet(d, "Pure CMS — you'll be able to swap which categories appear via Strapi (no code change needed).")
bullet(d, "Auto-hide a band if a category has zero articles — page never looks broken.")
para(d, "Build time once unblocked: about 2 hours.", italic=True, colour=MUTE)

# ─── 2. Events: Returning Circle ───
h1(d, "2. Events page — Returning Circle (2 locations weekly)")
quote(d, "I need to show both locations as I do 2 each week at 2 different locations so can I have space for 2 events? And then add more?")

h3(d, "Status")
para(d, "Already shipped on 8 June (commit reference: e6f71cd). The Returning Circle now supports unlimited recurring sessions — you can add as many as you like. They appear in date order with location + time per session.")

h3(d, "How to use it")
numbered(d, "Strapi → Generic Page → Returning Circle.")
numbered(d, "Scroll to the 'Sessions' repeatable field (or similar — labelled in plain English).")
numbered(d, "Click 'Add a session' for each location/time slot. Add as many as you need.")
numbered(d, "Each session has its own location, date, and time.")
numbered(d, "Save → live in 1-2 seconds.")
callout(d, "If you can't find the field:", "Tell me and send a screenshot — it might be hidden in a collapsed section. I can move it to the top if it's buried.")

# ─── 3. Preview while editing ───
h1(d, "3. Preview while editing")
quote(d, "Can you add a preview so I can see whilst I'm designing it so I know the correct layout and colours and fonts show up on the design?")
quote(d, "On the align and Amplify event page. You can't see any preview of the sales page — it jumps to the homepage.")

h3(d, "Status")
para(d, "There are now TWO ways to preview, both shipped on 9 June:")
bullet(d, "Green 'View live page →' button in the right sidebar of every edit screen. Opens the real published page in a new tab.")
bullet(d, "Strapi's built-in Preview tab (icon at the top of the edit screen).")
para(d, "The 'jumps to homepage' bug was a real issue — fixed on 9 Jun (commit b941b39). The preview function was missing the URL mapping for ~20 content types, so it fell through to /. Should be working now. Try the Align & Amplify page again with a hard refresh — if it still goes to homepage, screenshot the URL and send me.")

h3(d, "What 'View live' shows vs what Preview shows")
bullet(d, "View live page → published version. What real visitors see.")
bullet(d, "Preview tab → DRAFT version. What it will look like after you publish.")
para(d, "Use Preview when you're working on a draft and don't want it live yet. Use View Live to check what's already published.")

# ─── 4. Align & Amplify product / sales page ───
h1(d, "4. Align & Amplify — sales page + upsells")
quote(d, "This events page shouldn't have an add to cart or book a call to find out more but instead it just ends and they have to go to a separate product page which is long winded and unnecessary.")
quote(d, "On the Align and Amplify product page I see the page preview after you click save however there is no proper sales page with the sales content.")
quote(d, "May I suggest… Instead of you may also like — And add other services")

h3(d, "Status — DONE (Thursday afternoon)")
para(d, "Shipped. There is now ONE proper sales page per Experience entry, with full sales content and a Book button on it. Anna no longer needs the Product entry as a duplicate.")

h3(d, "What changed")
bullet(d, "New page at /experiences/<slug> for every Experience entry. e.g. /experiences/align-and-amplify. Renders: hero image + title + date + location + price, full description (your sales copy), Book Now button, reviews, FAQ, and 'where next' upsell cards.")
bullet(d, "Listing cards on /experiences and the sub-pages (/experiences/workshops etc.) now click through to this new detail page instead of jumping straight to Calendly. Customers see your sales copy first, decide, then book.")
bullet(d, "If you've set a booking_url (Calendly) on an experience, there's an extra 'Book directly →' button next to 'Read more →' on the listing card for power users who want to skip the sales page.")
bullet(d, "Replaced 'You may also like' with the upsell field you already control. For each experience, edit the upsells field to choose related SERVICES (not products) — that's now what shows at the bottom of each sales page. Products only get product upsells (jewellery), services only get service upsells.")

h3(d, "What you need to do")
numbered(d, "Open the Align & Amplify entry (Strapi → Experiences · Event → Align & Amplify).")
numbered(d, "Confirm the description field has your full sales copy. If not, paste it in. Separate paragraphs with a BLANK LINE between them.")
numbered(d, "Scroll to the 'upsells' field. Add 2-3 related services (e.g. Reset Room, REGULATED, 1:1 Sessions) as 'where next' cards. Leave empty if you don't want any.")
numbered(d, "Save. Visit https://staging.annalouwellness.com/experiences/align-and-amplify to see the live page.")
numbered(d, "Optional: delete the duplicate Align & Amplify Product entry from the Shop collection (it's no longer needed).")

# ─── 5. SEO ───
h1(d, "5. SEO — automate it!")
quote(d, "Can it not automatically fill it in what it thinks? AI should be doing this part so. Don't need to enter it each time on each page.")
quote(d, "It's also pulling a SEO description which isn't suitable for SEO and not what people would write in google or a llm.")

h3(d, "Status — live and tested")
para(d, "Two changes shipped and verified working on staging:")
numbered(d, "AUTOMATIC SEO generation. When you save ANY entry (Article, Programme, Experience, Coaching Session, Generic Page, Page Builder page) and the seo_title / seo_description are blank, the system calls Claude in the background and fills them in. You don't have to click anything. Takes 5–10 seconds — the right-sidebar 'Auto SEO' panel will auto-refresh the page for you so you see the new copy without manually F5-ing.")
numbered(d, "BETTER prompts. I rewrote the instructions Claude follows so descriptions sound like what real people Google. Less 'transformation journey' and more 'Nervous System Reset workshop for founders, Thames houseboat'. I used your own example as the gold standard.")

para(d, "Sample output from a real article in your CMS:")
para(d, "Title: 'I worked so hard not to feel overwhelmed. So why am I still holding everything?'", italic=True, colour=MUTE)
para(d, "→ Auto-generated SEO title: 'Still Holding Everything Even After the Work?'", italic=True, colour=MUTE)
para(d, "→ Auto-generated SEO description: 'A personal essay for women who've done the nervous system work but still carry the emotional load. On the difference between coping well and truly letting go.'", italic=True, colour=MUTE)

h3(d, "What you still control")
bullet(d, "If you've already filled in seo_title or seo_description manually, the system never overwrites you — your words always win.")
bullet(d, "After it auto-fills, you can edit further — your edits stick.")
bullet(d, "If you don't like what Claude wrote, clear both SEO fields and save again — it will rewrite.")

h3(d, "Existing articles (the catch-up question)")
para(d, "The auto-fill only fires on save going forward. Your existing articles (80 of them on staging right now, mostly with empty SEO) won't get filled automatically.")
para(d, "Two options for them:")
bullet(d, "Option A: I run a one-off backfill script that processes all 80 articles + all programmes / experiences / pages in one go. Takes about 10 minutes, fills every empty SEO field across the site. Your manual edits are skipped (never overwritten). I'd recommend this — it's a 'set everything up cleanly' operation.")
bullet(d, "Option B: You touch each article one by one over time. The next time you edit any article, auto-fill kicks in.")
para(d, "Reply and tell me 'run the backfill' if you want option A.")

callout(d, "Try it:", "Open ANY new article → save without filling SEO → the sidebar shows 'Auto-refresh in 10s' → page reloads → seo_title and seo_description are filled with clean copy.")

# ─── 6. Date picker ───
h1(d, "6. Date picker — can't see months ahead")
quote(d, "On the Align & Amplify workshop I was trying to change the date but can't see the months ahead to be able to choose a date.")

h3(d, "Status — partial. The picker is a real Strapi v5 bug; proper fix coming in a follow-up.")
para(d, "I attempted a proper replacement (native iOS / desktop date picker) but the Strapi v5 custom-field API took the CMS down at boot, so I had to revert. The picker itself is now back to the standard Strapi one until I figure out the right plumbing.")

h3(d, "Workaround for now — type the date directly")
para(d, "Strapi date fields accept typed input even when the calendar misbehaves:")
numbered(d, "Tap into the date field.")
numbered(d, "Type the date as YYYY-MM-DD, e.g. 2026-11-15.")
numbered(d, "Tab out. The field accepts it.")
para(d, "Works on every browser; sidesteps the broken picker completely.")

# ─── 7. Booking URL ───
h1(d, "7. Booking URL — what to put there")
quote(d, "And there's no booking url in there so people can't purchase. And I wouldn't know what to put there.")

h3(d, "Status — Calendly integration is live")
para(d, "Any Strapi entry with a 'booking_url' or 'ctaUrl' field can now hold a Calendly link. Paste the link in, and when a customer clicks the Book button on the live site, Calendly opens as a popup right on your page — they pick a time, confirm, and Calendly emails them. They never leave the site.")

h3(d, "What to do")
numbered(d, "Log into Calendly (annalouwellness77 account).")
numbered(d, "Create event types: e.g. 'Discovery Call (15 min)', 'Founder Reset Session (90 min)', 'Speaking Enquiry'.")
numbered(d, "Copy each event-type URL. Looks like: https://calendly.com/anna-annalouoflondon/discovery-call")
numbered(d, "In Strapi, paste it into the booking field on the relevant entry:")
bullet(d, "Programme entries (The Reset, Founder Reset, etc.) → 'ctaUrl' field at the bottom.")
bullet(d, "Experience entries (Align & Amplify, Houseboat workshop) → 'booking_url' field.")
bullet(d, "Coaching Session entries → new 'booking_url' field.")
numbered(d, "Save.")
para(d, "If the URL isn't a Calendly link (e.g. a Stripe checkout URL or an internal /contact page), it still works — it just opens as a regular link instead of a popup. So you can use the same field for anything.")

# ─── 8. Filenames / finding images ───
h1(d, "8. Media library — finding images")
quote(d, "When I upload images to the folder that come out in a very strange format like this so I can't find the images I'm looking for.")

h3(d, "Why we do this")
para(d, "I slugify (= simplify) the FILE NAME on upload — your iPhone 'IMG_4827.HEIC' becomes 'img-4827.jpg'. There's a real reason: Strapi has known bugs with spaces, accents, and apostrophes in filenames, and it breaks image rendering. Slugifying is the safe path.")

h3(d, "How to find your images going forward")
para(d, "Two options:")
numbered(d, "Use the ALT TEXT field every time you upload. That's what's searchable in the Media Library. Type 'Houseboat hero' or 'Anna portrait' as alt text and search will find it.")
numbered(d, "Use the upload date — Media Library sorts newest first by default. Today's uploads are at the top.")

h3(d, "What I can change")
para(d, "Right now we only show filename in the Media Library list. I can add a 'caption' field that shows whatever you type — separate from filename. Want me to do that? It's about 30 min.")

# ─── 9. Corporate Wellbeing page ───
h1(d, "9. Corporate Wellbeing — design not showing")
quote(d, "Corporate wellbeing pages I can't see the design and I see the secondary text but I can't see any artwork etc. Just text jumbled up.")

h3(d, "Status — needs me to inspect")
para(d, "I don't know yet whether this is a CMS data issue (no images attached to the entry) or a rendering bug. Need to check the entry directly.")

h3(d, "Two things I'll do today")
bullet(d, "Open the Corporate Wellbeing entry in Strapi and check if hero_image and other media fields are populated.")
bullet(d, "If they are, check the live page renders them. If not, surface the bug.")
para(d, "Will send a follow-up once I've looked. If you have a screenshot of what you're seeing, that'd help.")

# ─── 10. Content Type Builder ───
h1(d, "10. Content Type Builder — why it's locked")
quote(d, "Can't access previews and buttons on the left on desktop which says Content Type Builder. It is locked. Feels like these are the templates which I need to use to make changes to the templates and move pictures around and alter layout.")

h3(d, "It's not what you think")
para(d, "Content Type Builder is a DEVELOPER tool that lets you create new field types and content types from scratch — like building the schema of a new database table. It is intentionally locked in production because making a wrong change there would break the entire CMS for everyone, instantly.")

h3(d, "What you ACTUALLY want for design changes")
bullet(d, "Image placement / page layout → use the Page Builder (see section 12).")
bullet(d, "Section colours / fonts → those are coded into the design system; tell me what you want and I'll change it.")
bullet(d, "Moving photos around an existing page → edit the entry, swap the image field, save.")
para(d, "I left Content Type Builder visible because Strapi can't fully hide it. But it's locked because using it without knowing exactly what you're doing breaks things. Just ignore it.")

# ─── 11. Adding a new page ───
h1(d, "11. How to add a new page")
quote(d, "How do we add a new page?")

h3(d, "Use the Page Builder (shipped 8 June)")
numbered(d, "Strapi → Page (in the left sidebar — distinct from 'Generic Page').")
numbered(d, "Click 'Create new entry'.")
numbered(d, "Fill in: page title, slug (= URL — e.g. 'our-story' creates /p/our-story).")
numbered(d, "Click 'Add a section' and choose: hero, image+text split, full-bleed image, pull quote, FAQ, etc.")
numbered(d, "Each section has its own fields. Mix and match as you want.")
numbered(d, "Save → 'View live page' button (right sidebar) opens what you've built.")
para(d, "Pages built this way live at /p/<slug>. They can be linked from the nav, footer, or anywhere via just the URL.")

callout(d, "Documentation pending:", "I'll write a video walkthrough of the Page Builder this week. For now, try it on a test page first — call it 'test-page' so you can play without affecting public pages.")

# ─── 12. Image uploads from mobile ───
h1(d, "12. Image uploads — errors on mobile")
quote(d, "Errors uploading images from a mobile")
quote(d, "When I add assets in bulk it loads and then fails each time")

h3(d, "Status — partially fixed, still investigating")
para(d, "I fixed the most common cases on 8 June (commit 8143ebb):")
bullet(d, "iPhone HEIC photos are auto-converted to JPEG on upload.")
bullet(d, "Filename slugified so spaces/apostrophes don't crash the upload.")
bullet(d, "Files processed sequentially (not in parallel) so memory doesn't spike on big batches.")
bullet(d, "Upload size cap raised to handle larger photos.")

h3(d, "Bulk upload still failing — I need details")
para(d, "If bulk upload is still failing for you, send me:")
bullet(d, "How many images you were uploading at once.")
bullet(d, "Are they all photos, or mixed (PDFs, videos)?")
bullet(d, "The exact error message Strapi shows.")
bullet(d, "Whether it's Safari (iPhone) or another browser.")
para(d, "Most likely culprits: total payload too large (try 5-10 at a time instead of 20+), OR the server runs out of memory mid-upload (I can raise the limit if so).")

# ─── 13. FAQ page error ───
h1(d, "13. FAQ page error")
quote(d, "FAQ page error")

h3(d, "I need the error message")
para(d, "Your note just says 'FAQ page error' — I don't know which FAQ, what page, or what the error says. Send me:")
bullet(d, "Which page you were on (the URL, or a description).")
bullet(d, "What you were trying to do (edit, view, add a new FAQ?).")
bullet(d, "Screenshot of the error.")

# ─── 14. REGULATED page ───
h1(d, "14. REGULATED page — where's the design?")
quote(d, "Regulated - where do I add the design for page. I can still just see the fields with text. There are image space but how do I know which image goes where and for which heading and subtext area?")

h3(d, "Where the REGULATED page lives")
para(d, "The REGULATED sales page was rebuilt as a Page Builder page on 8 June (commit a7ee277). It lives in Strapi → Page → 'regulated' (slug).")

h3(d, "How its sections map to the live page")
para(d, "The REGULATED page is made of section components stacked top-to-bottom. Each section in the CMS = one block on the live page. The label on each section tells you what it is.")
bullet(d, "Hero section — top banner with image + title + subtitle.")
bullet(d, "Numbered list section — the '6 weeks, 6 sessions' or whatever structured list.")
bullet(d, "Image + text split — large feature image alongside body copy.")
bullet(d, "Pay-what-you-feel section — the unusual pricing block.")
bullet(d, "Buy Programme section — the final purchase CTA.")
para(d, "Each section is COLLAPSED by default in Strapi. Click on a section header to expand it. The first field inside is usually the heading; the image fields come below.")

h3(d, "I'll add field labels next")
para(d, "Right now field labels say things like 'image_1' which doesn't tell you what image goes where. I'll relabel them this week to say 'Hero image — large photo at top, behind the title' etc. Just need ~30 min.")
callout(d, "View Live", "If you save a draft and click 'View live page →' (green button, right sidebar) you can see exactly which image landed in which spot. Easiest way to learn the mapping until I add the labels.")

# ─── Closing — status by item ───
h1(d, "Status of the week's items — as of this update")

para(d, "Where each open item stands right now:")

h3(d, "DONE — live + tested")
bullet(d, "Automatic SEO on every save (Article, Programme, Experience, Coaching Session, Generic Page, Page Builder pages). Verified working on two real articles in your CMS as of Thursday 11 Jun.")
bullet(d, "Auto-refresh status panel in the right sidebar of every edit page — shows 'Auto-refresh in 10s' after you save so you see the new SEO copy without manually refreshing.")
bullet(d, "Better Google-style SEO description prompt (uses your own example as the gold standard).")
bullet(d, "Calendly popup integration on every Book button — paste the link in booking_url, popup opens.")
bullet(d, "This answers document.")
bullet(d, "Corporate Wellbeing diagnosis — root cause: hero_image is empty in the CMS. Upload an image to the Corporate Wellbeing entry and the page will look right.")
bullet(d, "Hero section + Experience sub-page field labels rewritten with plain-English position cues (TOP-LEFT, RIGHT, BACKGROUND IMAGE etc).")
bullet(d, "Old 'Generate SEO' button removed (it was broken). Replaced with the auto-refresh status panel — same job, no clicks needed.")

h3(d, "PAUSED — waiting on your reply")
bullet(d, "Homepage Cup-of-Jo 6-category magazine grid. I have 5 questions in your inbox to nail down the scope before I build. Quick reply unblocks ~3 hours of work.")
bullet(d, "If we are following Aneeza's full sitemap proposal — that affects the homepage build. Question 7 covers this.")

h3(d, "DONE Thursday afternoon — also live")
bullet(d, "Merged Align & Amplify event + sales into one page at /experiences/align-and-amplify. Same pattern for every Experience entry.")
bullet(d, "Service upsells on event pages — the upsells field on each Experience entry now controls cross-promotion. Anna picks which services to show.")

h3(d, "NEXT — I will start when homepage is unblocked")
bullet(d, "Homepage 6-category magazine grid (waiting on the 5 questions in section 1).")

h3(d, "BLOCKED — I need a screenshot or error message from you")
bullet(d, "Bulk image upload failure — what error does Strapi show, how many images, which browser?")
bullet(d, "FAQ page error — which FAQ, which page, screenshot of the error.")

h3(d, "Your turn — four quick tasks (10 minutes total)")
numbered(d, "Upload a hero image to the Corporate Wellbeing entry (Strapi → Experiences · Sub-page → Corporate Wellbeing → heroImage field).")
numbered(d, "Reply to the 5 (or 6) questions about the homepage rebuild — Sameer is waiting on these to build it.")
numbered(d, "Reply with 'run the backfill' if you want me to fill SEO on all 80 existing articles (and programmes, experiences, pages) in one go.")
numbered(d, "When the FAQ or bulk upload errors next happen, screenshot and send.")

para(d, "Anything I've misread or missed — reply with screenshots. The fewer questions you have to ask twice, the faster I can ship.",
     italic=True)
para(d, "Sameer x", italic=True, colour=MUTE)

out = "Docs/ALW_Answers_For_Anna_10-06-2026.docx"
d.save(out)
print(f"Wrote {out}")
