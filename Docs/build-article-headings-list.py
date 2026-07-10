"""
Build ALW_Article_Headings_For_Covers.docx + .csv — Anna's list of every
article title on the site so she can design thumbnail covers for each.
Grouped by section (Reset Stories, Life, Love & Relationships, Work &
Money), plus a section for "uncategorised" so she can see which need
categorising.
"""

import csv
import json
import urllib.request
from docx import Document
from docx.shared import Pt

API = 'https://cms.annalouwellness.com/api/articles'
FIELDS = (
    '?fields[0]=title&fields[1]=slug'
    '&populate[category][fields][0]=name'
    '&populate[category][fields][1]=section'
    '&pagination[pageSize]=200&sort=title:asc'
)

with urllib.request.urlopen(API + FIELDS, timeout=30) as r:
    payload = json.load(r)

articles = payload['data']

SECTION_LABELS = {
    'reset-stories': 'Reset Stories',
    'life': 'Life',
    'love-and-relationships': 'Love & Relationships',
    'work-and-money': 'Work & Money',
    None: 'Uncategorised',
    '': 'Uncategorised',
}

by_section = {}
for a in articles:
    cat = a.get('category') or {}
    section = cat.get('section') or 'Uncategorised'
    label = SECTION_LABELS.get(section, section)
    by_section.setdefault(label, []).append({
        'title': a['title'],
        'slug': a['slug'],
        'category': cat.get('name') or '',
    })

# ── CSV ───────────────────────────────────────────
csv_path = 'Docs/ALW_Article_Headings_For_Covers.csv'
with open(csv_path, 'w', encoding='utf-8', newline='') as f:
    w = csv.writer(f)
    w.writerow(['Section', 'Sub-category', 'Title', 'URL slug'])
    section_order = ['Reset Stories', 'Life', 'Love & Relationships', 'Work & Money', 'Uncategorised']
    for label in section_order:
        for a in sorted(by_section.get(label, []), key=lambda x: x['title'].lower()):
            w.writerow([label, a['category'], a['title'], a['slug']])
print(f'Wrote: {csv_path}')

# ── Docx ──────────────────────────────────────────
doc = Document()
styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal'].font.size = Pt(11)

h = doc.add_heading('Article Headings — For Cover Design', level=1)
for r in h.runs:
    r.font.size = Pt(20)

p = doc.add_paragraph()
r = p.add_run('Sameer for Anna · 10 July 2026')
r.italic = True
r.font.size = Pt(10)

p = doc.add_paragraph()
p.add_run(
    f'Every article currently in the CMS ({len(articles)} total), grouped '
    'by the editorial section it lives under. Use this list when designing '
    'the cover thumbnails.'
)

for label in ['Reset Stories', 'Life', 'Love & Relationships', 'Work & Money', 'Uncategorised']:
    items = by_section.get(label, [])
    if not items:
        continue
    h2 = doc.add_heading(f'{label}  ({len(items)})', level=2)
    for r in h2.runs:
        r.font.size = Pt(14)
    for a in sorted(items, key=lambda x: x['title'].lower()):
        p = doc.add_paragraph(a['title'], style='List Bullet')
        if a['category']:
            r = p.add_run(f'  ·  {a["category"]}')
            r.italic = True
            r.font.size = Pt(9)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run(
    'Note on "Uncategorised": these articles do not have a Category '
    'assigned in the CMS, so they only appear on their direct URL and '
    'not on any section landing page. Worth reviewing whether they need '
    'a category before you spend cover-design time on them.'
)
r.italic = True
r.font.size = Pt(10)

out = 'Docs/ALW_Article_Headings_For_Covers.docx'
doc.save(out)
print(f'Wrote: {out}')
