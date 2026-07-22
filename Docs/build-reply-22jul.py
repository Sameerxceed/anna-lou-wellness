"""Build ALW_Reply_To_Anna_22-07-2026.docx from Anna's 22 Jul feedback.

Item-by-item response with DONE / DO IN CMS / NEEDS-YOU triage.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


PLUM = RGBColor(0x6E, 0x3A, 0x5A)
INK = RGBColor(0x23, 0x1F, 0x20)
STONE = RGBColor(0x5A, 0x5A, 0x54)
GREEN = RGBColor(0x0A, 0x7A, 0x3B)
AMBER = RGBColor(0xC4, 0x70, 0x4A)


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(20)
    r.font.color.rgb = PLUM
    r.font.bold = True


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(14)
    r.font.color.rgb = INK
    r.font.bold = True


def para(doc, text, size=11, color=INK, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.bold = bold
    return p


def status_line(doc, status, colour, label):
    p = doc.add_paragraph()
    tag = p.add_run(f'[{status}] ')
    tag.font.name = 'Calibri'
    tag.font.size = Pt(11)
    tag.font.color.rgb = colour
    tag.font.bold = True
    txt = p.add_run(label)
    txt.font.name = 'Calibri'
    txt.font.size = Pt(11)
    txt.font.color.rgb = INK
    txt.font.bold = True


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(10.5)
    r.font.color.rgb = STONE


def div(doc):
    p = doc.add_paragraph()
    r = p.add_run('_' * 60)
    r.font.color.rgb = RGBColor(0xE0, 0xE0, 0xE0)


doc = Document()

h1(doc, 'Reply to your 22 July feedback')
para(doc, 'Hi Anna. Went through all 11 items. Here is what is done, what you can do now in CMS, and what needs you to send me one more thing.', color=STONE)
para(doc, ' ')

# ------------------------------------------------------------------------
h2(doc, '1. Retreat page still had the register interest button')
status_line(doc, 'DONE', GREEN, 'Removed 19 July')
para(doc, 'The button is gone. If you are still seeing it, hard-refresh the page (Ctrl+F5 on desktop, or pull-to-refresh on phone).')
div(doc)

# ------------------------------------------------------------------------
h2(doc, '2. Workshop page still had the enquiry button')
status_line(doc, 'DONE', GREEN, 'Removed 19 July')
para(doc, 'Replaced with a "View upcoming workshops" link that takes users to your events page. Same hard-refresh caveat as above.')
div(doc)

# ------------------------------------------------------------------------
h2(doc, '3. Press page still had the press enquiry button')
status_line(doc, 'DONE', GREEN, 'Replaced with a dedicated Press form')
para(doc, 'The generic enquiry button is gone. You now have a proper Press form on the page that sends the enquiry to your inbox (see item 8 below — from today onwards you get an email for every form submission, not just Practitioner).')
div(doc)

# ------------------------------------------------------------------------
h2(doc, '4 + 6. Programme "Book now" buttons — Stripe first, then Calendly')
status_line(doc, 'DONE (needs your action in CMS)', GREEN, 'Wiring shipped')
para(doc, 'This was two items combined. The Reset, Signal, and Signal & Build were all sending users to the Contact page. Now the Book button on each programme page goes to Stripe checkout, takes payment, and redirects them to Calendly to book their first session.')
para(doc, ' ')
para(doc, 'For each programme (The Reset, Signal, Signal & Build), open the Programme entry in Content Manager and fill in TWO fields:', bold=True)
bullet(doc, 'pricePence — the price in pence. £1,500 = 150000. £3,000 = 300000. £750 = 75000.')
bullet(doc, 'postCheckoutCalendlyUrl — the Calendly URL you want them to land on to book their first session AFTER paying.')
para(doc, ' ')
para(doc, 'Save. The Book button on that programme page flips to Stripe checkout automatically. If you leave the price blank, the button falls back to the Contact page (current behaviour) — so nothing breaks while you fill things in.')
div(doc)

# ------------------------------------------------------------------------
h2(doc, '5. Work with Anna footer version you liked')
status_line(doc, 'NEED FROM YOU', AMBER, 'Send me a screenshot')
para(doc, 'I could not find a Work with Anna specific footer variant in the code — the site uses one footer everywhere and it currently mirrors the top nav automatically. Can you WhatsApp me a screenshot of the version you remember? I will restore whichever elements you want back.')
div(doc)

# ------------------------------------------------------------------------
h2(doc, '7. Returning Circle — online recording purchase option')
status_line(doc, 'DONE', GREEN, 'Already on the page')
para(doc, 'The recording purchase block already exists on /community/the-returning-circle. It shows a small purchase box with email field + Buy button. It only appears when a Circle Recording entry is marked "is_available_for_purchase" in the CMS. If nothing is showing right now, it is because no recording is currently marked available. Any Tuesday you tick that flag on a recording, the block appears.')
para(doc, ' ')
para(doc, 'For in-person, the Hare & Moon link is already on the session card as you asked — no purchase flow on our end, they pay externally.')
div(doc)

# ------------------------------------------------------------------------
h2(doc, '8. Email to you after every enquiry form fill')
status_line(doc, 'DONE', GREEN, 'Shipped today')
para(doc, 'You were right — Help · Ask told you that only Practitioner sent me an email. That was true until today. It has been fixed. Every enquiry form on the site (Contact, Press, Retreat, Workshop, Speaking, Corporate, Returning Circle, One Day, Signal Collective, Recovery, Partnerships) now sends you an email at Hello@annalouoflondon.com within seconds of the form submission.')
para(doc, ' ')
para(doc, 'The email you get looks like:', bold=True)
bullet(doc, 'Subject: [Contact Enquiry] someone@email.com')
bullet(doc, 'Body: Name, email, phone, message, submitted time')
bullet(doc, 'From: Anna Lou Wellness <hello@annalouwellness.com>')
bullet(doc, 'Reply-To: your Hello@annalouoflondon.com inbox — so if you hit Reply it goes to you')
para(doc, ' ')
para(doc, 'You can edit the subject line and body wording in Content Manager → Email Template → find "admin_lead_notification". Same for the Practitioner-specific one (admin_practitioner_enquiry). Time to inbox is about 1 to 5 seconds.')
div(doc)

# ------------------------------------------------------------------------
h2(doc, '9. Returning Circle and REGULATED edits not reflecting')
status_line(doc, 'CHECK 3 THINGS', AMBER, 'The refresh code is working — worth ruling out these first')
para(doc, 'The refresh system is wired correctly for both. When you Save a change, the live site is supposed to update within 1-2 seconds. If your edit is not showing, check:')
bullet(doc, '1. Did you click Publish (not just Save)? Save creates a draft — the live site only shows Published entries. There is a Publish button in the top-right of the edit page.')
bullet(doc, '2. Are you looking at the CMS preview iframe, or the actual live URL? Sometimes the preview iframe shows a stale version. Open the site in a fresh tab (staging.annalouwellness.com/...) and hard-refresh.')
bullet(doc, '3. Are you editing "body" or "body_v2"? "body" is the old field, hidden — anything you type there does not show. "body_v2" (or intro_v2, description_v2 etc) is the new field the live site actually reads.')
para(doc, ' ')
para(doc, 'If you still cannot get an edit to reflect after all three checks, WhatsApp me the article/page name and I will look at the specific entry.')
div(doc)

# ------------------------------------------------------------------------
h2(doc, '10. Reset Room price — where to change it for testing')
status_line(doc, 'HERE IS THE PATH', AMBER, 'You are probably on the wrong content type')
para(doc, 'The price you want is on the Work · Membership singleton, NOT on the Reset Room Page (which is the marketing page copy).')
para(doc, ' ')
bullet(doc, 'Content Manager → left sidebar → scroll to "Work · Membership (Reset Room)"')
bullet(doc, 'Open it')
bullet(doc, 'Field name: pricePence')
bullet(doc, 'Set to 0 for a free test. To go back to £25, set to 2500 (it is stored in pence).')
bullet(doc, 'Save')
para(doc, ' ')
para(doc, 'Live checkout will use the new price immediately (Stripe is just a payment processor — the price is set here in your CMS).')
div(doc)

# ------------------------------------------------------------------------
h2(doc, '11. Programme booking links — the WhatsApp screenshots')
status_line(doc, 'DONE', GREEN, 'Covered by items 4/6 and 8 above')
para(doc, 'The screenshots you sent (Contact us for changes / needing to email programmes) are all addressed by the two changes above: (a) Book buttons now go to Stripe checkout, not Contact, and (b) any enquiry form on the site sends you an email now.')
div(doc)

# ------------------------------------------------------------------------
h2(doc, 'What you can do right now')
para(doc, 'Three things, roughly 10 minutes:')
bullet(doc, 'Item 4/6 — fill pricePence + postCheckoutCalendlyUrl on The Reset, Signal, Signal & Build in Content Manager. Book button flips to real checkout.')
bullet(doc, 'Item 5 — WhatsApp me the footer screenshot you remember.')
bullet(doc, 'Test — submit any form on the site with your own email address. You should receive an admin notification at Hello@annalouoflondon.com within seconds.')
para(doc, ' ')
para(doc, 'Everything else is done and live. Anything not behaving as expected, screenshot and send.')
para(doc, ' ')
para(doc, '— Sameer', color=STONE)


doc.save(r'E:\Xceed\Code\Clients\anna-lou-wellness\Docs\ALW_Reply_To_Anna_22-07-2026.docx')
print('OK: Docs/ALW_Reply_To_Anna_22-07-2026.docx')
