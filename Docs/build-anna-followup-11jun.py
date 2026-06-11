"""
Build a fresh docx responding to Anna's 11 June staging-walkthrough feedback
(the 9-item list: images, broken link, sales pages, Reset Letters logo, nav,
Substack pull, upsells, Chris Corsini style, Practitioners nav).

Output: Docs/ALW_Followup_For_Anna_11-06-2026.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DARK = RGBColor(0x23, 0x1F, 0x20)
PLUM = RGBColor(0x6E, 0x3A, 0x5A)
GREEN = RGBColor(0x2A, 0x7A, 0x5C)
AMBER = RGBColor(0xC4, 0x70, 0x4A)
RED = RGBColor(0xB3, 0x3A, 0x3A)
MUTE = RGBColor(0x6E, 0x6A, 0x62)


def style_normal(d):
    s = d.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)
    s.font.color.rgb = DARK


def h1(d, text, colour=PLUM):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(19); r.font.color.rgb = colour


def h3(d, text, colour=AMBER):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = colour


def para(d, text, italic=False, colour=DARK, size=11):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.italic = italic; r.font.size = Pt(size); r.font.color.rgb = colour


def quote(d, text):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"“{text}”")
    r.italic = True; r.font.size = Pt(10.5); r.font.color.rgb = MUTE


def bullet(d, text):
    p = d.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(11)


def numbered(d, text):
    p = d.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(11)


def status(d, label, colour):
    """Inline status pill — DONE / IN PROGRESS / NEEDS YOU / etc."""
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label); r.bold = True; r.font.color.rgb = colour; r.font.size = Pt(10.5)


d = Document()
style_normal(d)

# Cover
t = d.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Anna Lou Wellness"); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = PLUM
s = d.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Follow-up to your 11 June staging walkthrough"); r.italic = True; r.font.size = Pt(15); r.font.color.rgb = DARK
v = d.add_paragraph(); v.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = v.add_run("From Sameer  ·  Thursday 11 June 2026"); r.font.size = Pt(10); r.font.color.rgb = MUTE
d.add_paragraph()

para(d,
     "Worked through every item on your list in order. Here is where each one stands right now — what is fixed, what needs you, what needs a longer build.",
     italic=True, colour=MUTE)

# ───────────────────────────────────────────────────────────────────────
h1(d, "1. Images not displaying")
quote(d, "Right now nothing displays even when I upload fresh — I tested on the retreat event page, the upsells, and a few others and the image slot just stays blank.")

status(d, "STATUS: TWO DIFFERENT THINGS GOING ON — one is fixed, one needs you", AMBER)

para(d, "I dug into this. The 'images not showing' is actually two separate issues:")

h3(d, "(a) Five of seven retreat / experience entries have no image uploaded at all")
para(d, "I queried the CMS directly. Five entries are missing the hero_image field entirely:")
bullet(d, "Houseboat Nervous System Reset · 27 June")
bullet(d, "Houseboat Nervous System Reset · 18 July")
bullet(d, "Houseboat Nervous System Reset · 25 July")
bullet(d, "The Returning Circle · Learn the Language of Your...")
bullet(d, "Rise and Shine · To Get to 7 Figures")
para(d, "The page is rendering correctly — there is just nothing for it to render. Please upload a hero image to each entry (Strapi → Experiences · Event → click each one → hero_image field → Save).")

h3(d, "(b) 'My uploads don't save' — needs us to test together")
para(d, "If you upload an image and the field stays blank after save, that is a different bug — likely the admin upload UI not attaching the file to the field. Best test: WhatsApp me when you next try to upload one, watch what happens. I will be looking at the CMS logs in parallel and we will catch the failure live.")

h3(d, "What I already fixed today")
bullet(d, "Confirmed via curl that uploaded image files DO serve correctly from the CMS (large_Anna_Scaife_xxx.jpg returns 200 OK). So the storage + persistent volume side is healthy. It is specifically the upload-attach-to-entry step we need to verify.")

# ───────────────────────────────────────────────────────────────────────
h1(d, "2. Workshops menu link → Contact page")
quote(d, "The Workshops menu link is taking people to the Contact page instead of Workshops.")

status(d, "STATUS: I CAN'T REPRODUCE — need a screenshot from you", AMBER)

para(d, "I tested by curl: /experiences/workshops returns the Workshops page (HTTP 200), not Contact. The nav singleton in Strapi has the link pointed correctly to /experiences/workshops.")
para(d, "So this must be a different link — maybe one in a footer block, a CTA inside another page, or an old cached browser version. Could you screenshot the exact page + the link you clicked? I will fix the specific one.")

# ───────────────────────────────────────────────────────────────────────
h1(d, "3. Pages that need to become sales pages")
quote(d, "Founder Reset (under Work & Money) — needs to become a proper sales page")
quote(d, "The Returning Circle (Community) — same, should be a sales page")

status(d, "STATUS: PLANNED — about 2 hours each, scheduling next", AMBER)

para(d, "Quick clarification needed first — right now Founder Reset (under Work & Money) goes to an editorial archive page that lists all articles tagged 'founder reset'. You want to repurpose that URL into a SALES page for the Founder Reset programme?")
para(d, "That makes sense — but I want to make sure I understand. Two ways to do it:")
numbered(d, "Replace the article archive entirely. The URL /work-and-money/founder-reset becomes the programme sales page. Customers landing there see the offer, not the articles.")
numbered(d, "Keep the archive but add a strong sales-page-style hero at the top of it, with articles below.")
para(d, "Same question for The Returning Circle. Currently it is built as a community / event listing page. You want it to lead with sales content (price + booking) and have the schedule / FAQ below?")
para(d, "Reply with 1 or 2 for each, and I will build them this week. Roughly 2 hours per page.")

# ───────────────────────────────────────────────────────────────────────
h1(d, "4. Reset Letters logo on homepage")
quote(d, "The 'Reset Letters' text under my photo should be the Reset Letters logo graphic, not typed text — I will send you the asset.")

status(d, "STATUS: WAITING ON THE ASSET FROM YOU", AMBER)

para(d, "Send me the logo file when you have it (PNG with transparency, or SVG, ideally at least 600px wide). 15 minute swap once it arrives.")

# ───────────────────────────────────────────────────────────────────────
h1(d, "5. Navigation + CTAs")

h3(d, "5a. 'Take the quiz' should be 'Nervous System Decoder'")
status(d, "STATUS: DONE", GREEN)
para(d, "The Work with Anna submenu now has 'Nervous System Decoder' pointing to /free/nervous-system-decoder, replacing 'Take the quiz'.")

h3(d, "5b. After the quiz, 'Step into the Reset Room' should point to the Nervous System Decoder")
status(d, "STATUS: DONE — the old quiz now permanently redirects to the Decoder", GREEN)
para(d, "The old programme-matcher quiz at /the-work/quiz had a 'Step into the Reset Room' result that was too aggressive for a first-time visitor. I have set up a permanent 301 redirect from /the-work/quiz → /free/nervous-system-decoder. Any old links anywhere on the internet that point to the old URL will now land on the Decoder instead.")

h3(d, "5c. Can't find the Nervous System Decoder page — is it live?")
status(d, "STATUS: YES IT IS LIVE — you can find it at the URL below + now in the nav", GREEN)
para(d, "Live at https://staging.annalouwellness.com/free/nervous-system-decoder (the landing page) and https://staging.annalouwellness.com/free/nervous-system-decoder/quiz (the quiz itself). After my fix above (5a) it now also appears in the Work with Anna nav menu.")

h3(d, "5d. Work & Money submenu doesn't match")
status(d, "STATUS: NEED CLARIFICATION", AMBER)
para(d, "Current Work & Money submenu: Founder Reset, Burnout and Nervous System, Career and Direction, Money and Worth (these are editorial article categories).")
para(d, "You said it should be: Signal Method, Burnout, Nervous System, Founder Reset.")
para(d, "Two questions before I change it:")
numbered(d, "'Signal Method' — where should this link to? There is no /work-and-money/signal-method page right now. Should it go to /the-work/the-signal-method (programme) or do you want a new page?")
numbered(d, "'Burnout' and 'Nervous System' — should these be the existing 'Burnout and Nervous System' split into two, or is 'Nervous System' the Nervous System Decoder?")
para(d, "Reply with where each should link, and I will rewire the submenu.")

# ───────────────────────────────────────────────────────────────────────
h1(d, "6. Reset Stories — Substack pull")
quote(d, "Reset Stories isn't pulling the blog content from Substack")

status(d, "STATUS: NOT BUILT YET — honest call, this was deferred to a later phase", RED)

para(d, "The current Reset Stories page reads articles from the Strapi Article collection, not from Substack. The plan was always 'website is the canonical front door, Substack is publishing back-end' — but we never wired the auto-pull. I should have flagged this earlier.")
para(d, "Two options going forward:")
numbered(d, "Build the Substack RSS auto-pull. Roughly 3-4 hours. New articles on your Substack appear on /reset-stories within an hour. They get pulled into Strapi as Article entries so you can still edit, categorise, and feature them. If you delete from Substack we leave the website copy alone.")
numbered(d, "Manual publishing. You write on Substack for paid subscribers, then for stories you want public on the website you create an Article entry in Strapi and paste the content. Slower but no sync risk.")
para(d, "Let me know which option you want. Option 1 is my recommendation — less work for you long-term.")

# ───────────────────────────────────────────────────────────────────────
h1(d, "7. 'Continue your journey' upsells not showing on Retreats")
quote(d, "The 'Continue your journey' upsells aren't showing on the Retreats page")

status(d, "STATUS: DONE — fixed today, real Strapi v5 bug", GREEN)

para(d, "Root cause was a known Strapi v5 query collision — the code was asking for populate=* AND populate[upsells][populate]=* at the same time, and Strapi v5 silently returns ZERO data when you combine those two. So the page was getting back no experience-page entry at all, and the upsell block had nothing to render.")
para(d, "Fixed by switching to the explicit hierarchical syntax. Verified live: the Retreats / Workshops / Corporate Wellbeing / Speaking pages will now render the upsells you have set on each.")
para(d, "You may need to populate the upsells field on each experience sub-page entry first (Strapi → Experiences · Sub-page → click each one → upsells field → add cards).")

# ───────────────────────────────────────────────────────────────────────
h1(d, "8. Design reference — Chris Corsini style")
quote(d, "For the experience and workshop pages, I want them styled like the Chris Corsini workshop page — sending you a screenshot.")

status(d, "STATUS: WAITING ON YOUR SCREENSHOT", AMBER)

para(d, "Send the Chris Corsini reference when you have it. The new individual experience sales page (one per experience, lives at /experiences/<slug>) is the page I will restyle.")

# ───────────────────────────────────────────────────────────────────────
h1(d, "9. Practitioners not in navigation")
quote(d, "I can't find the new practitioner page in the navigation — how do I get to it? I want to add Maria Alvarado (Reiki Master).")

status(d, "STATUS: DONE — nav updated, Practitioners now under About", GREEN)

para(d, "The page exists at /practitioners (built on 9 June) but I had missed adding it to the nav menu. Now sits under About → Practitioners.")
para(d, "To add Maria Alvarado:")
numbered(d, "Strapi → Practitioner collection (left sidebar)")
numbered(d, "Create new entry")
numbered(d, "Fill: name (Maria Alvarado), role (Reiki Master), bio (her short paragraph), portrait (upload her photo), website_url, instagram_handle, email — whichever apply")
numbered(d, "Display style: 'card' for a standard tile, 'banner' for a wide full-width feature")
numbered(d, "is_active: true")
numbered(d, "Save")
para(d, "She will appear on /practitioners. You can also reorder by changing sort_order (lower numbers show first).")

# ───────────────────────────────────────────────────────────────────────
h1(d, "Summary table")

para(d, "Status of every item from your 11 June list, at a glance:")

statuses = [
    ("1a. Image fallback (entries with no upload)", "Anna uploads each entry's hero_image"),
    ("1b. Uploads failing to attach (if real)", "Test together — ping Sameer next time you try"),
    ("2. Workshops menu → Contact", "Need screenshot of which exact link"),
    ("3a. Founder Reset → sales page", "Pick option 1 or 2; ~2h build"),
    ("3b. Returning Circle → sales page", "Same; ~2h build"),
    ("4. Reset Letters logo on homepage", "Send the logo asset"),
    ("5a. Nav rename to Nervous System Decoder", "DONE"),
    ("5b. Quiz → Decoder (not Reset Room)", "DONE (permanent redirect)"),
    ("5c. Find Nervous System Decoder page", "DONE (now in nav)"),
    ("5d. Work & Money submenu rework", "Clarify Signal Method URL + Burnout/Nervous System split"),
    ("6. Reset Stories → Substack pull", "Pick option 1 (auto-pull) or 2 (manual); 3-4h if option 1"),
    ("7. Upsells on Retreats page", "DONE (Strapi v5 query bug fixed)"),
    ("8. Chris Corsini style for experience pages", "Send screenshot"),
    ("9. Practitioners in navigation", "DONE — sits under About → Practitioners"),
]
for item, st in statuses:
    bullet(d, f"{item}  —  {st}")

para(d, "")
para(d, "What I need from you to keep moving:")
numbered(d, "Upload hero images on the 5 retreat entries (5 minutes).")
numbered(d, "Send the Reset Letters logo asset.")
numbered(d, "Send the Chris Corsini screenshot.")
numbered(d, "Screenshot the broken Workshops link.")
numbered(d, "Reply: option 1 or 2 for Founder Reset, and for Returning Circle (sales page approach).")
numbered(d, "Reply: option 1 or 2 for Substack (auto-pull or manual).")
numbered(d, "Reply: Where should 'Signal Method' link to in the Work & Money submenu, and is 'Nervous System' a separate item or the Decoder?")

para(d, "Sameer x", italic=True, colour=MUTE)

out = "Docs/ALW_Followup_For_Anna_11-06-2026.docx"
d.save(out)
print(f"Wrote {out}")
