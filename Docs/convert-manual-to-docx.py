"""
Convert ANNA_USER_MANUAL.md to ANNA_USER_MANUAL.docx using python-docx.

Handles: H1-H4 headings, body paragraphs, bullet/numbered lists, markdown tables,
inline bold/italic/code, and horizontal rules. Output uses Word's built-in
styles so the doc has a navigable outline.

Much faster than Word COM (seconds vs minutes for an 800-line doc).

Usage:
    python convert-manual-to-docx.py
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(SCRIPT_DIR, 'ANNA_USER_MANUAL.md')
DOCX_PATH = os.path.join(SCRIPT_DIR, 'ANNA_USER_MANUAL.docx')


def add_inline_runs(paragraph, text):
    """Parse inline markdown (**bold**, *italic*, `code`) and add as runs."""
    # Pattern matches one of: **bold**, *italic*, `code`, or plain text run
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*\n]+\*|`[^`]+`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def flush_table(doc, rows):
    """Render accumulated table rows into a Word table."""
    if not rows:
        return
    # Strip separator rows like |---|---|
    data = [r for r in rows if not re.fullmatch(r'\s*\|?[\s\-:|]+\|?\s*', r)]
    if not data:
        return
    cells_per_row = []
    for raw in data:
        clean = raw.strip().strip('|').strip()
        cells = [c.strip() for c in re.split(r'\s*\|\s*', clean)]
        cells_per_row.append(cells)
    nrows = len(cells_per_row)
    ncols = max(len(r) for r in cells_per_row)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Light Grid Accent 1'
    for r, cells in enumerate(cells_per_row):
        for c in range(ncols):
            cell_text = cells[c] if c < len(cells) else ''
            # Strip inline markdown for cells (simple)
            cell_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
            cell_text = re.sub(r'`([^`]+)`', r'\1', cell_text)
            cell = table.cell(r, c)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            if r == 0:
                run.bold = True


def main():
    if not os.path.exists(MD_PATH):
        raise SystemExit(f'Source not found: {MD_PATH}')

    with open(MD_PATH, encoding='utf-8') as f:
        lines = f.read().splitlines()

    doc = Document()

    table_rows = []
    in_table = False

    for line in lines:
        # Table detection
        if line.lstrip().startswith('|'):
            in_table = True
            table_rows.append(line)
            continue
        elif in_table:
            flush_table(doc, table_rows)
            table_rows = []
            in_table = False

        if not line.strip():
            continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            level = min(len(m.group(1)), 4)
            heading = doc.add_heading('', level=level)
            add_inline_runs(heading, m.group(2))
            continue

        # Horizontal rule
        if re.match(r'^-{3,}\s*$', line) or re.match(r'^\*{3,}\s*$', line):
            p = doc.add_paragraph()
            p.add_run('_' * 60).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            continue

        # Bullet lists
        m = re.match(r'^[-*]\s+(.+)$', line)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            add_inline_runs(p, m.group(1))
            continue

        # Numbered lists
        m = re.match(r'^\d+\.\s+(.+)$', line)
        if m:
            p = doc.add_paragraph(style='List Number')
            add_inline_runs(p, m.group(1))
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, line)

    flush_table(doc, table_rows)

    doc.save(DOCX_PATH)
    print(f'Wrote: {DOCX_PATH} ({os.path.getsize(DOCX_PATH):,} bytes)')


if __name__ == '__main__':
    main()
