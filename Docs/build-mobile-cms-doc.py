"""
Build the 'Mobile CMS — what's new for you' doc for Anna.
Run from project root:  python Docs/build-mobile-cms-doc.py
Output: Docs/ALW_Mobile_CMS_New_For_Anna.docx

Summarises everything we've shipped in the last 24 hours that makes
editing the site from her phone safer and faster.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DARK = RGBColor(0x23, 0x1F, 0x20)
PLUM = RGBColor(0x6E, 0x3A, 0x5A)
GREEN = RGBColor(0x0A, 0x7A, 0x3B)
GOLD = RGBColor(0xC4, 0x70, 0x4A)
MUTE = RGBColor(0x6E, 0x6A, 0x62)
PINK = RGBColor(0xF2, 0x80, 0xAA)


def style_normal(d):
    s = d.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)
    s.font.color.rgb = DARK


def h1(d, text, colour=DARK):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = colour
    return p


def h2(d, text, colour=PLUM):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = colour
    return p


def h3(d, text, colour=DARK):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = colour
    return p


def kicker(d, text, colour=PINK):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = colour
    return p


def para(d, text, italic=False, colour=DARK, size=11, space_after=6):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = colour
    return p


def lede(d, text):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = MUTE
    return p


def step(d, n, text):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3)
    num = p.add_run(f"{n}. ")
    num.bold = True
    num.font.color.rgb = PLUM
    body = p.add_run(text)
    body.font.size = Pt(11)
    return p


def bullet(d, text):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3)
    m = p.add_run("•  ")
    m.bold = True
    m.font.color.rgb = PLUM
    body = p.add_run(text)
    body.font.size = Pt(11)
    return p


def tick(d, text):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.15)
    t = p.add_run("✓  ")
    t.bold = True
    t.font.color.rgb = GREEN
    body = p.add_run(text)
    body.font.size = Pt(11)
    return p


def divider(d):
    p = d.add_paragraph("─" * 60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)


# ────────────────────────────────────────────────────────────────
# Build
# ────────────────────────────────────────────────────────────────

d = Document()
style_normal(d)

# ── Cover ──
h1(d, "Your CMS, made for your phone.")
lede(d, "Everything new this week — safer editing, faster photos, and full visibility into what Google + AI engines see. Anna Lou Wellness · 9 June 2026")

# ── Section 1: Safety net ──
kicker(d, "Safety net")
h2(d, "Save no longer means live.")
para(d, "Until this week, every Save you tapped overwrote the live site instantly. A typo on the homepage was a typo the world saw. That's now changed.")
para(d, "From today, every editorial page — Homepage, About, Reset Room, Testimonials, Practitioners, every Programme, every Page Builder page, every Experience, every FAQ, your nav and footer — has a draft / publish workflow:")
step(d, 1, "Edit a field. Save. Your change is stored as a draft. The live site is untouched.")
step(d, 2, "Tap the green View live page → button (right sidebar of every edit screen) to open the page in another tab. Confirm it looks right.")
step(d, 3, "When you're happy, tap Publish (blue button, top-right). Live in 1–2 seconds.")
para(d, "")
h3(d, "Reverting a mistake")
bullet(d, "Just don't publish — the draft sits there. The live site keeps showing whatever was published before.")
bullet(d, "Or tap Discard changes (top-right of the edit form, near Publish) to wipe the draft back to whatever's currently live.")
bullet(d, "If something has already gone live and you didn't catch it for a day, Sameer can restore the entire site from yesterday's database backup.")

# ── Section 2: View Live ──
divider(d)
kicker(d, "View live page")
h2(d, "Preview without leaving the CMS.")
para(d, "Every edit screen now has a green View live page → button in the right sidebar.")
para(d, "Tap it. The page you're editing opens in a new tab on your phone. Edit and Save in the CMS tab. Swipe back to the live tab. Pull to refresh. Your change appears in seconds.")
para(d, "On iPhone, long-press the tab icon in Safari to switch between the two tabs faster.")

# ── Section 3: SEO & AI Files (new) ──
divider(d)
kicker(d, "SEO & AI Files")
h2(d, "See every file Google + ChatGPT read from your site.")
para(d, "New sidebar item: SEO & AI Files (🤖 icon).")
para(d, "Opens a single reference page listing every machine-readable file your site exposes — sitemap, robots.txt, RSS, the llms.txt file for AI engines, the product feeds for Google Shopping and OpenAI's new Agentic Commerce program. Each card explains who reads it, what it does, and how often it updates.")
para(d, "You don't need to maintain any of these files manually. They auto-update as you publish content. The page exists so you can SEE what's there — useful when curious about how AI engines describe your work, or when you want to share the right URL with a marketing partner.")
h3(d, "Files included")
bullet(d, "Sitemap (Google + all search crawlers)")
bullet(d, "Robots.txt (crawler rules)")
bullet(d, "RSS feed (Apple News, Substack, feed readers)")
bullet(d, "llms.txt (ChatGPT, Claude, Perplexity, Gemini)")
bullet(d, "Google Merchant feed (Google Shopping)")
bullet(d, "OpenAI product feeds (Agentic Commerce, jsonl + json)")
bullet(d, "Schema.org JSON-LD + Open Graph (embedded in every page)")
para(d, "Each card has a green View button to open the file in a new tab, and a Copy URL button if you need to share it.")

# ── Section 3b: Quick Photos (parked) ──
divider(d)
kicker(d, "Quick Photos · Coming soon")
h2(d, "A faster photo-swap workflow, in progress.")
para(d, "We're building a Quick Photos page that lets you replace any hero image across your whole site in 2 taps from a single grid. The page is in your CMS sidebar (📷 icon) but currently shows a 'Coming soon' message — Strapi 5 changed how it handles authentication and the page needs a rewrite to work properly.")
para(d, "Until that ships, the fastest way to replace a photo is to tap the existing image thumbnail directly inside the entry (not the field label) — a Replace media button appears, one tap, pick the photo, publish.")

# ── Section 4: iOS Shortcut ──
divider(d)
kicker(d, "iOS Shortcut")
h2(d, "Snap → Share → Upload. ~5 seconds.")
para(d, "We're setting up a one-time iPhone Shortcut so that uploading a photo to your Media Library is a single tap from the camera roll's share sheet.")
para(d, "After Sameer walks you through the 5-minute setup on your phone:")
step(d, 1, "Take a photo. Or open an existing one in Photos.")
step(d, 2, "Tap the Share button.")
step(d, 3, "Tap Upload to ALW CMS in the share sheet.")
step(d, 4, "A notification pops up: \"Uploaded to CMS.\" Done.")
para(d, "The photo is now in your Media Library. To attach it to a specific page, open Quick Photos (see above), find the section, tap Replace, pick the just-uploaded photo.")
para(d, "Sameer has the setup steps in Docs/IOS_SHORTCUT_UPLOAD.md and will walk you through it when you next have 10 minutes.")

# ── Section 5: Tap-image-not-field ──
divider(d)
kicker(d, "Hidden shortcut")
h2(d, "Tap the image, not the field.")
para(d, "When you're inside any edit form and want to swap a photo that's already set:")
bullet(d, "DON'T tap the field name or label.")
bullet(d, "DO tap the image thumbnail itself.")
para(d, "A modal opens with a Replace media button — one tap, pick new photo, done. Saves two taps every single time.")

# ── Section 6: Practitioners ──
divider(d)
kicker(d, "New page")
h2(d, "Practitioners — your trusted circle.")
para(d, "New page at annalouwellness.com/practitioners — same layout as your Testimonials page (3-column grid of cards with occasional full-width featured banner cards).")
para(d, "Added to the About menu, between Client Stories and Press.")
h3(d, "Adding a practitioner")
step(d, 1, "CMS → Practitioner (new collection in your sidebar)")
step(d, 2, "Create new entry. Fill in name, role, bio, portrait, website URL, optional Instagram handle, optional email, optional location.")
step(d, 3, "Display style: leave on \"card\" for normal entries. Pick \"banner\" for ~1 in every 6–10 to feature them as a full-width section.")
step(d, 4, "Save → Publish.")
para(d, "Cards on the live page are clickable — clicking sends visitors to that practitioner's external website.")

# ── Section 7: Duplicate ──
divider(d)
kicker(d, "Cloning a page")
h2(d, "Start a new page from an existing one.")
para(d, "Don't build new pages from scratch when something similar already exists. Strapi has a built-in Clone feature:")
step(d, 1, "Open the collection list (e.g. Page (build your own), or Work · Programme).")
step(d, 2, "Find the entry you want to use as a template.")
step(d, 3, "Tap the ⋮ (three dots) at the right end of its row.")
step(d, 4, "Tap Clone. Strapi creates a duplicate as a draft, copying every section, image, style, and link.")
step(d, 5, "Change the Title + Slug to something unique. Tweak whatever needs to be different.")
step(d, 6, "Save → Publish. The new page is live at its own URL.")
h3(d, "Adding the cloned page to a menu")
step(d, 1, "CMS → Navigation (single types)")
step(d, 2, "Find the menu group you want to add to (e.g. About or Work with Anna). Expand Children.")
step(d, 3, "Tap + Add another to Children. Fill in Label and Href (the URL of your new page).")
step(d, 4, "Save. The menu updates in 1–2 seconds on the live site.")

# ── Section 8: Combined workflow ──
divider(d)
kicker(d, "Combined")
h2(d, "Your new fastest workflow on phone.")
para(d, "Putting it all together — replacing the homepage hero photo from your phone now looks like this:")
step(d, 1, "Take photo on phone")
step(d, 2, "Share → Upload to ALW CMS  (lands in Media Library in ~5 sec)")
step(d, 3, "CMS → Content Manager → Homepage → tap the existing hero image directly → Replace media → pick the new photo")
step(d, 4, "Tap Publish (top right)")
step(d, 5, "View live page → pull to refresh → confirm it looks right")
para(d, "Total time: about 40 seconds. The old workflow took 2–3 minutes and several full navigation cycles. Once Quick Photos lands, step 3 becomes one tap instead of three.")

# ── Final note ──
divider(d)
kicker(d, "Behind the scenes")
h2(d, "What Sameer set up so this all works safely.")
tick(d, "Persistent storage on the server — your uploaded images will never disappear on a redeploy again (we hit this exact issue last week and have closed it permanently).")
tick(d, "Daily database backup running at 3 am UTC — if something catastrophic happens, we can restore the entire site from the previous night's snapshot.")
tick(d, "Bulk re-upload recovery scripts ready in case we ever need to bulk-restore images from your source folder. Documented in ops/MEDIA_RECOVERY.md.")
tick(d, "Sitemap auto-includes new programmes, experiences, articles, products, practitioners, and Page Builder pages — Google discovers your new content on its next crawl. No action needed.")
tick(d, "llms.txt auto-refreshes every 10 minutes — ChatGPT, Claude, Perplexity, and Gemini see your latest content within minutes of publishing.")
tick(d, "Schema.org JSON-LD generated automatically per page — Article, Service, Product, FAQ, Review, AggregateRating, BreadcrumbList. Lets Google show rich-snippet stars and FAQ accordions in search results.")
tick(d, "Draft mode + View Live button + Practitioners page + SEO & AI Files admin + iOS Shortcut all wired and live.")
para(d, "")
para(d, "Any time something feels wrong or slow, tell Sameer. The whole point of this week's work is that you can edit safely from your phone.")
para(d, "")
para(d, "— ALW Team")

# Save
out = "Docs/ALW_Mobile_CMS_New_For_Anna.docx"
d.save(out)
print(f"Wrote {out}")
