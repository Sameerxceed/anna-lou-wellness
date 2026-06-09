"""
Build Anna's Order Operations Playbook docx.

Run from project root:  python Docs/build-order-playbook.py
Output: Docs/ALW_Order_Playbook_For_Anna.docx

Covers every workflow Anna needs once real orders start landing:
  - The order lifecycle (paid → shipped → completed)
  - Cancellations + refunds (auto-refunds via Stripe)
  - Returns (customer-initiated, you approve/refund)
  - Customer accounts (auto-created on first order, single login)
  - Email templates (edit subject + body in CMS, never code)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DARK = RGBColor(0x23, 0x1F, 0x20)
PLUM = RGBColor(0x6E, 0x3A, 0x5A)
GREEN = RGBColor(0x2A, 0x7A, 0x5C)
AMBER = RGBColor(0xC4, 0x70, 0x4A)
MUTE = RGBColor(0x6E, 0x6A, 0x62)
LIGHT_BG = RGBColor(0xF5, 0xF0, 0xE8)


def style_normal(d):
    s = d.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)
    s.font.color.rgb = DARK


def h1(d, text, colour=PLUM):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = colour
    r.font.name = "Calibri"


def h2(d, text, colour=DARK):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = colour


def h3(d, text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = AMBER


def para(d, text, italic=False, colour=DARK, size=11):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = colour


def bullet(d, text, indent=0):
    p = d.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(11)


def numbered(d, text):
    p = d.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(11)


def callout(d, label, body, colour=AMBER):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r1 = p.add_run(f"{label} ")
    r1.bold = True
    r1.font.color.rgb = colour
    r1.font.size = Pt(11)
    r2 = p.add_run(body)
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARK


def divider(d):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("— — —")
    r.font.color.rgb = MUTE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ───────────────────────────────────────────────────────────────────────
d = Document()
style_normal(d)

# Cover
title = d.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Anna Lou Wellness")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = PLUM

sub = d.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Order Operations Playbook")
r.font.size = Pt(18)
r.font.color.rgb = DARK
r.italic = True

ver = d.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ver.add_run("Version 1.0  ·  June 2026")
r.font.size = Pt(10)
r.font.color.rgb = MUTE

d.add_paragraph()
para(
    d,
    "This is your playbook for everything that happens after a customer places an order: "
    "shipping, returns, refunds, cancellations, customer accounts, and the emails they receive. "
    "Every step is designed so you only have to change ONE thing in Strapi — the status — and "
    "the right things happen automatically (Stripe refund fires, customer email sends, stock updates).",
    italic=True,
    colour=MUTE,
)
para(
    d,
    "Keep this open while you process your first few orders. After that you will know it by heart.",
    italic=True,
    colour=MUTE,
)

# ───────────────────────────────────────────────────────────────────────
h1(d, "1. The big picture")

para(
    d,
    "When a customer places an order, this is what happens — most of it without you doing anything.",
)

h3(d, "Card payment (Stripe)")
numbered(d, "Customer pays on Stripe's secure page. Card charged.")
numbered(d, "Order appears in Strapi → Shop · Order with status PAID.")
numbered(d, "Stock decremented on each product.")
numbered(d, "If first-time buyer: a customer account is created automatically using their email.")
numbered(d, "Customer receives: order confirmation email + (if new) welcome email with link to set their password.")
numbered(d, "You receive: '[New order] ALW-XXXXXXXX · £XX.XX · PAID' email.")
numbered(d, "Stripe generates a downloadable invoice PDF in your Stripe dashboard.")

h3(d, "Bank transfer")
numbered(d, "Customer chooses bank transfer at checkout. Sees bank details on screen.")
numbered(d, "Order appears in Strapi → Shop · Order with status PENDING.")
numbered(d, "Customer receives: order received email WITH bank details and reference number.")
numbered(d, "If first-time buyer: a customer account is created + welcome email sent.")
numbered(d, "You receive: '[New order] ALW-XXXXXXXX · £XX.XX · BANK PENDING' email.")
numbered(d, "When the money lands in your bank, you manually mark the order PAID in Strapi.")

callout(
    d,
    "Key idea:",
    "One thing creates the customer account — the first order. The same login works for the shop, the Reset Room, and any course they buy. You never have to manage accounts manually.",
)

# ───────────────────────────────────────────────────────────────────────
h1(d, "2. Marking an order shipped")

para(d, "When you post the parcel, do this in Strapi (it takes 30 seconds).")

h3(d, "Step-by-step")
numbered(d, "Open the order in Strapi → Shop · Order.")
numbered(d, "Fill in TRACKING NUMBER (e.g. RX123456789GB).")
numbered(d, "Fill in TRACKING URL (e.g. https://track.royalmail.com/RX123456789GB).")
numbered(d, "Fill in SHIPPING CARRIER (e.g. Royal Mail, DPD, Evri).")
numbered(d, "Change STATUS from 'paid' to 'shipped'.")
numbered(d, "Save.")

para(d, "The customer immediately gets the 'Your order is on its way' email. The email includes all three fields you just filled in, plus a Track-Your-Parcel button that links to the URL.")

callout(
    d,
    "Important:",
    "Fill in the tracking fields BEFORE you change the status. If you change the status first and add the tracking number later, the email has already been sent with empty placeholders.",
    colour=AMBER,
)

# ───────────────────────────────────────────────────────────────────────
h1(d, "3. Marking an order complete (delivered)")

para(
    d,
    "About a week after shipping (or when you know the parcel has been delivered), mark it complete. "
    "This sends the customer a friendly 'Hope you are enjoying your order' email asking them to leave a review.",
)

h3(d, "Step-by-step")
numbered(d, "Open the order in Strapi.")
numbered(d, "Change STATUS from 'shipped' to 'completed'.")
numbered(d, "Save.")

para(
    d,
    "That is it. The review-request email fires automatically. There is no penalty for marking late — do it once a week if you prefer to batch.",
    italic=True,
    colour=MUTE,
)

# ───────────────────────────────────────────────────────────────────────
h1(d, "4. Cancelling an order")

para(
    d,
    "Use this when you cannot fulfil an order (out of stock for too long, customer mistake, etc.). "
    "Cancellation by itself does NOT issue a refund — see section 5 for refunds. If you want to "
    "cancel AND refund, do them as separate status changes (cancel first, then refund).",
)

h3(d, "Step-by-step")
numbered(d, "Open the order in Strapi.")
numbered(d, "Fill in CANCELLATION REASON — write what the customer will read (e.g. \"We have run out of stock and a restock will take 4+ weeks.\")")
numbered(d, "Change STATUS to 'cancelled'.")
numbered(d, "Save.")

para(d, "Customer gets the 'Your order has been cancelled' email, including your reason text. If you also need to refund their card, do step 5 next.")

# ───────────────────────────────────────────────────────────────────────
h1(d, "5. Refunding an order")

para(
    d,
    "This is the magic one. Changing the status to 'refunded' AUTOMATICALLY fires the Stripe refund "
    "for you — you do not need to open the Stripe dashboard. The money goes back to the customer's "
    "card within minutes, the customer gets a 'Your refund has been processed' email, and the order "
    "is stamped with the Stripe refund ID so we have a permanent record.",
)

h3(d, "Full refund (most common)")
numbered(d, "Open the order in Strapi.")
numbered(d, "Leave REFUND AMOUNT blank.")
numbered(d, "Change STATUS to 'refunded'.")
numbered(d, "Save.")
para(d, "The full order total is refunded to the customer's card. Done.")

h3(d, "Partial refund (e.g. customer keeps £40 of a £100 order, refund £60)")
numbered(d, "Open the order in Strapi.")
numbered(d, "Fill in REFUND AMOUNT — enter the AMOUNT TO REFUND, not what they keep. e.g. 60.00")
numbered(d, "Change STATUS to 'refunded'.")
numbered(d, "Save.")

h3(d, "Bank transfer order refund")
para(
    d,
    "Bank transfer orders cannot be auto-refunded (Stripe was never involved). The status change still "
    "sends the customer the refund email, but YOU need to send the money back via your bank manually.",
)

callout(
    d,
    "Safety net:",
    "Once an order is refunded, it cannot be refunded again — the system stamps the Stripe refund ID and blocks any further attempts. So if you accidentally hit save twice, nothing bad happens.",
    colour=GREEN,
)
callout(
    d,
    "If Stripe rejects the refund:",
    "(e.g. card expired, payment too old) the status will still say 'refunded' in Strapi, but you will get an admin email with the error so you can refund manually via Stripe dashboard. The customer email will NOT have been sent so they are not confused.",
    colour=AMBER,
)

# ───────────────────────────────────────────────────────────────────────
h1(d, "6. Handling a return")

para(
    d,
    "Customers can request a return themselves from their account page on the website (Account → "
    "click 'Request a return' next to the order). When they submit, a Return Request appears in "
    "Strapi and you get an admin email.",
)

para(
    d,
    "Returns are eligible only for orders in status PAID, SHIPPED, or COMPLETED. Cancelled and "
    "already-refunded orders cannot be returned.",
)

h3(d, "When a return request comes in")
numbered(d, "You receive '[Return requested] ALW-XXXXXXXX' email with the reason.")
numbered(d, "Customer receives 'We have received your return request' acknowledgement.")
numbered(d, "Open Strapi → Shop · Return Request → click into the new one (status 'requested').")
numbered(d, "Decide: approve or reject.")

h3(d, "To approve a return")
numbered(d, "Fill in NOTES — write the shipping instructions the customer needs (where to send it, how to package, who pays postage, any special instructions). This text goes into their email.")
numbered(d, "Change STATUS to 'approved'.")
numbered(d, "Save.")
para(d, "Customer gets 'Your return is approved — here is how to ship it back' email with your notes.")

h3(d, "When the items arrive back to you")
numbered(d, "Optional: change STATUS to 'received' so you know they are back. No email is sent for this transition.")

h3(d, "To process the refund")
numbered(d, "Fill in REFUND AMOUNT — how much to refund (you may keep an amount for return shipping or restocking).")
numbered(d, "Change STATUS to 'refunded'.")
numbered(d, "Save.")
para(d, "The Stripe refund fires automatically for that amount, the customer gets the refund email, and the parent order's status flips to 'refunded' too so you can see the linkage in your orders list.")

# ───────────────────────────────────────────────────────────────────────
h1(d, "7. Customer accounts")

para(
    d,
    "Every customer who places an order gets a free account on the site, automatically. They use it "
    "to track orders, request returns, and access anything else they buy (Reset Room, REGULATED, "
    "future courses). One login for everything.",
)

h3(d, "How they sign in")
numbered(d, "First-time buyer → gets an automatic 'Welcome — set your password' email after their order. They click the link, choose a password, and they are signed in.")
numbered(d, "Returning customer → goes to /login (top-right of the site) and signs in with their email + password.")
numbered(d, "Forgot password → 'Forgot your password?' link on the sign-in page sends them a fresh reset link.")

h3(d, "What they see at /account")
bullet(d, "Their order history — every order ever placed on this email.")
bullet(d, "Tracking links on shipped orders.")
bullet(d, "'Request a return' button on eligible orders.")
bullet(d, "Quick link to the Reset Room dashboard (if they are a member).")
bullet(d, "Quick link to REGULATED course (if they have access).")

callout(
    d,
    "Guest checkout still works:",
    "Customers don't HAVE to log in to buy. They can place an order as a guest. We still create the account silently so when they get the welcome email they can set a password and use the account later. Nothing forced.",
    colour=GREEN,
)

# ───────────────────────────────────────────────────────────────────────
h1(d, "8. Editing your customer emails")

para(
    d,
    "Every email the site sends — order confirmations, shipped, refunded, return-approved, welcome, "
    "password reset — has its subject and body in the Email Template collection in Strapi. You edit "
    "the wording any time without needing Sameer to change code.",
)

h3(d, "Where to find them")
numbered(d, "Strapi → Email Template (in the left sidebar).")
numbered(d, "You will see 10+ rows, one per email. The 'When it fires' field on each row explains when that email gets sent.")

h3(d, "What you can change")
bullet(d, "SUBJECT — what appears in the customer's inbox subject line.")
bullet(d, "PREHEADER — the short preview text under the subject line in their inbox.")
bullet(d, "INTRO — the first paragraph(s) of the email body.")
bullet(d, "OUTRO — the closing paragraph(s) of the email body.")
bullet(d, "CTA LABEL + CTA URL — the button under the outro (e.g. 'Track your parcel').")

h3(d, "What you CANNOT change")
bullet(d, "KEY — the internal name (do not touch — it is tied to the code).")
bullet(d, "AUDIENCE — whether this email goes to the customer or to you.")

h3(d, "Merge tags (auto-fill at send time)")
para(d, "Put any of these in any field and they get filled in with the real value when the email goes out:")
bullet(d, "{{order_number}} — e.g. ALW-A8K3M2P1")
bullet(d, "{{customer_name}} — the customer's name")
bullet(d, "{{first_name}} — just the first name")
bullet(d, "{{total}} — the order total (e.g. 115.00)")
bullet(d, "{{tracking_number}}, {{tracking_url}}, {{shipping_carrier}} — for shipped emails")
bullet(d, "{{cancellation_reason}} — for cancelled emails")
bullet(d, "{{refund_amount}} — for refunded emails")
bullet(d, "{{return_reason}}, {{return_notes}} — for return emails")
bullet(d, "{{set_password_url}}, {{reset_password_url}} — for welcome / password reset")

h3(d, "Disabling an email")
para(
    d,
    "Untick the ENABLED checkbox on any template row. That email will stop being sent. The system "
    "logs it and moves on — nothing breaks. Use this if you would rather handle a specific stage "
    "manually (e.g. you prefer to write 'shipped' emails by hand).",
)

callout(
    d,
    "Test your edits:",
    "Place a test order with your own email, change its status, and watch what arrives in your inbox. Edit the template again until you are happy. You will not break anything — the system always has the previous version saved.",
    colour=GREEN,
)

# ───────────────────────────────────────────────────────────────────────
h1(d, "9. Tracking the money (matching Stripe to Strapi)")

h3(d, "Card payments")
bullet(d, "Every paid Stripe order has STRIPE PAYMENT ID stamped on the Strapi order (starts with 'pi_').")
bullet(d, "Every refunded Stripe order has STRIPE REFUND ID stamped (starts with 're_').")
bullet(d, "You can search Stripe dashboard by these IDs to find the exact charge/refund.")
bullet(d, "Stripe automatically generates an Invoice PDF for every paid order — download from Stripe → Payments → click into the charge → Invoice tab.")

h3(d, "Bank transfers")
bullet(d, "Your bank statement will show 'ALW-XXXXXXXX' as the reference customers used.")
bullet(d, "Match each one to a Strapi order by ORDER NUMBER, then change its status from 'pending' to 'paid'.")
bullet(d, "Once you mark it paid: stock is decremented (if not already), Mailchimp tag fires, customer journey starts.")

# ───────────────────────────────────────────────────────────────────────
h1(d, "10. Quick reference: status meanings")

h3(d, "Order statuses")
bullet(d, "PENDING — bank transfer placed, awaiting your reconciliation.")
bullet(d, "PAID — card charge succeeded OR you marked a bank transfer paid.")
bullet(d, "SHIPPED — you posted the parcel. Customer got the shipping email.")
bullet(d, "COMPLETED — delivered. Customer got the review-request email.")
bullet(d, "CANCELLED — order voided. Customer got the cancellation email.")
bullet(d, "REFUNDED — money returned. Stripe refund fired automatically (or manually for bank transfers).")

h3(d, "Return request statuses")
bullet(d, "REQUESTED — customer just submitted. Action: review and decide.")
bullet(d, "APPROVED — you said yes. Customer got shipping instructions (your notes).")
bullet(d, "RECEIVED — optional: you got the items back. No email.")
bullet(d, "REFUNDED — you processed the refund. Triggers Stripe refund + email.")
bullet(d, "REJECTED — return declined. No automated email — write to the customer manually.")

# ───────────────────────────────────────────────────────────────────────
h1(d, "11. If something goes wrong")

h3(d, "An email didn't send")
para(d, "Check three things, in order:")
numbered(d, "Is the template ENABLED in Strapi Email Template? (Untick by mistake → enable.)")
numbered(d, "Is the right env var set in Coolify? (RESEND_API_KEY for sending; ORDER_EVENT_SECRET for lifecycle triggers.)")
numbered(d, "Email Sameer the order number and what you expected to happen.")

h3(d, "Stripe refund didn't fire")
para(d, "If you change a paid Stripe order to 'refunded' but the customer says no refund showed up:")
numbered(d, "Check the order in Strapi — does STRIPE REFUND ID have a value (re_xxx)? If yes, the refund did fire — wait 5–10 minutes; refunds sometimes take time to appear on the customer's statement.")
numbered(d, "If STRIPE REFUND ID is empty, the Stripe call failed. Check your admin inbox for an admin email with the error.")
numbered(d, "Worst case: refund manually in Stripe dashboard, write to the customer.")

h3(d, "A customer can't log in")
numbered(d, "Ask them to click 'Forgot your password?' on /login — that sends a fresh reset link.")
numbered(d, "If that doesn't work, check the Strapi User collection for their email — does the row exist? If not, they probably never completed a purchase (account creation only happens on first order). Tell them to place a small order to set the account up.")

# ───────────────────────────────────────────────────────────────────────
h1(d, "12. Done.")

para(
    d,
    "That is the whole order lifecycle. Once you have done each step once, the muscle memory takes "
    "over and it is mostly invisible — change a status, the right thing happens.",
    italic=True,
)
para(
    d,
    "If you find yourself wanting to do something the system doesn't already do — email me. Anything "
    "you change repeatedly is a candidate for automation.",
    italic=True,
)

# ───────────────────────────────────────────────────────────────────────
out_path = "Docs/ALW_Order_Playbook_For_Anna.docx"
d.save(out_path)
print(f"Wrote {out_path}")
