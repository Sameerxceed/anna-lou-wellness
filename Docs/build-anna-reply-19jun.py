"""
Build a documented reply to Anna's 19 Jun feedback (Onenote_19-06-2026.docx
and Onenote_email_19-06-2026.docx).

For each item: quote the line from her notes, give a status pill, and an
honest plain-English response with what to do next.

Output: Docs/ALW_Reply_To_Anna_19-06-2026.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x23, 0x1F, 0x20)
PLUM = RGBColor(0x6E, 0x3A, 0x5A)
GREEN = RGBColor(0x2A, 0x7A, 0x5C)
AMBER = RGBColor(0xC4, 0x70, 0x4A)
RED = RGBColor(0xB3, 0x3A, 0x3A)
BLUE = RGBColor(0x3B, 0x5C, 0x8C)
MUTE = RGBColor(0x6E, 0x6A, 0x62)


def style_normal(d):
    s = d.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)
    s.font.color.rgb = DARK


def h1(d, text, colour=PLUM):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = colour


def h2(d, text, colour=AMBER):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = colour


def h3(d, text, colour=PLUM):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = colour


def para(d, text, italic=False, colour=DARK, size=11):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text); r.italic = italic; r.font.size = Pt(size); r.font.color.rgb = colour


def quote(d, text):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"“{text}”")
    r.italic = True; r.font.size = Pt(10.5); r.font.color.rgb = MUTE


def bullet(d, text):
    p = d.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)


def status(d, label, colour):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label); r.bold = True; r.font.color.rgb = colour; r.font.size = Pt(10.5)


# ---------- BUILD ----------

d = Document()
style_normal(d)

# Cover
h1(d, "Reply to your 19 June notes")
para(d, "I went through both your Onenote docs and grouped everything below. Each item has a status pill (DONE / NEEDS RETEST / NEEDS YOU / IN BUILD / PLANNED) and a plain-English reply.", italic=True, colour=MUTE)
para(d, "Quick read: skim PART 1 for the big items, then dip into PART 2 for the item-by-item.", italic=True, colour=MUTE)

# ============================================================================
h1(d, "PART 1 — The big questions first", colour=PLUM)

# ------ BIG 1: Images not rendering
h2(d, "1. Images aren’t showing anywhere")
quote(d, "I tested on the retreat event page, the upsells, and a few others and the image slot just stays blank.")
status(d, "DONE — root cause fixed, needs retest with fresh uploads", AMBER)
para(d, "The underlying bug was that the CMS server didn’t have a persistent volume mounted, so every redeploy wiped /uploads. We fixed this on 8 June (Coolify Volume Mount, verified across a deliberate redeploy). Since then, all NEW uploads stay put.")
para(d, "Why you might still be seeing blanks:")
bullet(d, "Old entries that had images BEFORE the volume crisis lost them. The image field shows a broken-link icon or empty slot. Fix: re-upload to the entry, save, publish. That image will now persist forever.")
bullet(d, "For event cards on /experiences/retreats etc., I rewrote the renderer 16 Jun to LEAD with the event’s hero_image. If you haven’t uploaded a hero_image to the Experience event entry, the card shows an “Add a hero image in CMS” placeholder, not nothing.")
bullet(d, "For upsells, the image is OPTIONAL on each card. If left empty the card renders text-only — that’s by design, not a bug.")
para(d, "Action from you: re-upload the cover image on one event (say, 13th June) + on one upsell card. If those still don’t render, screenshot the broken page + the CMS entry side-by-side and send to me. I’ll dig further.")

# ------ BIG 2: Mobile upload errors
h2(d, "2. Mobile upload errors / bulk upload failing")
quote(d, "Errors uploading images from a mobile. When I add assets in bulk it loads and then fails each time.")
status(d, "PARTIALLY FIXED — needs your exact error screenshots", AMBER)
para(d, "On 8 June I shipped fixes for the most common mobile upload causes: filename sanitisation (spaces / unicode), HEIC → JPEG conversion, sequential processing to avoid memory spikes, and raised the upload size cap.")
para(d, "If you’re still seeing errors:")
bullet(d, "Send me a screenshot of the exact error message you see.")
bullet(d, "Tell me roughly the file size you’re uploading (under 8 MB usually safe; over 15 MB can fail).")
bullet(d, "For bulk: try 3 at a time instead of 10+ — the mobile network can choke on parallel large uploads.")
para(d, "Quick-Photo workaround for mobile: there’s a new sidebar item called Quick Photos — it lists every hero image with a one-tap Replace button. Much faster than navigating the full Content Manager on a phone.")

# ------ BIG 3: Homepage Cup of Jo redesign
h2(d, "3. Homepage Cup of Jo magazine layout (6 categories featured)")
quote(d, "It needs the layout please to be clear like Cup of Jo from the start. Image, category, preview text, Read More. 6 different categories. Home page text to go elsewhere.")
status(d, "PLANNED — waiting on your reply to the 5 questions", BLUE)
para(d, "I sent you 5 questions on 10 June so I can design this properly: (a) full revamp vs mobile-only, (b) which 6 of 22 sub-categories to feature, (c) where the homepage body copy goes if moved off, (d) keep/drop the non-article strips (products / testimonials / press / certs), (e) one big featured above vs jump straight in.")
para(d, "Once I have those answers it’s a 2–3 hour build. Without them I’m guessing.")

# ------ BIG 4: Mobile readability
h2(d, "4. Tiny text on mobile")
quote(d, "Ours has very tiny text on mobile so you can barely read it.")
status(d, "DONE — site-wide font bump", GREEN)
para(d, "12 June I bumped the root font-size across the site (desktop and mobile). Top nav links ~26 percent bigger, body text ~9 percent bigger, footer chunkier. Should be very obvious if you hard-refresh on your phone.")

# ------ BIG 5: Sales pages for editorial pages
h2(d, "5. Pages that need to be SALES pages (currently editorial / magazine)")
quote(d, "Founder Reset (Work & Money) needs to be a proper sales page. The Returning Circle (Community) same, should be a sales page.")
status(d, "NEEDS SCOPE CHOICE FROM YOU", AMBER)
para(d, "I flagged this in the 11 June follow-up too. Both can be sales pages, the question is whether you want them as:")
bullet(d, "OPTION A: Full Page Builder pages styled like the REGULATED sales page (long-form, sections you drag in, image-led). ~3–4 hours to build each. You get pixel control via drag-blocks in CMS.")
bullet(d, "OPTION B: Programme-style pages like The Reset / Signal (a structured template you fill in). ~1 hour each. Faster but less editorial flexibility.")
para(d, "Reply A or B per page and I’ll build them within 48 hours.")

# ------ BIG 6: Reset Stories from Substack
h2(d, "6. Reset Stories not pulling from Substack")
quote(d, "Reset Stories isn’t pulling the blog content from Substack. New blog post I wrote in Work and Money isn’t in the Substack reset letters.")
status(d, "AUTO-PULL IS BUILT — needs your config", AMBER)
para(d, "The hourly Substack RSS pull job IS built and running on the live CMS (cron at :15 past every hour). What I need from you to verify the config:")
bullet(d, "Your exact Substack URL. The job currently points at https://annalouwellness.substack.com/feed.")
bullet(d, "Your Substack section names. The job maps them to Reset Stories categories: “Sunday Cosmic Forecast” → Cosmic Forecast, “Wednesday Signal Check” → Signal vs Noise, “Reset Stories” → Holding Everything, “Contributor Features” → Contributors. If yours differ, posts pull but land in the default “Holding Everything” bucket.")
bullet(d, "Have you published any public (free) posts? Paid-only Substack posts show only a preview in the public RSS so the job marks them is_free=false.")
para(d, "Also worth confirming: the FLOW IS ONE-WAY (Substack → site, NOT site → Substack). The Work & Money blog post you wrote in the CMS doesn’t go to Substack. If you want CMS → Substack too, that’s a separate build — worth discussing whether you write in Substack or in CMS as the primary surface.")

# ------ BIG 7: Take the quiz → Nervous System Decoder + quiz CTA destination
h2(d, "7. “Take the Quiz” navigation + post-quiz CTA")
quote(d, "On Work with Anna “Take the quiz” should be “Nervous System Decoder”. After the quiz it sends people to “Step into the Reset Room” — should point to the Nervous System Decoder instead.")
status(d, "BOTH DONE — nav 11 Jun, post-quiz CTA 19 Jun", GREEN)
para(d, "Nav link is renamed and pointed at /free/nervous-system-decoder. Hard-refresh (Ctrl+Shift+R) on your phone if you're seeing the old label cached.")
para(d, "Post-quiz CTA: the Clear-result fallback was pointing at /community/reset-room with label 'Step inside The Reset Room'. Changed 19 Jun to /free/nervous-system-decoder with label 'Get your free Nervous System Decoder'. Your CMS-set copy (Strapi -> Decoder Quiz Page -> results) still wins if filled — this only affects entries left blank.")

# ------ BIG 8: Practitioner page in nav
h2(d, "8. Can’t find the new Practitioner page in the navigation")
quote(d, "I can’t find the new practitioner page in the navigation. I want to add Maria Alvarado (Reiki Master).")
status(d, "DONE — it’s under About → Practitioners", GREEN)
para(d, "Live at staging.annalouwellness.com/practitioners. Hover About in the top nav, last item in the dropdown.")
para(d, "To add Maria: CMS → Practitioner → + Create → fill Name (Maria Alvarado), Role (Reiki Master), Bio, Portrait, optional Website URL, Instagram handle, Location. Pick Display style = card (or banner for a featured row). Save + Publish.")
para(d, "Also new on the page: a CTA called “Apply to be listed”. Wellness professionals can fill a form (name, email, phone, practice, message) and you get an email + they’re tagged in Mailchimp.")

# ============================================================================
h1(d, "PART 2 — Everything else, item by item", colour=PLUM)

# ============ Onenote_19-06-2026.docx items ============
h2(d, "Onenote_19-06-2026 — staging walk-through items")

h3(d, "Returning Circle — needs 2 locations support")
quote(d, "I do 2 each week at 2 different locations. Can I have space for 2 events? Then add more?")
status(d, "DONE — shipped 19 Jun", GREEN)
para(d, "Built it. The Returning Circle page now renders the recurring sessions as side-by-side cards (one per location/time). Set up: CMS -> Community Event -> The Returning Circle -> Sessions -> +Add an entry per location. Fields per session: day_of_week, time, location_label, optional location_url (Google Maps or Zoom), optional notes.")
para(d, "Cards render as a responsive grid (3-2-1 across desktop/tablet/mobile). Leave empty and the page falls back to a list pointing you at the CMS field. Also removed the hardcoded 'Every Wednesday evening' copy that was stale now you run multiple per week.")

h3(d, "Design preview while editing")
quote(d, "Can you add a preview so I can see whilst I’m designing it so I know the correct layout and colours and fonts?")
status(d, "BUILT — the View live page button", GREEN)
para(d, "Right side of any edit screen there’s a green button labelled “View live page”. Click it after saving and the page opens in a new tab with your edits applied. Within 1–2 seconds of saving.")

h3(d, "Align & Amplify — sales page preview jumps to homepage")
quote(d, "On the align and Amplify event page. You can’t see any preview of the sales page it jumps to the homepage.")
status(d, "NEEDS RETEST — was a populate bug fixed 11 Jun", AMBER)
para(d, "The Strapi query was silently returning zero data (combining populate=* with hierarchical sub-populate is a known v5 silent bug). Fixed 11 June. The /experiences/align-and-amplify page should now render properly. Hard-refresh and confirm.")

h3(d, "Event pages — no add-to-cart, link to product page")
quote(d, "This events page shouldn’t have an add to cart or book a call to find out more but instead it just ends and they have to go to a separate product page.")
status(d, "NEEDS DISCUSSION", BLUE)
para(d, "Currently event cards have two buttons: Read more (→ sales page) and Book directly (→ Calendly if set). I can remove the Book directly button so cards only ‘Read more’ — your call. Confirm.")

h3(d, "Upsells on the product page — services not ‘you may also like’")
quote(d, "There’s no upsells on the product page for other services. ‘You may also like’ is only relevant for jewellery.")
status(d, "DONE — upsells now everywhere", GREEN)
para(d, "12–16 June I added the Upsells field (max 3 cards) to every page singleton + every editable collection. Anna picks Label, Link (searchable dropdown of every site URL), Eyebrow, Blurb, Image. Renders at the bottom of the page as a ‘Where next’ block. Use this for service-to-service upsells.")
para(d, "Shop product detail page still uses ‘You Might Also Like’ (same-category products, automatic) — you mentioned that was fine.")

h3(d, "Auto-SEO — don’t make me enter it on every page")
quote(d, "Can it not automatically fill it in what it thinks? AI should be doing this part.")
status(d, "DONE — auto-fills on Save", GREEN)
para(d, "10 June I shipped auto-SEO on save: every time you Save an article / programme / experience / coaching session / generic page / Page Builder page, Claude reads the entry’s name + body and writes seo_title + seo_description if they’re blank. Your manual edits are NEVER overwritten.")
para(d, "For all the OLD entries that pre-date this: sidebar → SEO & AI Files → the new button at the top called “Run bulk SEO backfill”. One click fills every empty SEO across all entries.")

h3(d, "Date picker — can’t see months ahead")
quote(d, "I was trying to change the date but can’t see the months ahead to be able to choose a date.")
status(d, "DONE — native date picker", GREEN)
para(d, "Replaced the broken Strapi date picker with a native HTML5 one. Works on every browser including Safari iOS. Tap the month/year header to jump months/years. Tap Today or Clear if needed.")

h3(d, "SEO description should be Anna-style example given")
quote(d, "Eg: An exclusive one-day immersion for founders and leaders on a private Thames houseboat.")
status(d, "DONE — your example is the gold standard in the prompt", GREEN)
para(d, "I updated the auto-SEO prompt 10 June using your exact “An exclusive one-day immersion” sentence as the example Claude is given. The system is told to: lead with the topic, use words people search, include 2 of (audience / format / place / outcome), and avoid filler verbs (discover, unlock, transform, embark, journey).")
para(d, "Run the bulk backfill button on the SEO Files page to retroactively rewrite descriptions for old entries with this improved prompt.")

h3(d, "No booking_url on the event — where do I put it?")
quote(d, "And there’s no booking url in there so people can’t purchase. And I wouldn’t know what to put there.")
status(d, "FIELD EXISTS — see help text now added", GREEN)
para(d, "On every Experience event entry the booking_url field has explicit help text under it (added 16 June): “Paste a Calendly link (e.g. https://calendly.com/anna-annalouoflondon/discovery-call) and it opens as a popup. Other URLs (Stripe checkout, internal pages) open normally. Leave blank to hide the button.”")

h3(d, "Folder upload format / image library is messy")
quote(d, "When I upload images to the folder that come out in a very strange format like this so I can’t find the images I’m looking for.")
status(d, "STRAPI MEDIA LIBRARY UX — some helpers added", BLUE)
para(d, "Strapi’s built-in Media Library sorts by upload date by default. Two ways to make it easier:")
bullet(d, "Use Folders. Create folders by topic (Retreats / Workshops / Anna portraits) and upload INTO the folder.")
bullet(d, "Use the Quick Photos sidebar item I shipped — it lists every photo IN-USE on the site as a thumbnail, with one-tap Replace. Bypasses the Media Library entirely.")

h3(d, "Corporate Wellbeing — can’t see design, jumbled text")
quote(d, "Just text jumbled up. I can’t see the design.")
status(d, "DIAGNOSED — needs hero_image upload", AMBER)
para(d, "Curled the live CMS — heroImage is null on that entry. The page falls back to a stock Unsplash image while you fill it in, but layout-wise the hero band collapses without a real image. Upload a landscape JPG (~1600x1000px) to Corporate Wellbeing → hero_image and it’ll snap into the proper layout.")

h3(d, "Content Type Builder locked")
quote(d, "Can’t access previews and buttons on the left on desktop which says Content Type Builder. It is locked.")
status(d, "BY DESIGN — schemas are code-controlled", BLUE)
para(d, "Content Type Builder lets you ADD or REMOVE fields from a collection. That’s code-controlled in production for safety — a mis-click can drop columns and lose data. If you want a new field on any content type, tell me what you want and I’ll add it (~10–20 min per field).")

h3(d, "How do we add a new page?")
status(d, "TWO WAYS", GREEN)
para(d, "OPTION A: For a freeform page like a sales page, use Page Builder. Sidebar → Pages (build your own) → + Create. Drag in section components (Hero, Text Block, Image-Text Split, Buy Programme, Pay-What-You-Feel, FAQ, Card Grid, etc.). URL becomes /p/{slug}. The REGULATED sales page is an example.")
para(d, "OPTION B: For a page that looks like an existing one (e.g. a new programme like another version of The Reset), use Clone. Open the existing programme → three-dot menu top right → Clone → rename. Less work than building from scratch.")

h3(d, "FAQ page error")
quote(d, "FAQ page error.")
status(d, "NEED SCREENSHOT", AMBER)
para(d, "FAQ system is shipped (per-page Q&A across 31 page slugs). Could you screenshot the exact error you’re seeing? I can’t reproduce it from here.")

h3(d, "REGULATED — where to add design for page")
quote(d, "Regulated — where do I add the design for page. I can still just see the fields with text.")
status(d, "PAGE EXISTS — it’s a Page Builder page", BLUE)
para(d, "REGULATED is at sidebar → Pages (build your own) → entry with slug ‘regulated’. Open it and you’ll see all the section components stacked (numbered list, anchor band, pay-what-you-feel, etc.). Edit each section individually. Image fields are on each section component that supports one (Hero, Image-Text Split, Full-bleed Image, Image Pair, Image with Caption).")
para(d, "Hint: in each section, the field labels (now with help text added 16 June) tell you the visual position the image will appear in. Example: “HERO IMAGE — large image on the RIGHT side”.")

h3(d, "Workshops menu → Contact page (broken link)")
quote(d, "The Workshops menu link is taking people to the Contact page instead of Workshops.")
status(d, "FIX IN NAV — ~2 min", AMBER)
para(d, "Navigation → Experiences (or wherever Workshops sits) → children → find Workshops → set href to /experiences/workshops → Save + Publish. Or send me which top-level menu it’s under and I’ll do it.")

h3(d, "Reset Letters logo — should be graphic not typed")
status(d, "BLOCKED ON YOU — send me the logo asset", AMBER)
para(d, "I’m waiting on the Reset Letters wordmark / logo graphic. Once you send it I’ll swap the typed-text block for the image (~15 minutes).")

h3(d, "Work & Money submenu doesn’t match")
status(d, "NEEDS CLARIFICATION", AMBER)
para(d, "Could you tell me exactly what the menu should be? Current children under Work & Money. Send me the desired list of sub-items and which URLs each should point to (or pick from the Site URLs sidebar lookup).")

h3(d, "Chris Corsini design reference")
status(d, "BLOCKED ON YOU — send screenshot", AMBER)
para(d, "Waiting on the screenshot of his workshop page so I can match the style for experience / workshop pages.")

h3(d, "CMS page search — can’t find Align and Amplify")
status(d, "DIAGNOSED + CLARIFIED", GREEN)
para(d, "Align & Amplify lives in the collection now renamed ‘Experiences · Event Bookings’ (previously ‘Experiences · Event’). The other collection ‘Experiences · Category Pages’ (previously ‘Experiences · Sub-page’) is for the 4 main category landings (Retreats / Workshops / Corporate / Speaking). Each collection’s description now tells you what’s in it so you can’t mix them up.")
para(d, "Also: use the Site URLs sidebar item for a master list of every page on the site with search + Copy buttons.")

h3(d, "On your side: image uploads, tagging, Houseboat Life tags, Mailchimp, ChatGPT shopping")
status(d, "NOTED", BLUE)
para(d, "Acknowledged. Tell me when you’re ready on any of those or if you’d like help.")

# ============ Onenote_email_19-06-2026.docx items ============
h2(d, "Onenote_email_19-06-2026 — follow-on items")

h3(d, "One Day Intensive email — pull date + format from Calendly")
quote(d, "On the Email for the one day intensive, it’s got to pull the date and the format from callendly.")
status(d, "DONE — shipped 19 Jun (needs ~10 min Calendly+Mailchimp setup)", GREEN)
para(d, "Built. New webhook at /api/calendly/webhook receives every booking from Calendly, extracts event date + time + name + location, and stamps them onto the Mailchimp contact as merge fields EVENT_DATE / EVENT_TIME / EVENT_NAME / EVENT_LOC. The webhook also tags the contact based on the Calendly event slug (one-day-intensive -> 'One Day Booked', discovery-call -> 'Discovery Booked', etc.).")
para(d, "Your setup work (one-time, when you're back):")
bullet(d, "Calendly -> Account -> Integrations -> Webhooks -> Create. URL https://annalouwellness.com/api/calendly/webhook. Subscribe to invitee.created. Copy the signing key.")
bullet(d, "Coolify web app env vars -> add CALENDLY_WEBHOOK_SECRET with that signing key.")
bullet(d, "Mailchimp -> Audience -> Settings -> Audience fields and |MERGE| tags. Add 4 text merge fields: EVENT_DATE, EVENT_TIME, EVENT_NAME, EVENT_LOC.")
bullet(d, "Mailchimp -> the One Day welcome email body -> reference |EVENT_DATE| |EVENT_TIME| |EVENT_NAME| |EVENT_LOC| wherever you want the details to render per-recipient.")
para(d, "Works for every Calendly-booked service, not just One Day. Each event slug fires a different tag so you can build per-service journeys.")

h3(d, "Substack subscribers → Mailchimp tag")
quote(d, "We need to find out who came to the website from Substack so they get signed up to our welcome flow.")
status(d, "DONE — UTM path shipped 19 Jun", GREEN)
para(d, "Built the UTM-aware path. When a visitor lands on any page with ?utm_source=substack in the URL, the site captures that in sessionStorage. Any signup form they then submit (Reset Letters, Decoder, enquiry forms) automatically tags the contact with a SECOND tag 'Origin: substack' alongside the main signup tag.")
para(d, "Your setup work (one-time): append ?utm_source=substack to every Substack share link you post. That's the only manual step. Then in Mailchimp build a Customer Journey triggered by the 'Origin: substack' tag — that's your Substack-arrival welcome flow.")
para(d, "Same pattern works for ?utm_source=instagram, ?utm_source=podcast, anything else. The tag is always 'Origin: <whatever>' (lowercased).")

h3(d, "Where is the welcome email on Mailchimp?")
status(d, "DOC ALREADY SENT", GREEN)
para(d, "I emailed you ALW_Mailchimp_Journey_Entry_Points.docx on 12 June — it lists every journey, the URL on staging that triggers it, and the trigger tag. The welcome flow for Reset Letters is the ‘Reset Letters — Founding Welcome’ journey in Mailchimp → Automations → Customer Journeys.")
para(d, "Also: I uploaded the 23 ALW email templates to your Mailchimp account on 25 May (Email templates → Saved tab). Each journey can switch from the placeholder ‘ALW Master’ template to its dedicated one (ALW - 01 - Newsletter welcome, etc.) and you edit copy in your voice.")

h3(d, "Banner image on W&M blog post not showing in list")
quote(d, "I edited the banner image for the Work and Money blog post and it’s not showing up in the list of blog posts.")
status(d, "PINNED ORDER ISSUE", AMBER)
para(d, "The Articles list is sorted by is_homepage_pinned + homepage_pin_order + publishedAt. Tick is_featured or is_homepage_pinned on the new entry, set homepage_pin_order to 1 (lower = higher), and save. It’ll jump to the top.")
para(d, "If you wanted JUST the new entry to appear in the section’s article list and not the homepage, that’s a different question — send a screenshot of where you expected to see it.")

h3(d, "AI answers — remove em-dashes")
quote(d, "All of the AI answers ideally needs to remove the EM DASHES.")
status(d, "DONE — shipped 19 Jun", GREEN)
para(d, "Added 'NEVER use em-dashes or en-dashes; use commas, full stops, or hyphens (-) instead' to all 4 AI prompts (auto-SEO on save, manual Generate SEO, bulk SEO backfill script, AskAnna chatbot). New AI output is clean from now on.")
para(d, "For OLD SEO already written with em-dashes: open the SEO & AI Files sidebar -> click 'Run bulk SEO backfill'. It re-runs through every entry, blank-only by default so it never touches anything you've edited manually. (If you want it to rewrite all SEO including filled ones, tell me and I'll add a force-overwrite toggle.)")

h3(d, "Programme covers — image-led like Chris Corsini")
quote(d, "How do I add covers onto the programs that it images like not as much text?")
status(d, "NEEDS BUILD — ~2 hrs (gated on Corsini screenshot)", BLUE)
para(d, "Once I have the Chris Corsini reference, I can build an image-led programme card / page template that uses a single big cover image instead of the structured intro band. Will swap the Programme list view too. Will revise the Programme content type to have a cover_image field if it doesn’t already.")

h3(d, "Drag-reorder programmes")
quote(d, "Can I drag them around to move them into place?")
status(d, "USE sort_order FIELD", GREEN)
para(d, "Open each Programme entry, edit the sort_order number. Lower = appears first. Use 10, 20, 30 (not 1, 2, 3) so you can drop new entries between existing ones without renumbering. Strapi v5 doesn’t support drag-reorder in the list view yet but sort_order achieves the same outcome.")

h3(d, "Regulation programme first, low-price first")
status(d, "SET SORT_ORDER", GREEN)
para(d, "Donation-only Regulation programme: sort_order = 10. The Reset: 20. Signal: 30. Signal & Build: 40. One Day: 50. Signal Collective: 60. Save each. They’ll order automatically.")

h3(d, "Signal link goes to wrong service")
status(d, "NEEDS YOUR SCREENSHOT", AMBER)
para(d, "Which Signal link — from the nav, from the homepage, or from another page? Send me the page URL + the screenshot of where it lands wrongly. I’ll fix in minutes.")

h3(d, "Testimonials — text + video, per-page reviews")
quote(d, "I want to edit the text and add videos like the page I showed you on Shoshana Raven. I have videos and text written reviews I want to add to relevant pages for each service.")
status(d, "DONE — testimonials are CMS-driven with video support", GREEN)
para(d, "Sidebar → Testimonial / Review collection. Each entry can have text AND a video URL (YouTube / Vimeo / direct mp4). To attach a review to a service:")
bullet(d, "Tick the relevant Programme / Experience in the entry’s Programme or Experience field.")
bullet(d, "The matching service page will show that review in its ReviewsSection.")
para(d, "Plus the testimonials page (/testimonials) lists every review in a grid — ones with video play inline, ones without show as quote cards.")

h3(d, "SEO preview per page")
quote(d, "I can’t see what SEO needs to go in here as there’s no preview showing me on each page so I don’t know what it’s picking up and using.")
status(d, "AUTO-FILL ON SAVE + REFRESH SHOWS IT", GREEN)
para(d, "When you Save, the seo_title and seo_description fields fill in automatically (if blank). The Refresh button you clicked on the right side of the edit view triggers the page to re-read and show the filled values. After Save → wait 5–10 seconds → click Refresh → the SEO fields populate with what Claude wrote.")

h3(d, "What does the Refresh button do?")
status(d, "EXPLAINED — same as above", GREEN)
para(d, "It re-reads the entry from the database and shows you the SEO that auto-SEO wrote. Use it after Save when you want to confirm what got generated.")

h3(d, "Can’t access testimonial space to add pic + text")
quote(d, "I can’t see how to access this testimonial space to be able to put a picture and text on it.")
status(d, "DIRECTION", GREEN)
para(d, "Sidebar → Testimonial / Review → + Create new entry. Fill: Reviewer name, Reviewer title, Quote, optional Photo, optional Video URL, optional Rating (1–5). Save + Publish. The review surfaces on the testimonials page AND (if you tick the Programme / Experience relation) on the matching service page.")

h3(d, "Trauma-informed Abuse recovery coaching page")
quote(d, "I can’t seem to find this page. https://annalouwellness.com/somatic-trauma-informed-narcisstic-abuse-coaching/")
status(d, "NEEDS BUILD OR REDIRECT", AMBER)
para(d, "The current Squarespace site has this page; the new staging site doesn’t have a direct equivalent yet. Two options:")
bullet(d, "Build it as a Page Builder page at the same URL (/somatic-trauma-informed-narcisstic-abuse-coaching) — ~2 hours, keeps your existing inbound links working.")
bullet(d, "Make it the Recovery Coaching programme page (which already exists at /the-work/recovery) and set a 301 redirect from the old URL. Cleaner long-term but you’d update marketing links.")
para(d, "Confirm preference.")

h3(d, "‘Holding everything / strong mind / signal vs noise’ headings confusing")
status(d, "DONE — shipped 19 Jun", GREEN)
para(d, "Removed from cards on the section landing pages (Reset Stories / Life / Love & Relationships / Work & Money) and from the kicker on individual article detail pages. The page itself already establishes the section, so the sub-category was redundant.")
para(d, "Homepage card badges still appear but now show the SECTION name (Reset Stories / Life / Love & Relationships / Work & Money) rather than the sub-category — useful on the homepage where cards span sections.")

h3(d, "Empty section help text")
quote(d, "I don’t know what is meant to go in these sections.")
status(d, "PARTIALLY DONE — field help text added 16 June", AMBER)
para(d, "I added plain-English help text to 72 fields across 38 collections on 16 June. Every common field (seo_title, slug, hero_image, kicker, booking_url, mailchimpTag, etc.) now has an example below it. If there are SPECIFIC fields you’re still confused by, screenshot them and I’ll add help text + an example.")

h3(d, "Substack write-back from CMS")
quote(d, "I wrote a new blog post in the Work and money section of the website but it hasn’t appeared in the reset letters Substack. All of the reset stories are meant to be over on the Substack reset letters.")
status(d, "FLOW IS ONE-WAY — needs decision", BLUE)
para(d, "Currently the flow is Substack → site, not the other way. Three options:")
bullet(d, "OPTION A: Keep one-way. You write in Substack (primary), site auto-pulls.")
bullet(d, "OPTION B: Reverse it. You write in CMS (primary), and we build a CMS → Substack publisher. ~4 hours. Requires a Substack API token you generate.")
bullet(d, "OPTION C: Both ways. Choose which surface you wrote on first and the other auto-mirrors. Most complex, can cause sync conflicts. ~6 hours.")
para(d, "OPTION A is the cleanest — what I assumed from the original plan. Confirm which you want.")

h3(d, "Reset Room sales page — couldn’t add link in article body")
quote(d, "I went to add a link to the Reset Room sales page here in the final paragraph but when I clicked the add.")
status(d, "NEEDS SCREENSHOT", AMBER)
para(d, "Could you screenshot the article editor with the link-adding panel open? The richtext editor in Strapi supports links (highlight text → link icon in toolbar → paste URL). If that’s broken send me what you saw.")

h3(d, "Relation field for category — show all blog categories")
quote(d, "In the add a relation section it should show all the blog posts categories so I can decide which section it needs to go in to be featured.")
status(d, "DROPDOWN SHOULD ALREADY LIST THEM", AMBER)
para(d, "On any Article entry, the Category field is a dropdown of all Article Categories. It should auto-search by typing. If it’s empty for you, screenshot — sometimes Strapi v5 hides the dropdown until you click into the field.")

# Section: summary checklist
h2(d, "Summary — what I need from you this week")
para(d, "Reply on any of these so I can unblock + ship:")
bullet(d, "Homepage 6-category grid — reply to my 5 questions (10 June).")
bullet(d, "Sales pages — Option A or B for Founder Reset and Returning Circle.")
bullet(d, "Substack — your URL, your section names, and which direction the sync should run (A / B / C).")
bullet(d, "Work & Money submenu — the list you want with each URL.")
bullet(d, "Chris Corsini screenshot — for the design reference.")
bullet(d, "Reset Letters logo asset — to swap the typed text.")
bullet(d, "Signal link goes to wrong service — send screenshot + URL.")
bullet(d, "Image upload errors — send a screenshot of the exact error.")
bullet(d, "Trauma-informed Abuse Recovery — build new page or redirect to existing Recovery Coaching?")
bullet(d, "Trauma-informed Abuse Recovery — build new page or 301 redirect to existing Recovery Coaching?")

para(d, "", colour=DARK)
para(d, "Generated " + __import__('datetime').datetime.now().strftime('%d %b %Y'), italic=True, colour=MUTE, size=9)

out = "Docs/ALW_Reply_To_Anna_19-06-2026.docx"
d.save(out)
print(f"Wrote {out}")
