"""
Reply doc for Anna's 14 July 2026 feedback (2 documents combined).
Item-by-item status with DONE / NEEDS YOU labels.
"""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor

doc = Document()

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.size = Pt(20)

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.size = Pt(14)

def para(text, italic=False, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold

def bullet(text, bold=False):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.bold = bold

def blank():
    doc.add_paragraph()

def status(label, colour):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.font.size = Pt(10)
    r.bold = True
    r.font.color.rgb = colour
    return p


DONE = RGBColor(0x2E, 0x7D, 0x32)     # green
NEEDS = RGBColor(0xE6, 0x5C, 0x00)    # orange
CMS = RGBColor(0x1A, 0x52, 0x76)      # blue


# Cover
h1('Your Two Docs — What Shipped, What You Do Next')
para('Anna Lou Wellness · 14 July 2026 review', italic=True, size=11)
para('Sameer', italic=True, size=10)
blank()

para(
    'This covers both documents you sent — the Website Fixes list and the '
    'Homepage & Footer Review. Everything below is grouped by whether it '
    'is DONE (I shipped it), CMS (you can edit yourself in Strapi), or '
    'NEEDS YOU (an action outside the website — Mailchimp, Calendly, etc.).'
)
blank()

# ============ DONE ============
h1('Shipped in code (already live after deploy)')

h2('1. Contact form was completely broken')
para(
    'You reported the Contact form "becomes unresponsive after completing the fields". '
    'Root cause: the form was static HTML with no submit handler — the Send Message button did nothing. '
    'Replaced with a working form that posts to your Mailchimp with a "Contact Enquiry" tag.'
)

h2('2. Preview vs Live discrepancy')
para(
    'Pay buttons and forms breaking in CMS Preview but working on the Live site: '
    'this is the CMS preview iframe. Stripe and CAPTCHA refuse to run inside iframes for security — '
    'nothing to do with your code. Now when you view a page in CMS Preview, interactive elements are '
    'replaced with a friendly note: "Preview mode: this action is disabled. Click Open live to test the real button."'
)
para('So going forward: use Open live (top-right in Strapi) to test payments and form submits.', italic=True)

h2('3. Cart icon missing from main website navigation')
para(
    'You were right — the cart icon was hidden below 1340px screen width by a CSS rule. Now visible '
    'across all screen sizes, always. It stays as a compact icon on mobile so nav space stays clean.'
)

h2('4. "4 ways to begin" section — 5 items, all Learn More linking to same page')
para(
    'Each Learn More button now goes to that programme\'s own page. Section title is still editable in CMS '
    '(Content Manager -> Work with Anna Page -> programmesTitle) — if you have 5 programmes now, change '
    '"Four ways to begin" to "Five ways to begin".'
)

h2('5. Retreats — remove Register Interest button')
para('Done. Retreats page now shows only the actual retreats.')

h2('6. Events Calendar — remove Enquire About Events link')
para('Done.')

h2('7. Workshops — remove Enquiry button')
para(
    'Checked the code — there is no Enquiry button on /experiences/workshops currently, '
    'only the standard experience cards. Might you have been looking at /experiences (the landing '
    'grid) filtered by Workshop? Let me know if you still see one and I will hunt it down.'
)

h2('8. Returning Circle — remove enquiry form, add Book Now')
para(
    'Done. The RSVP form is gone. In its place is a Book Now button that opens whatever URL you '
    'set as CTA URL on Community Event -> The Returning Circle in CMS. Paste your Calendly link '
    'there and it will open as a popup so people stay on your site.'
)

h2('9. Press page — dedicated Press Enquiry form')
para(
    'Done. /about/press now has its own form (Your name, Publication, Email, Deadline, What are you working on). '
    'No more redirect to Contact. Posts get tagged in Mailchimp as "Press Enquiry".'
)

h2('10. Work With Me page — dedicated form')
para(
    'Done. /about/partnerships now has its own form (Your name, Organisation, Email, Website, Collaboration). '
    'No more redirect to Contact. Tagged in Mailchimp as "Partnership Enquiry".'
)

h2('11. Discovery Call rename -> 15-minute 1 to 1 chat')
para(
    'Renamed every user-facing occurrence on the website — programme pages, quiz results, chat quick '
    'questions, FAQ answers, Ask Anna page, CMS field descriptions and default copy. 17 files updated.'
)
para('Website is DONE. You still need to update Mailchimp + Calendly to match — see NEEDS YOU section below.', italic=True)

h2('12. REGULATED — pay-what-you-can dropdown')
para(
    'Done. New feature: opt-in per programme via 3 new CMS fields (see CMS section below). When set, '
    'the sales page shows a dropdown of allowed amounts + a Pay button. Server-side enforces that '
    'the price submitted matches one of your allowed amounts, so no one can pay £1 by tampering.'
)

h2('13. REGULATED membership flow (pay -> become member -> access)')
para(
    'This was already fully wired. Payment through the REGULATED programme entry auto-grants access. '
    'Members see a REGULATED card in their /community/reset-room/dashboard and can visit '
    '/the-work/regulated/access to view course modules.'
)

h2('14. Corporate + Speaking "No upcoming items" empty state')
para(
    'From the Homepage/Footer review. When someone filters the /experiences grid by Corporate or Speaking '
    'and there are no scheduled events, the page used to say a generic "check back soon" message. '
    'Now for Corporate it points at /experiences/corporate-wellbeing for private bookings, and Speaking '
    'points at /experiences/speaking to book Anna. Retreats + Workshops keep the mailing-list nudge.'
)

# ============ CMS ============
doc.add_page_break()
h1('CMS — quick things you can do yourself in Strapi')

h2('A. Reset Session £0 for testing')
para(
    'Content Manager -> Coaching Session -> "1:1 Reset Session" (or whichever session you want free '
    'for testing) -> set price to 0 -> Save + Publish.'
)

h2('B. Remove £10 payment gate from Contact page temporarily')
para(
    'Content Manager -> Contact Page -> clear the "discovery_calendly_url" field -> Save + Publish. '
    'The whole 15-minute 1 to 1 chat block disappears from the page. Refill the URL when you want it back.'
)

h2('C. REGULATED pay-what-you-can amounts')
para(
    'Content Manager -> Work · Programme -> REGULATED entry. Set:'
)
bullet('pwycOptions -- example: "50, 100, 200, 500" (comma-separated pounds)')
bullet('pwycDefault -- example: 100 (pre-selected amount)')
bullet('pwycLabel -- example: "Choose your amount" (shown above the dropdown)')
para(
    'Save + Publish. The dropdown appears on /the-work/regulated. To turn OFF pay-what-you-can and '
    'go back to a fixed price, clear pwycOptions.'
)

h2('D. Returning Circle booking URL')
para(
    'Content Manager -> Community Event -> The Returning Circle -> "CTA URL" -> paste your Calendly '
    'link for the Circle (or Stripe URL if you sell tickets). "CTA Label" defaults to "Join the Circle" '
    '— change to "Book now" if you prefer.'
)

h2('E. Programme booking URLs')
para(
    'For each programme (Reset, Signal, Signal & Build, One Day, Signal Collective, Recovery): '
    'Content Manager -> Work · Programme -> [programme] -> "ctaLabel" + "ctaUrl". Paste the specific '
    'Calendly URL for that programme so buyers can book directly. Right now they all default to /contact.'
)

h2('F. Nav consistency (Homepage vs Footer)')
para(
    'You mentioned the Footer under "About" only shows Contact / Press / Work With Me, but the Homepage '
    'nav under About shows Anna\'s Story + Press + Work With Me + Practitioners. Both are CMS-editable:'
)
bullet('Content Manager -> Navigation (top nav) -> About -> add/remove children as needed.')
bullet('Content Manager -> Footer -> Connect Links -> add "Anna\'s Story" + "Practitioners" to match.')
para(
    'The Footer under "Work with Anna" currently says "What do you need?" and "Ask Anna for a recommendation" — '
    'that is coming from the Footer content type Connect Links or Explore Links. You can rewrite those to match '
    'whatever the Homepage shows.'
)

h2('G. Homepage: category sections that show "No upcoming items"')
para(
    'Motherhood and Home & Space sections have no articles yet. Content Manager -> Story · Article -> '
    '+ Create new entry -> pick the right category. Or leave those categories empty and I can hide the '
    'empty sections from the homepage — let me know.'
)

h2('H. Returning Circle spec gaps')
para(
    'From the Homepage/Footer review, RC is missing: Tuesday day label, The Hare and the Moon venue, '
    'Twickenham location, membership info, FAQ. Content Manager -> Community Event -> The Returning Circle:'
)
bullet('For each location/session -> "Sessions" repeatable component -> +Add entry with day, time, location.')
bullet('Membership info -> add a section under intro if you want it visible.')
bullet('FAQ -> Content Manager -> FAQ -> + Create with page = "the-returning-circle".')

h2('I. About page navigation vs Side nav inconsistency')
para(
    'You said the About page in main nav shows different content than About page in the side navigation. '
    'Both point to /about — same page. Check if you\'re looking at the Homepage nav dropdown (Anna\'s Story, Press, '
    'Work With Me, Practitioners) vs mobile menu (same list). If they still look different, screenshot to me — '
    'might be a caching issue.'
)

h2('J. Community page — display only blog posts')
para(
    'Fixes doc item 11. Two ways to read this — which do you want?'
)
bullet(
    'OPTION 1 (keep 4-section landing): current /community landing shows Returning Circle + Reset Room + Events + Resources. If you want to also add a Blog Posts strip pulling from /reset-stories, tell me and I add it.'
)
bullet(
    'OPTION 2 (strip it back): remove the RC + Reset Room + Events + Resources sections and show ONLY blog posts. That is a bigger change and I want to confirm before I do it — those 4 sections currently drive a lot of Reset Room signups.'
)
para('Reply with 1 or 2. Meanwhile, everything above is CMS-editable via Content Manager -> Community Page.')

h2('K. Resource Library — only free resources')
para(
    'Fixes doc item 11 (second half). Content Manager -> Resource Library entries have an "is_free" toggle. '
    'Right now some entries may be marked paid. Filter to those that are paid, either delete them (if they were seeded '
    'placeholders) or move the paid resource content into the Reset Room dashboard instead. I can help audit — '
    'send me the list of resources you want to keep and I will clean up the rest.'
)

h2('L. Practitioner in homepage navigation')
para(
    'Fixes doc item 16. Practitioners IS in the nav today, but only inside the About dropdown — not visible unless you hover About. '
    'Two options:'
)
bullet(
    'OPTION A (recommended, no work): keep it in About dropdown but I add it to Footer -> Connect Links so it always appears. Anna does this in CMS -> Footer -> Connect Links -> +Add link (label: Practitioners, href: /practitioners).'
)
bullet(
    'OPTION B (top-level nav): I add "Practitioners" as its own top-level nav item next to Shop/Community/About. Makes the nav 10 items instead of 9 — might feel crowded on mid-size laptops.'
)
para('Reply which option. Or leave as-is if you are ok with it being nested under About.')


# ============ NEEDS YOU ============
doc.add_page_break()
h1('NEEDS YOU — outside the website')

h2('1. Rename Discovery Call in Calendly')
para(
    'Log into Calendly -> Event Types -> find your existing Discovery Call event -> Edit -> '
    'change the name to "15-minute 1 to 1 chat" (or your preferred wording). '
    'Also update the URL slug so the link makes sense (e.g. calendly.com/anna/15-min-chat). '
    'Then paste the new Calendly URL into the CMS Contact Page discovery_calendly_url field.'
)

h2('2. Rename Discovery Call tags in Mailchimp')
para(
    'Mailchimp -> Audience -> Manage tags. Any tag currently named "Discovery Booked" or similar '
    '-> rename to "1 to 1 Chat Booked". After you rename, tell me and I will update the Calendly '
    'webhook so the new tag fires for future bookings.'
)

h2('3. Decoder Quiz Mailchimp Customer Journeys')
para(
    'The quiz already sends the right tags to Mailchimp on every submission:'
)
bullet('Every quiz completer: "Decoder Subscriber"')
bullet('Signal Clear result: "Decoder · Signal Clear"')
bullet('Signal Scrambled result: "Decoder · Signal Scrambled"')
bullet('Signal Faint result: "Decoder · Signal Faint"')
para('What you need to do in Mailchimp:', bold=True)
bullet('Automations -> Customer Journeys -> Create.')
bullet('Trigger: "Tag added" -> pick each of the 4 tags above.')
bullet('For "Decoder Subscriber" -> send your welcome email with the free PDF.')
bullet(
    'For each "Signal ___" tag -> tailored follow-up. Signal Scrambled + Signal Faint '
    'people should get the REGULATED upsell email; Signal Clear gets a meditation download instead.'
)
para(
    'When your journeys are live, tags will fire automatically for every quiz completion. '
    'Nothing else to wire on the code side.', italic=True
)

h2('4. Confirm form recipient email addresses')
para(
    'The following forms tag people in Mailchimp when submitted. Confirm which email address should '
    'ALSO receive a notification email (currently only Practitioner Enquiry does):'
)
bullet('Corporate Enquiry')
bullet('Speaking Enquiry')
bullet('Press Enquiry (new)')
bullet('Partnership Enquiry (new)')
bullet('Contact form (general)')
para(
    'Reply with which of these should send you an email (and if the address is different from hello@annalouwellness.com, '
    'tell me the address).'
)

h2('5. Blog posts with broken images')
para(
    'From the Homepage/Footer review — some blog posts have broken/missing hero images. Anna, the '
    'safest thing is for you to audit article by article: Content Manager -> Story · Article -> '
    'sort by "Updated" -> open each -> if hero_image is blank or shows an error, re-upload from the '
    'Media Library. If a whole batch of images is broken, screenshot me one so I can diagnose.'
)


blank()
para('That is the full list. WhatsApp me when you have hit each of the CMS + NEEDS YOU items so I can verify.', italic=True)


# ============ Content policy change ============
doc.add_page_break()
h1('One important change to how the website reads content')

para(
    "You told me you were seeing hardcoded copy on pages that you thought "
    "should be coming from the CMS — and once you edited part of a section "
    "in Strapi, the rest of that section was still showing default text you "
    "could not edit. That is confusing and it is fixed."
)

h2('New rule: what you see on the live site is what is in the CMS')

para(
    "Every page now reads its content strictly from Strapi. There is no "
    "hardcoded copy hiding underneath. Three consequences:"
)
bullet(
    "IF you have filled a section in the CMS -> that content shows on live. "
    "Edit it, save, and the live site updates in a few seconds."
)
bullet(
    "IF you clear a field in the CMS -> that field disappears from the live "
    "site. The rest of the page adjusts around the gap."
)
bullet(
    "IF you clear a whole section (all its fields blank) -> the entire "
    "section is hidden from the live site. The page becomes shorter."
)

h2('What we did so you do not lose anything')

para(
    "Before switching the site to this stricter mode, we copied every piece "
    "of content that was previously hardcoded INTO the CMS. So on the day "
    "we ship this change, the live site looks identical to today, but every "
    "piece of copy is now editable in your admin panel. Specifically:"
)
bullet("About page — kicker, title, roles tagline, both story paragraphs, additional bio")
bullet("Community page — full header + Returning Circle + Reset Room + Events + Resources copy")
bullet("Work with Anna hub — header, Ways-to-work-with-me section, Programmes section")
bullet("Corporate Wellbeing + Speaking sub-pages — eyebrow + tagline in the hero band")
bullet("All 6 programme pages (Reset, Signal, Signal & Build, One Day, Signal Collective, Recovery) were already fully in CMS — no change there")
bullet("Retreats, Workshops, Events Calendar, Resource Library, Press, Work With Me — already in CMS, unchanged")

h2('What this means for you day-to-day')

para(
    "You can now confidently edit any section without wondering whether "
    "some hidden default is going to leak through. If you delete a section "
    "you do not want any more, it goes away entirely. If you fill it, it "
    "appears. Anna in control, no surprises."
)
para(
    "Only three things stay as safety-net fallbacks (they never render as "
    "text, only as visual scaffolding):",
    italic=True,
)
bullet("Hero background images -- if you have not uploaded one, a stock photo shows instead of a blank white box")
bullet("Programme accent colours -- the CSS needs a colour to be set for the layout to work")
bullet("The site nav (top of every page) -- this stays coded because if it disappeared, visitors could not navigate")

blank()
para('Anna x', italic=True)
para('Sameer', italic=True, size=10)

out = 'Docs/ALW_Reply_To_Anna_14-07-2026_v2.docx'
doc.save(out)
print(f'Wrote: {out}')
