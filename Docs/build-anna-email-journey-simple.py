"""
Anna-friendly version of the email journey map.
Same information as ALW_Email_Journey_Map.docx but rewritten without
technical language — no endpoint paths, no code, no jargon.
Reads like a conversation, not an audit.
"""

from docx import Document
from docx.shared import Pt

doc = Document()

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11.5)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.size = Pt(22)


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.size = Pt(16)


def h3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.size = Pt(13)


def para(text, italic=False, bold=False, size=11.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold


def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(11.5)


def blank():
    doc.add_paragraph()


# ── Cover ──────────────────────────────────────────────
h1('Your emails, explained')
para('What happens when someone signs up, books, or buys on your site.',
     italic=True, size=13)
para('For Anna. Written 10 July 2026.', italic=True, size=10)
blank()

para(
    "I've written a technical version of this for developers. This one is "
    "for you. It's a friendly walkthrough of what actually happens when "
    "someone does something on your site, so you can picture the journey a "
    "real person goes on."
)
blank()
para('Prefer to just ASK?', bold=True)
para(
    "Everything in this document is also loaded into your in-CMS chatbot. "
    "Click the little Help - Ask bubble at the bottom-right of any admin "
    "page (or open the Help - Ask page from the sidebar) and type things "
    "like \"what emails does someone get when they sign up for Reset "
    "Letters?\" or \"who edits the order receipt?\" or \"if I refund an "
    "order, what happens to the emails?\". It will answer in plain "
    "English, just like this document -- and it knows your specific setup, "
    "so it can also tell you which Mailchimp tag or CMS field is behind "
    "each email."
)


# ── Two systems ─────────────────────────────────────────
h2('Two things send emails from your site')

para('Mailchimp', bold=True)
para(
    'This is where your welcome sequences and nurture emails live. Any time '
    'someone signs up for Reset Letters, takes the Decoder quiz, or books a '
    'call, your site quietly tells Mailchimp what happened. Mailchimp then '
    'runs whichever Customer Journey you have set up for that action.'
)
para(
    "You write and edit these emails INSIDE Mailchimp. The site never "
    "touches the actual wording. If you want to change what your welcome "
    "email says, you log into Mailchimp and edit it there."
)
blank()

para('Resend', bold=True)
para(
    "Resend handles the boring-but-essential emails. Order confirmations, "
    "receipts, refund notifications, password reset links, and the emails "
    "that go to YOU when someone new orders something. These are the ones "
    "that have to arrive within seconds because someone is waiting for them."
)
para(
    "You edit these emails INSIDE the CMS, in the Email Template collection. "
    "Change the wording there, save, and the next email sent uses your new "
    "copy — no code change needed."
)


# ── Journeys ────────────────────────────────────────────
doc.add_page_break()
h1('What happens when someone...')


h2('...signs up for Reset Letters')
para(
    'You get a new subscriber in two places: Mailchimp (for your welcome '
    'sequence) and Substack (so they start receiving the newsletter itself).'
)
para('The emails they get:', bold=True)
bullet(
    'Your welcome sequence from Mailchimp. If they signed up BEFORE launch day '
    '(22 June 2026) they get the Founding Members welcome. After that, they '
    'get the Standard Subscribers welcome.'
)
bullet(
    'The Substack newsletter itself, whenever you publish a new letter.'
)
para(
    'You get nothing sent to you personally by the site — you just see them '
    'appear in your Mailchimp audience and Substack subscriber list.',
    italic=True,
)


h2('...takes the Nervous System Decoder quiz')
para(
    'They enter their email, get tagged as a Decoder Subscriber in Mailchimp, '
    'and start your Decoder journey.'
)
para('The emails they get:', bold=True)
bullet('Your welcome email with the free Decoder PDF.')
bullet(
    'Two days later, an upsell email inviting them to REGULATED (the paid '
    'programme). If they buy, they exit this sequence — you never sell them '
    'something they already own.'
)
bullet(
    'Four days later, one final gentle nudge to REGULATED, then the sequence '
    'ends. No more emails unless they take another action.'
)


h2('...fills in any enquiry form (Contact, Practitioner, Corporate, Speaking, etc.)')
para(
    'They get tagged in Mailchimp with a tag that matches the form — for '
    'example, someone filling in the Practitioner form gets the '
    '"Practitioner Enquiry" tag.'
)
para('The emails they get:', bold=True)
bullet(
    'Whichever journey you have set up in Mailchimp for that tag. Every form '
    'has its own tag, so you can send a different welcome depending on why '
    'they got in touch.'
)
para('The emails YOU get:', bold=True)
bullet(
    'Only for the Practitioner enquiry form — you get a notification email so '
    'you can review the application. For the other forms, you just see the '
    'new contact appear in Mailchimp with the right tag.'
)


h2('...buys something from the shop (paying by card)')
para(
    'This is the busiest flow because it kicks off several things at once.'
)
para('The emails they get:', bold=True)
bullet(
    'If this is their first ever purchase, a "welcome, set your password" '
    'email so they can log in later and see their orders.'
)
bullet(
    'Their order receipt with the order number and what they bought.'
)
bullet(
    'If the item is part of a paid programme like REGULATED, they also enter '
    'that programme\'s follow-up sequence in Mailchimp — course prompts, '
    'check-ins, etc.'
)
para('The emails YOU get:', bold=True)
bullet(
    'A "new order" notification so you know a sale happened. Every order, '
    'every time.'
)
para(
    'Everyone who buys anything also gets tagged as a "Shop Customer" in '
    'Mailchimp, so you can build a "past customers" journey if you want to.',
    italic=True,
)


h2('...buys something from the shop (paying by bank transfer)')
para(
    'They get a slightly different flow because you have to manually mark '
    'the order paid once the transfer lands.'
)
para('The emails they get right away:', bold=True)
bullet('The "welcome, set your password" email if it is their first order.')
bullet(
    'A bank transfer instructions email with your bank details and a reference '
    'number so you can match the payment.'
)
para('The emails they get once you mark the order paid in the CMS:', bold=True)
bullet(
    'The normal order receipt. From that point on, this becomes the same as a '
    'card order — Shop Customer tag, programme journey if applicable.'
)


h2('...you process a refund')
para(
    'You do this in the CMS when they attend their Discovery Call, cancel, '
    'or you just decide to refund.'
)
para('The emails they get:', bold=True)
bullet(
    'A refund confirmation with the amount and how long it takes to reach '
    'their card (usually 5-10 business days).'
)
para('The emails YOU get:', bold=True)
bullet(
    'The same "new order" notification with an updated status, so you have '
    'a paper trail.'
)
para(
    'If they had entered a programme-specific journey (like the REGULATED '
    'follow-up), we quietly remove that tag so they stop getting the course '
    'emails. Their "Shop Customer" tag stays — they were, after all, a '
    'customer at one point.',
    italic=True,
)


h2('...ships or completes an order')
para(
    'When you mark an order as shipped or completed in the CMS, the customer '
    'gets an email you can fully control.'
)
para('The emails they get:', bold=True)
bullet('"Your order has shipped" with the tracking link you added.')
bullet('"Your order arrived" when you mark it complete.')
bullet(
    'Return-related emails if a customer requests a return, if you approve '
    'it, receive the item, or reject it — one email per stage.'
)
para(
    'Every single one of these emails is editable by you in the CMS, in the '
    'Email Template collection. Change the tone, add a personal note, '
    'shorten them — whatever you want.',
    italic=True,
)


h2('...books a call on your Calendly')
para(
    'Discovery Call, One Day, Reset Session, Signal Scoping — any of them.'
)
para('The emails they get:', bold=True)
bullet(
    'The standard Calendly booking confirmation with the date, time, and any '
    'Zoom link (this comes from Calendly, not from your site).'
)
bullet(
    'Whichever pre-call sequence you set up in Mailchimp for that specific '
    'call type. You can reference the exact date and time of their booking '
    'inside the email — so it feels personal even though it is automated.'
)


h2('...forgets their password')
para(
    'They click "forgot password" on the login screen, enter their email, and '
    'we send a one-time reset link. No tags, no journey — just a link to reset '
    'their password.'
)


# ── Anna's job list ─────────────────────────────────────
doc.add_page_break()
h1('What you actually need to do')

para(
    'Everything the SITE does is automatic. What is left is what YOU need to '
    'set up in Mailchimp so the tags actually cause emails to send.'
)
blank()

h2('In Mailchimp')

para('Customer Journeys', bold=True)
para(
    'For each of the tags mentioned above, you need a Customer Journey '
    'configured with that tag as the trigger. If you have already built some, '
    'this is a matter of activating each one. If not, they are draft copies '
    'waiting for you to write the emails and click Start.'
)
para(
    'Log into Mailchimp -> Automations -> Customer Journeys. Open each one, '
    'draft the emails using Mailchimp\'s merge tags like *|FNAME|* for the '
    'first name, and click Start when you are happy.'
)
blank()

para('Merge fields for booking emails', bold=True)
para(
    'For the pre-call sequences after someone books on Calendly, you can use '
    'merge tags like |EVENT_DATE|, |EVENT_TIME|, |EVENT_NAME|, |EVENT_LOC| '
    'in the email body — so the email says "your Reset Session is on 15 '
    'August at 3pm" without you writing that manually.'
)
para(
    'For those to work, go to Audience -> Settings -> Audience fields and '
    'add four TEXT fields with tag names EVENT_DATE, EVENT_TIME, EVENT_NAME, '
    'EVENT_LOC. Exact spelling, all caps.'
)
blank()

h2('In the CMS')

para('Email Template collection', bold=True)
para(
    'Content Manager -> Email Template. You will see rows for order '
    'confirmations, refunds, shipped notifications, return updates, and '
    'a few others. Open each one, read the current copy, and adjust it to '
    'sound like you. These emails send automatically when the trigger '
    'happens — you never send them by hand.'
)
para(
    'Any changes save immediately. The next order that goes through uses '
    'your new wording.'
)
blank()

para(
    'That is genuinely the whole list. Everything else is already wired.',
    italic=True,
)


# ── FAQs ────────────────────────────────────────────────
doc.add_page_break()
h1('Questions you might have')

h3('If I edit an email in Mailchimp, when does the change go live?')
para(
    'Immediately, for everyone who enters the journey after you save. Anyone '
    'already IN the journey continues on the version they started on. Rule of '
    'thumb: your edits affect NEW subscribers, not existing ones mid-sequence.'
)

h3('If I edit an email in the CMS Email Template, when does the change go live?')
para(
    'The very next email that goes out. There is no queue and no delay.'
)

h3('Can I stop someone from receiving emails without deleting them?')
para(
    'Yes. In Mailchimp, unsubscribe them from the audience or remove the '
    'tag. In the CMS, transactional emails only go when the trigger fires, '
    'so if you cancel their order, no refund email sends.'
)

h3('Can I see if someone opened / clicked an email?')
para(
    'For Mailchimp emails, yes — Mailchimp tracks opens and clicks on every '
    'campaign and journey email. Log in and check Reports.'
)
para(
    'For Resend transactional emails (order receipts etc.), Resend tracks '
    'delivery status and bounces but not opens by default.'
)

h3('What happens if my Anthropic bill or Mailchimp bill runs out?')
para(
    'Mailchimp: if you exceed your monthly send quota, sends pause but no '
    'contacts are lost. Journey emails queue up and resume when you upgrade.'
)
para(
    'Anthropic: only affects the auto-generated SEO and the in-CMS chatbot. '
    'Zero impact on customer-facing emails.'
)

h3('Who do I ask if something breaks?')
para(
    'WhatsApp Sameer with a screenshot. He can trace exactly which email '
    'was supposed to send and why it did or did not.'
)

blank()
para('Anna x', italic=True)

out = 'Docs/ALW_Email_Journey_For_Anna.docx'
doc.save(out)
print(f'Wrote: {out}')
