"""Convert ANNA_CMS_GUIDE.md to a clean .docx for Anna."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r"e:\Xceed\Code\Clients\anna-lou-wellness\Docs\ANNA_CMS_GUIDE.md"
DST = r"e:\Xceed\Code\Clients\anna-lou-wellness\Docs\ANNA_CMS_GUIDE.docx"

PLUM = RGBColor(0x6E, 0x3A, 0x5A)
DARK = RGBColor(0x23, 0x1F, 0x20)
GREY = RGBColor(0x55, 0x55, 0x55)


def strip_emoji(s: str) -> str:
    # Drop leading emoji + space from headings (Word doesn't render the colour glyphs cleanly)
    return re.sub(r"^[\U0001F300-\U0001FAFF☀-➿]+\s*", "", s).strip()


def add_inline_runs(paragraph, text: str, base_size=11, base_color=DARK):
    """Handle **bold**, `code`, and plain text inline."""
    # Tokenise by **bold**, `code`, and the rest
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        run.font.size = Pt(base_size)
        run.font.color.rgb = base_color
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = PLUM
        else:
            run.text = part


def main():
    with open(SRC, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Horizontal rule
        if line.strip() == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("―" * 30)
            run.font.color.rgb = GREY
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            text = strip_emoji(line[2:])
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = PLUM
            i += 1
            continue
        if line.startswith("## "):
            text = strip_emoji(line[3:])
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = PLUM
            i += 1
            continue
        if line.startswith("### "):
            text = strip_emoji(line[4:])
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = DARK
            i += 1
            continue

        # Tables (markdown)
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
            # Collect table rows
            header_cells = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2  # skip header + separator
            body_rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                body_rows.append(cells)
                i += 1
            table = doc.add_table(rows=1 + len(body_rows), cols=len(header_cells))
            table.style = "Light Grid Accent 1"
            # Header
            for idx, cell_text in enumerate(header_cells):
                cell = table.rows[0].cells[idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = PLUM
            # Body
            for r_idx, row in enumerate(body_rows, start=1):
                for c_idx, cell_text in enumerate(row):
                    if c_idx >= len(table.rows[r_idx].cells):
                        continue
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_inline_runs(p, cell_text)
            # spacing after table
            doc.add_paragraph()
            continue

        # Bullet list
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            add_inline_runs(p, line[2:])
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\d+)\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2)
            add_inline_runs(p, m.group(2))
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_inline_runs(p, line)
        i += 1

    doc.save(DST)
    print(f"Wrote {DST}")


if __name__ == "__main__":
    main()
