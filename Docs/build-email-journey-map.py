"""
Build ALW_Email_Journey_Map.docx — a complete, machine-parseable audit
of every email a visitor/customer can receive, mapped by trigger.

Sections:
  1. Overview + how to read the doc
  2. Trigger index (every action -> tags + emails)
  3. Mailchimp tag catalogue (every tag + which journey listens)
  4. Transactional email catalogue (every email-template + when it fires)
  5. Anna's activation checklist
  6. Data flow diagrams (ASCII)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Base styles ─────────────────────────────────────────────
styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.size = Pt(20)
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.size = Pt(15)
    return p


def h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.size = Pt(13)
    return p


def para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p


def numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    return p


def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)


# ── Cover ──────────────────────────────────────────────────
h1('Email Journey Map')
para('Anna Lou Wellness site + Mailchimp + Resend transactional email',
     italic=True, size=12)
para('Version 1.0  ·  Generated 8 July 2026  ·  For Anna Lou Scaife', italic=True, size=10)
doc.add_paragraph()

# ── Section 1: Overview ────────────────────────────────────
h1('1. Overview')

para(
    'This document maps every email that can be sent from the Anna Lou '
    'Wellness site. Two systems send emails:'
)
bullet(
    'Mailchimp — for marketing journeys triggered by tags (e.g. "you signed up '
    'for Reset Letters, get the welcome sequence"). Anna edits these inside '
    'Mailchimp; the site only fires the tag.'
)
bullet(
    'Resend — for transactional emails triggered by specific events '
    '(e.g. "your order shipped"). Anna edits these inside the CMS in the '
    'Email Template collection; the site pulls the template at send time.'
)

h2('How to read this document')
numbered(
    'Section 2 (Trigger Index) tells you what happens the moment a visitor '
    'takes any given action. Start here if you are asking "if someone does X, '
    'what emails do they get?"'
)
numbered(
    'Section 3 (Mailchimp Tag Catalogue) is the source of truth for every '
    'tag the site can apply. Each tag maps to a specific Mailchimp Customer '
    'Journey Anna configures. Use this when editing journeys in Mailchimp.'
)
numbered(
    'Section 4 (Transactional Email Catalogue) lists every Resend email by '
    'template key. Use this when editing the wording of a specific email in '
    'the CMS Email Template collection.'
)
numbered(
    'Section 5 (Activation Checklist) is the pre-launch punch list — what '
    'Anna needs to configure in Mailchimp before the site is fully live.'
)

# ── Section 2: Trigger Index ───────────────────────────────
doc.add_page_break()
h1('2. Trigger Index')
para(
    'Every action a visitor can take on the site, and exactly what happens next.'
)


# ---- 2.1 Reset Letters signup ----
h2('2.1  Reset Letters signup')
para('Trigger: visitor submits the form on /reset-letters', bold=True)
para('Endpoint: POST /api/subscribe-reset-letters')

para('What happens, in order:')
numbered('Turnstile CAPTCHA verified (blocks bots).')
numbered(
    'Email is upserted into Mailchimp audience with merge field FNAME = first name.'
)
numbered(
    'Mailchimp tag applied — "Founding Members" if the submission is BEFORE '
    'the RESET_LETTERS_LAUNCH_DATE env var (currently 2026-06-22), otherwise '
    '"Standard Subscribers".'
)
numbered(
    'Email is ALSO subscribed to the Substack list (annalouwellness.substack.com) '
    'in parallel so they start receiving the newsletter itself immediately.'
)
numbered(
    'If the visitor arrived via a UTM link (e.g. ?utm_source=substack), a SECOND '
    'tag "Origin: <source>" is applied. This lets Anna send a different welcome '
    'to visitors who came from Substack vs Instagram vs a podcast vs elsewhere.'
)
numbered(
    'Success response sent to the browser — form shows the confirmation state.'
)

para('Journeys this triggers in Mailchimp:', bold=True)
bullet('Founding Members welcome sequence (if pre-launch)')
bullet('Standard Subscribers welcome sequence (if post-launch)')
bullet('Optional: fork by Origin tag (e.g. Substack-arrival journey)')

para('Resend emails: NONE.', italic=True)


# ---- 2.2 Nervous System Decoder signup ----
h2('2.2  Nervous System Decoder signup')
para('Trigger: visitor submits the email form on /free/nervous-system-decoder',
     bold=True)
para('Endpoint: POST /api/lead/decoder')
para('Also fires from: the email gate on /free/nervous-system-decoder/quiz')

para('What happens, in order:')
numbered('Turnstile CAPTCHA verified.')
numbered('Email is upserted into Mailchimp with FNAME merge field.')
numbered('Mailchimp tag applied — "Decoder Subscriber".')
numbered(
    'If UTM present, second tag "Origin: <source>" applied same as Reset Letters.'
)
numbered(
    'Redirect / success shown to the visitor (URL depends on which entry point).'
)

para('Journeys this triggers in Mailchimp:', bold=True)
bullet(
    'Decoder Subscriber welcome — delivers the free Nervous System Decoder PDF'
)
bullet(
    'Decoder Upsell (day 2) — offers REGULATED at the paid tier. Exits if the '
    'REGULATED Buyer tag is applied (visitor purchased).'
)
bullet(
    'Decoder Upsell Final (day 4) — last nudge to REGULATED, then journey ends.'
)

para('Resend emails: NONE.', italic=True)


# ---- 2.3 Enquiry forms ----
h2('2.3  Enquiry forms (Contact, Practitioner apply, Corporate, Speaking, etc.)')
para('Trigger: visitor submits any enquiry form on the site', bold=True)
para('Endpoint: POST /api/lead/[type]')

para(
    'The [type] slug in the URL controls the Mailchimp tag. Known types:'
)
table(
    ['Form type (URL slug)', 'Tag applied', 'Also sends Anna an email?'],
    [
        ['returning-circle', 'Returning Circle Enquiry', 'No'],
        ['signal-collective', 'Signal Collective Enquiry', 'No'],
        ['recovery', 'Recovery Coaching Enquiry', 'No'],
        ['one-day', 'One Day Enquiry', 'No'],
        ['speaking', 'Speaking Enquiry', 'No'],
        ['corporate', 'Corporate Wellbeing Enquiry', 'No'],
        ['practitioner-enquiry', 'Practitioner Enquiry',
         'Yes — admin_practitioner_enquiry template'],
        ['anything else',
         '<Capitalised words from slug> + " Enquiry"',
         'No'],
    ]
)

para('What happens, in order:')
numbered('Turnstile CAPTCHA verified (except decoder, which has its own route).')
numbered('Email upserted into Mailchimp with FNAME.')
numbered('Tag applied per the table above.')
numbered('If Origin UTM present, second Origin tag applied.')
numbered(
    'If the form type is in the admin-email list (currently only '
    'practitioner-enquiry), a Resend email fires to Anna via the '
    'email-template collection.'
)

para('Journeys this can trigger in Mailchimp:', bold=True)
bullet('Practitioner welcome / next-steps sequence')
bullet('Corporate Wellbeing enquiry follow-up')
bullet('Speaking enquiry follow-up')
bullet('Returning Circle enquiry welcome')
bullet('Any custom journey Anna wires to the tag')


# ---- 2.4 Shop purchase (card) ----
h2('2.4  Shop purchase (Stripe checkout)')
para('Trigger: visitor completes checkout with a card on /checkout', bold=True)
para('Endpoint: Stripe webhook -> POST /api/stripe/webhook (checkout.session.completed)')

para('What happens, in order:')
numbered('Stripe webhook signature verified.')
numbered('Order status flipped from "pending" to "paid" in Strapi.')
numbered(
    'If the customer is new, Strapi creates a User record + sets a one-time '
    'password token, then Resend fires the shop_account_invite email '
    '("welcome, set your password").'
)
numbered(
    'Mailchimp tag "Shop Customers" applied to the email.'
)
numbered(
    'If the item bought has a mailchimpTag on its Programme / Product entry '
    '(e.g. REGULATED programme has mailchimpTag = "REGULATED Buyer"), that '
    'tag is ALSO applied.'
)
numbered(
    'Resend fires order_paid to the customer with the receipt.'
)
numbered(
    'Resend fires admin_new_order to Anna so she knows a sale happened.'
)

para('Journeys this triggers in Mailchimp:', bold=True)
bullet('Shop welcome / thank-you journey (any journey listening for "Shop Customers")')
bullet(
    'Programme-specific onboarding — e.g. REGULATED Follow-up journey triggered '
    'by the REGULATED Buyer tag. This journey ALSO exits the Decoder Upsell '
    'journey (Mailchimp exit condition) so the customer stops getting sold '
    'something they just bought.'
)

para('Resend emails sent:', bold=True)
bullet('shop_account_invite (new customers only)')
bullet('order_paid (every purchase)')
bullet('admin_new_order (every purchase — Anna gets it)')


# ---- 2.5 Shop purchase (bank transfer) ----
h2('2.5  Shop purchase (bank transfer, no card)')
para('Trigger: visitor picks bank transfer at checkout', bold=True)
para('Endpoint: POST /api/checkout')

para('What happens, in order:')
numbered('Order created in Strapi with status "pending_bank_transfer".')
numbered(
    'If new customer, shop_account_invite email fires.'
)
numbered(
    'Resend fires order_bank_transfer to the customer with bank details + '
    'reference number.'
)
numbered(
    'Resend fires admin_new_order to Anna.'
)
numbered(
    'NO Mailchimp tags applied yet — those wait until Anna marks the order paid '
    'in the CMS (which fires the paid webhook flow).'
)

para(
    'Once Anna manually marks the order as paid in Strapi, the full paid-flow '
    'from section 2.4 runs — tags applied, order_paid email sent.',
    italic=True,
)


# ---- 2.6 Refund ----
h2('2.6  Refund')
para('Trigger: Anna clicks Refund on an order in the CMS', bold=True)
para('Endpoint: POST /api/order-refund')

para('What happens, in order:')
numbered('Stripe refund API called with the amount.')
numbered('Order status flipped to "refunded" in Strapi.')
numbered(
    'If the item had a mailchimpTag, that tag is REMOVED (so the customer '
    'stops receiving that programme\'s follow-up journey).'
)
numbered(
    'Resend fires order_refunded to the customer with the refund amount + '
    'expected settlement time.'
)
numbered(
    'Resend fires admin_new_order (repurposed as a change notification) to '
    'Anna as the audit trail.'
)

para('Journeys affected in Mailchimp:', bold=True)
bullet('Programme-specific journey (e.g. REGULATED Follow-up) — customer removed.')
bullet('Shop Customers tag is KEPT — refund does not un-mark them as a customer.')


# ---- 2.7 Order status change (shipped, completed, etc.) ----
h2('2.7  Order status change (shipped / completed / return)')
para('Trigger: Anna changes an order status in the CMS', bold=True)
para('Endpoint: POST /api/order-event')

para(
    'This is the flexible email switchboard. The request body carries a '
    'template_key, which points at an Email Template row in the CMS.'
)

para('Known template keys used here:')
table(
    ['template_key', 'When Anna triggers it', 'Who receives'],
    [
        ['order_shipped', 'On dispatch, adds tracking URL', 'Customer'],
        ['order_completed', 'When delivery is confirmed', 'Customer'],
        ['return_requested', 'When customer requests a return', 'Customer'],
        ['return_approved', 'When Anna approves the return', 'Customer'],
        ['return_received', 'When Anna marks return as received', 'Customer'],
        ['return_rejected', 'When Anna rejects a return', 'Customer'],
        ['admin_return_requested', 'When customer submits return request', 'Anna'],
    ]
)

para(
    'Anna edits the actual wording of each of these in the CMS Email Template '
    'collection. She does not have to touch any code to change the tone.',
    italic=True,
)


# ---- 2.8 Calendly booking ----
h2('2.8  Calendly booking (Discovery, One Day, Reset Session, etc.)')
para('Trigger: visitor books a call via Anna\'s Calendly links', bold=True)
para('Endpoint: Calendly webhook -> POST /api/calendly/webhook')

para('What happens, in order:')
numbered('Calendly webhook signature verified (HMAC-SHA256).')
numbered(
    'Event type slug from the Calendly URL is looked up in EVENT_TAG_MAP.'
)
numbered('Email upserted in Mailchimp with FNAME.')
numbered(
    'Mailchimp tag applied per the map below.'
)
numbered(
    'Mailchimp merge fields set on this member: EVENT_DATE, EVENT_TIME, '
    'EVENT_NAME, EVENT_LOC. These let Anna reference "you booked {{EVENT_NAME}} '
    'on |EVENT_DATE| at |EVENT_TIME|" inside her Mailchimp email bodies.'
)

para('Calendly event slug -> Mailchimp tag map:')
table(
    ['Calendly URL slug', 'Mailchimp tag applied'],
    [
        ['one-day-intensive OR one-day', 'One Day Booked'],
        ['discovery-call OR discovery', 'Discovery Booked'],
        ['signal-scoping', 'Signal Scoping Booked'],
        ['reset-session', 'Reset Session Booked'],
        ['anything else', 'Calendly Booked (fallback)'],
    ]
)

para('Journeys this triggers in Mailchimp:', bold=True)
bullet(
    'Booking-specific pre-call sequence (e.g. "your Discovery Call is on '
    '|EVENT_DATE| — here is how to prepare")'
)
bullet(
    'Post-call follow-up sequence — Anna wires a "call happened" step in the '
    'journey.'
)

para('Resend emails: NONE (Calendly sends its own booking confirmation).', italic=True)


# ---- 2.9 Password reset ----
h2('2.9  Password reset (forgot password link)')
para('Trigger: visitor clicks "forgot password" on /login', bold=True)
para('Endpoint: POST /api/auth/forgot-password')

para('What happens, in order:')
numbered('Email address validated + rate-limited.')
numbered('One-time reset token generated + stored on the user in Strapi.')
numbered('Resend fires password_reset email with the reset link.')

para('Mailchimp: no tag applied — this is a security event, not a marketing event.',
     italic=True)


# ── Section 3: Mailchimp Tag Catalogue ──────────────────────
doc.add_page_break()
h1('3. Mailchimp Tag Catalogue')
para(
    'Every tag the site can apply, and what journey it should trigger. This '
    'is the source of truth Anna maintains in Mailchimp.'
)

table(
    ['Tag name', 'Fires when', 'Journey Anna wires in Mailchimp'],
    [
        ['Founding Members',
         'Reset Letters signup before RESET_LETTERS_LAUNCH_DATE',
         'Founding Members welcome sequence'],
        ['Standard Subscribers',
         'Reset Letters signup after launch date OR after 500-subscriber cap',
         'Standard Subscribers welcome sequence'],
        ['Decoder Subscriber',
         'Nervous System Decoder signup on landing or quiz gate',
         'Decoder welcome + PDF delivery'],
        ['REGULATED Buyer',
         'REGULATED programme purchased (Stripe webhook, mailchimpTag on entry)',
         'REGULATED Follow-up — exits Decoder Upsell journey'],
        ['Shop Customers',
         'Any completed shop purchase (Stripe webhook)',
         'Shop welcome / thank-you'],
        ['Practitioner Enquiry',
         'Practitioner enquiry form submitted',
         'Practitioner review + application follow-up'],
        ['Returning Circle Enquiry',
         'Returning Circle enquiry form',
         'Returning Circle next steps'],
        ['Signal Collective Enquiry',
         'Signal Collective enquiry form',
         'Signal Collective next steps'],
        ['Recovery Coaching Enquiry',
         'Recovery Coaching enquiry form',
         'Recovery Coaching intake sequence'],
        ['One Day Enquiry',
         'One Day enquiry form',
         'One Day pre-booking nurture'],
        ['Speaking Enquiry',
         'Speaking enquiry form',
         'Speaking follow-up'],
        ['Corporate Wellbeing Enquiry',
         'Corporate Wellbeing enquiry form',
         'Corporate proposal + follow-up'],
        ['One Day Booked',
         'One Day booked in Calendly',
         'One Day pre-call sequence (uses EVENT_DATE merge field)'],
        ['Discovery Booked',
         'Discovery call booked in Calendly',
         'Discovery pre-call preparation email'],
        ['Signal Scoping Booked',
         'Signal Scoping call booked in Calendly',
         'Signal Scoping pre-call sequence'],
        ['Reset Session Booked',
         'Reset Session booked in Calendly',
         'Reset Session pre-call sequence'],
        ['Calendly Booked',
         'Any Calendly booking without a specific tag mapping',
         'Generic Calendly pre-call sequence (fallback)'],
        ['Origin: <source>',
         'Second tag on ANY signup when the visitor arrived via ?utm_source=',
         'Optional segment / different welcome per source (Substack, Instagram, podcast, etc.)'],
    ],
    col_widths=[1.8, 2.6, 2.6],
)


# ── Section 4: Transactional Email Catalogue ────────────────
doc.add_page_break()
h1('4. Transactional Email Catalogue')
para(
    'Every email sent by Resend. Each maps to a row in the Email Template '
    'collection in the CMS — Anna edits subject/preheader/intro/outro/CTA there '
    'and the changes go live in seconds. She does NOT edit any code to change '
    'the copy.'
)

table(
    ['Template key (CMS row)', 'Sent when', 'Sent to'],
    [
        ['shop_account_invite',
         'A new customer completes their first purchase',
         'Customer'],
        ['order_paid',
         'Stripe confirms payment received (card or bank transfer marked paid)',
         'Customer'],
        ['order_bank_transfer',
         'Customer selects bank transfer at checkout',
         'Customer'],
        ['order_shipped',
         'Anna dispatches the order and enters tracking',
         'Customer'],
        ['order_completed',
         'Order marked complete / delivered',
         'Customer'],
        ['order_refunded',
         'Anna processes a refund on an order',
         'Customer'],
        ['return_requested',
         'Customer submits a return request',
         'Customer'],
        ['return_approved',
         'Anna approves the return',
         'Customer'],
        ['return_received',
         'Anna marks the returned item as received',
         'Customer'],
        ['return_rejected',
         'Anna rejects the return',
         'Customer'],
        ['password_reset',
         'Customer clicks Forgot Password on /login',
         'Customer'],
        ['admin_new_order',
         'Any new order (card + bank transfer + refund)',
         'Anna (owner)'],
        ['admin_return_requested',
         'Customer submits a return request',
         'Anna (owner)'],
        ['admin_practitioner_enquiry',
         'A wellness practitioner applies via the /practitioners form',
         'Anna (owner)'],
    ],
    col_widths=[2.2, 3.4, 1.2],
)

para('Available merge tokens inside every template:', bold=True)
code_block([
    '{{order_number}}    - e.g. ALW-3F9K2P',
    '{{customer_name}}   - the customer first name',
    '{{total}}           - order total in GBP',
    '{{tracking_number}} - carrier tracking number (order_shipped only)',
    '{{tracking_url}}    - clickable tracking link (order_shipped only)',
    '{{refund_amount}}   - amount refunded in GBP (order_refunded only)',
    '{{cancellation_reason}} - if provided',
    '{{site_url}}        - annalouwellness.com',
])


# ── Section 5: Activation Checklist ─────────────────────────
doc.add_page_break()
h1('5. Anna\'s Activation Checklist')
para(
    'Everything the SITE will do automatically is done. What is left is what '
    'Anna needs to activate INSIDE Mailchimp so the tags actually cause emails '
    'to send. Rough order for pre-launch:'
)

h2('5.1  In Mailchimp — Customer Journeys')
para(
    'For each tag in Section 3, Anna needs a Customer Journey configured with '
    'that tag as the trigger. If she has already built the shells, this is a '
    'matter of activating each one.'
)
numbered('Log in to Mailchimp -> Automations -> Customer Journeys.')
numbered(
    'For each row in Section 3, open the corresponding journey. If it does not '
    'exist yet, create it. Trigger: "When a contact is tagged with X".'
)
numbered(
    'Draft the emails inside the journey. Use merge tokens like *|FNAME|* '
    '(Mailchimp syntax, not the {{ }} used in Resend transactional templates).'
)
numbered(
    'For journeys that reference Calendly booking events, use the merge fields '
    '|EVENT_DATE|, |EVENT_TIME|, |EVENT_NAME|, |EVENT_LOC|.'
)
numbered('Set exit conditions where relevant (e.g. Decoder Upsell exits on REGULATED Buyer).')
numbered('Click "Start journey" — Mailchimp begins listening for the tag.')

h2('5.2  In Mailchimp — merge fields (one-time setup)')
para(
    'Before the Calendly webhook journeys will work with merge fields, Anna '
    'must create the merge fields in her Mailchimp audience settings:'
)
numbered('Audience -> Settings -> Audience fields and *|MERGE|* tags.')
numbered(
    'Add TEXT merge fields with tag names: EVENT_DATE, EVENT_TIME, EVENT_NAME, '
    'EVENT_LOC. Match the tag names exactly — case-sensitive.'
)

h2('5.3  In the CMS — Email Template copy')
para(
    'Every transactional template in Section 4 has a row in the Email Template '
    'collection. Some rows may still have placeholder copy from setup. Anna '
    'should proofread each one and adjust tone / wording to match her voice.'
)
numbered('Content Manager -> Email Template.')
numbered(
    'Open each row (14 in total per Section 4). Verify subject / preheader / '
    'intro / outro / CTA read the way Anna wants.'
)
numbered('Save and Publish — takes effect immediately, no deploy needed.')

h2('5.4  In Coolify — environment variables')
para('Confirm these are set on the Next.js frontend service in Coolify:')
code_block([
    'MAILCHIMP_API_KEY          - required, from Mailchimp -> Profile -> Extras -> API keys',
    'MAILCHIMP_LIST_ID          - required, from Audience -> Settings',
    'MAILCHIMP_TAG_FOUNDING     - optional (default "Founding Members")',
    'MAILCHIMP_TAG_STANDARD     - optional (default "Standard Subscribers")',
    'RESET_LETTERS_LAUNCH_DATE  - optional (default 2026-06-22)',
    'SUBSTACK_PUBLICATION_URL   - optional (default annalouwellness.substack.com)',
    'CALENDLY_WEBHOOK_SECRET    - required for the Calendly integration',
    'RESEND_API_KEY             - required for every transactional email',
    'EMAIL_FROM                 - required (default "onboarding@resend.dev" for dev — CHANGE for prod)',
    'STRIPE_SECRET_KEY          - required for shop + refund flow',
    'STRIPE_WEBHOOK_SECRET      - required for the /api/stripe/webhook signature check',
])

h2('5.5  In Calendly — webhook (one-time setup)')
numbered('Calendly -> Integrations -> Webhooks -> Create webhook.')
numbered(
    'URL: https://annalouwellness.com/api/calendly/webhook (or the staging URL).'
)
numbered('Subscribe to: invitee.created.')
numbered(
    'Copy the signing key that Calendly generates -> paste into Coolify as '
    'CALENDLY_WEBHOOK_SECRET on the frontend service.'
)


# ── Section 6: Data flow diagrams ───────────────────────────
doc.add_page_break()
h1('6. Data Flow Diagrams (ASCII)')
para('Quick visual reference for each of the main journeys.')

h2('6.1  Reset Letters')
code_block([
    '    Visitor fills form on /reset-letters',
    '                |',
    '                v',
    '    POST /api/subscribe-reset-letters',
    '                |',
    '     +----------+----------+',
    '     |                     |',
    '     v                     v',
    '  Mailchimp             Substack',
    '  (upsert + tag)        (upsert)',
    '     |',
    '     v',
    '  Founding Members  OR  Standard Subscribers  tag',
    '     |',
    '     v',
    '  Mailchimp Customer Journey ── welcome sequence',
])

h2('6.2  Nervous System Decoder + REGULATED upsell')
code_block([
    '    Visitor submits Decoder email on /free/nervous-system-decoder',
    '                |',
    '                v',
    '    POST /api/lead/decoder',
    '                |',
    '                v',
    '    Mailchimp tag: Decoder Subscriber',
    '                |',
    '                v',
    '    Journey: Decoder welcome + PDF delivery',
    '                |',
    '                v (day 2)',
    '    Journey: Decoder Upsell to REGULATED',
    '                |',
    '           +----+----+',
    '           |         |',
    '     purchases   ignores',
    '           |         |',
    '           v         v (day 4)',
    '    REGULATED  Final upsell nudge',
    '    Buyer tag       |',
    '           |         v',
    '           |    Journey ends',
    '           v',
    '    EXITS Decoder Upsell journey',
    '           |',
    '           v',
    '    Journey: REGULATED Follow-up (onboarding + course prompts)',
])

h2('6.3  Shop purchase')
code_block([
    '    Customer completes card checkout',
    '                |',
    '                v',
    '    Stripe -> POST /api/stripe/webhook',
    '                |',
    '     +----------+-------------+---------------+',
    '     |          |             |               |',
    '     v          v             v               v',
    '  Strapi:    Mailchimp:   Resend:         Resend:',
    '  order      Shop         order_paid      admin_new_order',
    '  paid       Customers    (to customer)   (to Anna)',
    '             + programme',
    '             tag if any',
    '     |',
    '     v (new customer only)',
    '  Resend: shop_account_invite (set password link)',
])

h2('6.4  Calendly booking')
code_block([
    '    Visitor books via Calendly link',
    '                |',
    '                v',
    '    Calendly -> POST /api/calendly/webhook',
    '                |',
    '                v',
    '    Look up event slug in EVENT_TAG_MAP',
    '                |',
    '                v',
    '    Mailchimp: tag applied + merge fields set',
    '    (EVENT_DATE, EVENT_TIME, EVENT_NAME, EVENT_LOC)',
    '                |',
    '                v',
    '    Journey: pre-call preparation email uses |EVENT_DATE| etc.',
])


# ── Save ────────────────────────────────────────────────────
out = 'Docs/ALW_Email_Journey_Map.docx'
doc.save(out)
print(f'Wrote: {out}')
