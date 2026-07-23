"""
Convert the single-source manual (cms/src/data/anna-manual.md) into a
Word .docx for Anna to read outside the CMS.

Run: python scripts/manual-to-docx.py
Output: Docs/ANNA_USER_MANUAL.docx (overwrites)

The .md is the source of truth — the Help · Ask chatbot reads it directly.
This script exists only so Anna can read the same content in Word.
"""

import io
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / 'cms' / 'src' / 'data' / 'anna-manual.md'
OUT = ROOT / 'Docs' / 'ANNA_USER_MANUAL.docx'

BOLD_RE = re.compile(r'\*\*(.+?)\*\*')
ITALIC_RE = re.compile(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)')
CODE_RE = re.compile(r'`([^`]+)`')
LINK_RE = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')


def add_runs(paragraph, text):
    """Add runs to a paragraph, honouring **bold**, *italic*, `code`, [link](url)."""
    # Simple tokenisation — split on inline markers, keep them as spans.
    # Not a full markdown parser, but good enough for a straight read.
    cursor = 0
    tokens = []
    for m in re.finditer(r'(\*\*.+?\*\*|(?<!\*)\*(?!\*).+?(?<!\*)\*(?!\*)|`[^`]+`|\[[^\]]+\]\([^)]+\))', text):
        if m.start() > cursor:
            tokens.append(('text', text[cursor:m.start()]))
        chunk = m.group(0)
        if chunk.startswith('**'):
            tokens.append(('bold', chunk[2:-2]))
        elif chunk.startswith('`'):
            tokens.append(('code', chunk[1:-1]))
        elif chunk.startswith('['):
            link = LINK_RE.match(chunk)
            tokens.append(('link', link.group(1), link.group(2)))
        else:
            tokens.append(('italic', chunk[1:-1]))
        cursor = m.end()
    if cursor < len(text):
        tokens.append(('text', text[cursor:]))

    for tok in tokens:
        if tok[0] == 'text':
            r = paragraph.add_run(tok[1])
        elif tok[0] == 'bold':
            r = paragraph.add_run(tok[1]); r.bold = True
        elif tok[0] == 'italic':
            r = paragraph.add_run(tok[1]); r.italic = True
        elif tok[0] == 'code':
            r = paragraph.add_run(tok[1])
            r.font.name = 'Consolas'
            r.font.color.rgb = RGBColor(0x88, 0x11, 0x88)
        elif tok[0] == 'link':
            r = paragraph.add_run(tok[1])
            r.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
            r.underline = True


def convert(md_text: str) -> Document:
    doc = Document()
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.startswith('```'):
            j = i + 1
            code_lines = []
            while j < len(lines) and not lines[j].startswith('```'):
                code_lines.append(lines[j]); j += 1
            p = doc.add_paragraph()
            r = p.add_run('\n'.join(code_lines))
            r.font.name = 'Consolas'
            r.font.size = Pt(9)
            i = j + 1
            continue

        # Table (markdown pipe table)
        if line.strip().startswith('|') and i + 1 < len(lines) and re.match(r'^\s*\|[\s\-:|]+\|\s*$', lines[i + 1]):
            header = [c.strip() for c in line.strip().strip('|').split('|')]
            j = i + 2
            rows = []
            while j < len(lines) and lines[j].strip().startswith('|'):
                rows.append([c.strip() for c in lines[j].strip().strip('|').split('|')])
                j += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(header))
            table.style = 'Light Grid Accent 1'
            for c, h in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.text = ''
                add_runs(cell.paragraphs[0], h)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
            for ri, row in enumerate(rows):
                for c, val in enumerate(row):
                    if c < len(header):
                        cell = table.rows[1 + ri].cells[c]
                        cell.text = ''
                        add_runs(cell.paragraphs[0], val)
            i = j
            continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            level = min(len(m.group(1)), 4)
            p = doc.add_heading(level=level)
            add_runs(p, m.group(2))
            i += 1
            continue

        # Bullet
        if re.match(r'^\s*[-*]\s+', line):
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, re.sub(r'^\s*[-*]\s+', '', line))
            i += 1
            continue

        # Numbered
        if re.match(r'^\s*\d+\.\s+', line):
            p = doc.add_paragraph(style='List Number')
            add_runs(p, re.sub(r'^\s*\d+\.\s+', '', line))
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph — collapse consecutive non-blank lines
        buf = [line]
        j = i + 1
        while j < len(lines) and lines[j].strip() and not re.match(r'^(#|```|\||\s*[-*]\s|\s*\d+\.\s)', lines[j]):
            buf.append(lines[j])
            j += 1
        p = doc.add_paragraph()
        add_runs(p, ' '.join(buf))
        i = j

    return doc


def main():
    if not SRC.exists():
        print(f'ERROR: source not found at {SRC}')
        sys.exit(1)
    md_text = SRC.read_text(encoding='utf-8')
    doc = convert(md_text)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f'Wrote {OUT} ({OUT.stat().st_size // 1024} KB)')


if __name__ == '__main__':
    main()
