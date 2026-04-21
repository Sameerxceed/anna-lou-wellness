from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style definitions ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

DEEP_PURPLE = RGBColor(0x1E, 0x10, 0x40)
MID_PURPLE = RGBColor(0x4A, 0x2D, 0x8A)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
GREY_TEXT = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DEEP_PURPLE if level <= 2 else MID_PURPLE
        if level == 1:
            run.font.size = Pt(22)
        elif level == 2:
            run.font.size = Pt(16)
        elif level == 3:
            run.font.size = Pt(13)
    h.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    h.paragraph_format.space_after = Pt(8)
    return h

def add_para(text, bold=False, italic=False, size=None, color=None, align=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10.5)
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    else:
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    return p

def add_table_with_header(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h_text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h_text)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E1040"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            # Alternate row shading
            if r_idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4EEF9"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # spacer
    return table

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('─' * 60)
    run.font.color.rgb = RGBColor(0xC8, 0xB8, 0xE8)
    run.font.size = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ANNA LOU WELLNESS')
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = DEEP_PURPLE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Website Rebuild Proposal')
run.font.size = Pt(18)
run.font.color.rgb = MID_PURPLE
p.paragraph_format.space_after = Pt(30)

add_separator()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('annalouwellness.com')
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = GOLD

for _ in range(4):
    doc.add_paragraph()

add_para('Prepared by: Xceed Code', size=11, color=GREY_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Date: 28 March 2026', size=11, color=GREY_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Version: 1.0', size=11, color=GREY_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════

add_heading_styled('Table of Contents', 1)
toc_items = [
    '1. Executive Summary',
    '2. Understanding the Brief',
    '3. Complete Feature List',
    '4. Technology Stack & Code Reuse',
    '5. Three-Phase Development Plan',
    '6. Design Direction',
    '7. Investment Summary',
    '8. Why This Approach',
    '9. Next Steps',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.color.rgb = MID_PURPLE
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════

add_heading_styled('1. Executive Summary', 1)

add_para(
    'This proposal outlines the complete rebuild of annalouwellness.com as a story-led, editorial wellness '
    'platform built around the Reset Stories brand positioning. The site will serve as the coaching, healing, '
    'and transformation hub for Anna Lou Wellness \u2014 distinct from annalouoflondon.com (the jewellery brand), '
    'yet cross-linked as a coherent brand ecosystem.',
    space_after=8
)

add_para(
    'The design will be magazine-like with a rustic/retro editorial aesthetic \u2014 warm, textured, emotionally '
    'resonant \u2014 inspired by the clarity of thejamiesea.com but differentiated through the Reset Stories '
    'narrative engine and a curated symbolic shop.',
    space_after=8
)

add_para(
    'We propose a three-phase development approach, leveraging our proven Astro + Strapi technology stack '
    '(successfully delivered for Ardea Gardens) to deliver a fast, scalable, and content-editor-friendly platform.',
    space_after=8
)

add_separator()

# ════════════════════════════════════════════════════════════
# 2. UNDERSTANDING THE BRIEF
# ════════════════════════════════════════════════════════════

add_heading_styled('2. Understanding the Brief', 1)

add_heading_styled('The Client', 2)
add_para(
    'Anna Lou is a somatic trauma-informed coach, ICF-accredited mentor, and founder of Anna Lou of London '
    '(jewellery brand sold at Harrods, Selfridges, Harvey Nichols, and internationally). She has consciously '
    'transitioned from luxury retail to coaching and healing work \u2014 and that transformation journey is the brand story.'
)

add_heading_styled('The Core Positioning: Reset Stories', 2)
add_para(
    'The new site leads with stories, not services. Every element \u2014 from the homepage hero to the shop \u2014 '
    'is anchored in the concept of personal transformation ("resets"). The jewellery and symbolic pieces on this '
    'site are not fashion products \u2014 they are physical anchors for transformation moments.'
)

add_heading_styled('Primary Audience', 2)
add_para(
    'UK-based female founders, entrepreneurs, and leaders (30s\u201340s). High-achieving women experiencing burnout, '
    'nervous system overload, or the gap between external success and internal wellbeing. Spiritually curious, '
    'self-aware, ready to invest. Core geography: London and surrounding areas, with global reach via online '
    'programmes and memberships.'
)

add_heading_styled('Design Direction', 2)
add_bullet('Magazine-like editorial feel \u2014 reads like a person, not a menu')
add_bullet('Rustic/retro aesthetic \u2014 warm, textured, film-grain photography feel')
add_bullet('Inspired by thejamiesea.com\'s minimal navigation, moody visuals, and emotional-first approach')
add_bullet('Brand palette: Deep Purple, Gold, Lavender, Blush \u2014 applied with warmth and intentionality')
add_bullet('Typography: Serif-forward (editorial authority with warmth), clean sans-serif for UI elements')

add_heading_styled('Brand Palette', 2)
add_table_with_header(
    ['Colour', 'Hex', 'Usage'],
    [
        ['Deep Purple', '#1E1040', 'Headers, accents, depth'],
        ['Mid Purple', '#4A2D8A', 'Interactive elements, links'],
        ['Gold', '#C9A84C', 'CTAs, highlights, luxury feel'],
        ['Lavender', '#C8B8E8', 'Backgrounds, soft sections'],
        ['Blush', '#F4EEF9', 'Page backgrounds, cards'],
        ['Light Gold', '#FAF3E0', 'Alternate backgrounds, warmth'],
    ],
    col_widths=[3.5, 3, 9]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 3. COMPLETE FEATURE LIST
# ════════════════════════════════════════════════════════════

add_heading_styled('3. Complete Feature List', 1)

# 3.1 Site Structure
add_heading_styled('3.1 Site Structure & Navigation', 2)
add_table_with_header(
    ['#', 'Feature', 'Description'],
    [
        ['1', 'Minimal Top Navigation', 'Maximum 6 items: Home, Work With Me, Shop, Events, About, Community. No deep nesting.'],
        ['2', 'Mobile-First Responsive', 'Full mobile optimisation for booking and shopping journeys'],
        ['3', 'Sticky/Animated Header', 'Elegant header with scroll behaviour'],
        ['4', 'Footer with Signature Offers', 'Clear links to key services, social icons, newsletter signup, cross-link to annalouoflondon.com'],
    ],
    col_widths=[1, 4, 11]
)

# 3.2 Homepage
add_heading_styled('3.2 Homepage', 2)
add_table_with_header(
    ['#', 'Feature', 'Description'],
    [
        ['5', 'Hero Section', 'Single powerful headline + atmospheric video/image background with moody overlay. "Reset Stories" positioning immediately felt.'],
        ['6', 'Three Visitor Pathways', 'Clear CTAs: Work With Me, Shop, Events \u2014 immediate clarity for any visitor'],
        ['7', 'Reset Stories Intro', 'Story-led introduction section \u2014 transformation narrative, not a sales pitch'],
        ['8', 'Testimonials / Social Proof', 'Real client testimonials with name, context, and genuine quotes'],
        ['9', 'Press & Credentials Strip', 'ICF, CPD, TRE certifications + media logos (Harrods, Selfridges, Harvey Nichols, Disney, etc.)'],
        ['10', 'Curated Instagram', 'Handpicked 3\u20134 posts chosen for emotional resonance (not auto-feed)'],
        ['11', 'Newsletter / Lead Magnet CTA', 'Email capture tied to free resources (nervous system recalibration download)'],
    ],
    col_widths=[1, 4, 11]
)

# 3.3 Work With Me
add_heading_styled('3.3 Work With Me (Services)', 2)
add_table_with_header(
    ['#', 'Feature', 'Description'],
    [
        ['12', 'Consolidated Services Page', 'One clean page with 4 categories, no duplication'],
        ['13', 'Healing & Somatics', 'TRE, breathwork, Flash EMDR, Brainspotting, nervous system regulation'],
        ['14', 'Coaching & Mentorship', '1:1 sessions, intensive days, 12-week programmes'],
        ['15', 'Corporate & Retreats', 'Corporate Wellness Retreat (half/full day) + Monthly Wellbeing Retainer \u2014 two distinct offerings with separate enquiry flows'],
        ['16', 'Gatherings & Events', 'The Returning Circle, crystal parties, houseboat retreat days'],
        ['17', 'Pricing Display', 'Visible where appropriate, or clear enquiry CTA'],
        ['18', 'Calendly Integration', 'Free 15-minute discovery call booking (primary CTA site-wide)'],
        ['19', 'Stripe Checkout', 'Seamless on-brand payment for paid programmes and services'],
    ],
    col_widths=[1, 4, 11]
)

# 3.4 Shop
add_heading_styled('3.4 Shop (Curated Symbolic Pieces)', 2)
add_table_with_header(
    ['#', 'Feature', 'Description'],
    [
        ['20', 'Shop Landing Page', 'Beautiful, elevated product browsing \u2014 not an afterthought'],
        ['21', 'Category Navigation', 'Clean categories: Crystals, Crystal Jewellery, Sage & Palo Santo, Gifts'],
        ['22', 'Product Detail Pages', 'Story-connected product pages \u2014 each item earns its place through its narrative meaning'],
        ['23', 'Digital Products', 'Purchasable programmes and digital products alongside physical items'],
        ['24', 'Cart & Checkout', 'Seamless cart with Stripe-integrated payment'],
        ['25', 'Stock Management', 'Real-time inventory via CMS'],
        ['26', 'Cross-Link to annalouoflondon.com', 'Clear pathway for visitors wanting the full jewellery range'],
    ],
    col_widths=[1, 4, 11]
)

# 3.5 Events
add_heading_styled('3.5 Events', 2)
add_table_with_header(
    ['#', 'Feature', 'Description'],
    [
        ['27', 'Events Page', 'Clean, scannable layout: date, brief description, ticket link'],
        ['28', 'Event Categories', 'Workshops, Returning Circle gatherings, crystal parties, houseboat retreats'],
        ['29', 'Ticket Integration', 'External ticket links (Eventbrite, etc.) or direct Stripe purchase'],
        ['30', 'Past Events Archive', 'Gallery/archive of past events for social proof'],
    ],
    col_widths=[1, 4, 11]
)

# 3.6 About
add_heading_styled('3.6 About', 2)
add_table_with_header(
    ['#', 'Feature', 'Description'],
    [
        ['31', 'Story-Led About Page', 'Full editorial narrative: Anna Lou of London journey \u2192 Portobello Market to Harrods \u2192 conscious reset \u2192 coaching & healing'],
        ['32', 'Credentials & Press', 'ICF, CPD, TRE certifications, podcast appearances, media features'],
        ['33', 'Zen Hustler Philosophy', 'Woven into the About narrative as a framing device'],
        ['34', 'Professional Photography', 'Editorial, emotionally resonant imagery (not standard headshots)'],
    ],
    col_widths=[1, 4, 11]
)

# 3.7 Community
add_heading_styled('3.7 Community (Key Differentiator)', 2)
add_table_with_header(
    ['#', 'Feature', 'Description'],
    [
        ['35', 'Community Hub', 'Central page for free resources, membership, and user-generated content'],
        ['36', 'Crystal Clear Collective', 'Membership signup with 14-day free trial (recurring revenue)'],
        ['37', 'Free Lead Magnets', 'Nervous system recalibration download + narcissistic abuse ebook \u2014 consolidated in one place'],
        ['38', 'User Experience Sharing', 'Members can share their personal reset stories / transformation experiences'],
        ['39', 'Commenting System', 'Other users can comment on and engage with shared stories'],
        ['40', 'Community Moderation', 'Admin tools for approving/moderating user posts and comments'],
        ['41', 'User Profiles', 'Basic profiles for community members (avatar, bio, join date)'],
    ],
    col_widths=[1, 4, 11]
)

# 3.8 Blog
add_heading_styled('3.8 Blog / Reset Stories Content', 2)
add_table_with_header(
    ['#', 'Feature', 'Description'],
    [
        ['42', 'Blog / Stories Section', 'Home for Reset Stories content \u2014 client stories (anonymised), Anna\'s journey, transformation narratives'],
        ['43', 'Blog Categories / Tags', 'Organised by theme (healing, business, crystals, personal growth)'],
        ['44', 'SEO-Optimised Posts', 'Rich meta tags, Open Graph, structured data for every post'],
        ['45', 'Related Content', 'Stories connected to relevant services and shop items'],
    ],
    col_widths=[1, 4, 11]
)

# 3.9 Podcast
add_heading_styled('3.9 Podcast', 2)
add_table_with_header(
    ['#', 'Feature', 'Description'],
    [
        ['46', 'Podcast Page', 'Dedicated page with episode listings (similar to Jamie Sea\'s "The Jamie Sea Show")'],
        ['47', 'Episode Player', 'Embedded audio player for each episode'],
        ['48', 'Episode Details', 'Show notes, guest info, timestamps, related links'],
        ['49', 'Subscribe Links', 'Apple Podcasts, Spotify, etc.'],
    ],
    col_widths=[1, 4, 11]
)

# 3.10 Contact
add_heading_styled('3.10 Contact', 2)
add_table_with_header(
    ['#', 'Feature', 'Description'],
    [
        ['50', 'Contact Page', 'Clean form, email, phone, social links'],
        ['51', 'Enquiry Routing', 'Different enquiry types (general, corporate, press)'],
        ['52', 'Map / Location', 'Twickenham area reference for in-person sessions'],
    ],
    col_widths=[1, 4, 11]
)

# 3.11 Cross-Cutting
add_heading_styled('3.11 Cross-Cutting Features', 2)
add_table_with_header(
    ['#', 'Feature', 'Description'],
    [
        ['53', 'SEO Infrastructure', 'Sitemap, robots.txt, structured data, Open Graph for all pages'],
        ['54', 'Cookie Consent (GDPR)', 'Cookie banner with consent management'],
        ['55', 'Analytics', 'Google Analytics integration'],
        ['56', 'Email Integration', 'Resend or similar for transactional emails (order confirmations, community notifications)'],
        ['57', 'Image Optimisation', 'Automatic image processing with Sharp (WebP, lazy loading)'],
        ['58', 'Page Transitions', 'Smooth View Transitions API for editorial feel'],
        ['59', 'Scroll Animations', 'Subtle fade-up, reveal-on-scroll for magazine-like reading experience'],
        ['60', 'Lightbox Galleries', 'Fullscreen image viewer with keyboard navigation'],
        ['61', 'Brand Ecosystem Links', 'Coherent cross-linking between annalouwellness.com and annalouoflondon.com'],
        ['62', '404 Page', 'On-brand not-found page'],
        ['63', 'Privacy Policy & Terms', 'GDPR-compliant legal pages'],
    ],
    col_widths=[1, 4, 11]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 4. TECHNOLOGY STACK & CODE REUSE
# ════════════════════════════════════════════════════════════

add_heading_styled('4. Technology Stack & Code Reuse', 1)

add_heading_styled('Proven Stack (from Ardea Gardens)', 2)
add_para(
    'We will leverage the same technology stack successfully delivered for ardea-gardens.com, '
    'providing a proven, battle-tested foundation:'
)

add_table_with_header(
    ['Layer', 'Technology', 'Why'],
    [
        ['Frontend', 'Astro 4 + React Islands', 'Static generation for speed + React for interactivity (cart, checkout, community)'],
        ['CMS', 'Strapi 5 (Headless)', 'Content editor-friendly, REST API, media management'],
        ['Styling', 'Tailwind CSS', 'Rapid theming with custom palette, responsive design'],
        ['Database', 'SQLite (dev) / PostgreSQL (prod)', 'Proven, scalable, cost-effective'],
        ['Payments', 'Stripe', 'Checkout for services, shop, and events'],
        ['Hosting', 'Vercel + Railway', 'Auto-deploy, CDN, low cost (~\u00a36/month)'],
        ['Email', 'Resend', 'Transactional emails, community notifications'],
    ],
    col_widths=[2.5, 5, 8.5]
)

add_heading_styled('Code Reuse from Ardea Gardens (~60\u201370%)', 2)
add_para(
    'Our existing codebase provides substantial reusable components, significantly reducing '
    'development time, cost, and risk:'
)

add_table_with_header(
    ['Component', 'Reuse Level', 'Notes'],
    [
        ['Astro Project Structure', 'Full reuse', 'Same layout, routing, config patterns'],
        ['Strapi Backend Architecture', 'Full reuse', 'Same API patterns, permissions, seeding'],
        ['Base Layout', 'Adapt', 'Retheme with Anna Lou palette/typography'],
        ['Navigation Component', 'Adapt', 'Simplify to 6-item structure'],
        ['Footer Component', 'Adapt', 'Retheme + add signature offers'],
        ['Shop Pages (listing + detail)', 'Full reuse', 'Same product grid, filtering, detail layout'],
        ['Cart System', 'Full reuse', 'sessionStorage cart, cross-component sync'],
        ['Checkout + Stripe', 'Full reuse', 'Payment flow, order creation, webhooks'],
        ['CMS Data Layer', 'Full reuse', 'API client, caching, fallback logic'],
        ['SEO Head Component', 'Full reuse', 'Meta tags, Open Graph, structured data'],
        ['Cookie Banner', 'Full reuse', 'GDPR consent management'],
        ['Lightbox Gallery', 'Full reuse', 'Fullscreen viewer, keyboard navigation'],
        ['Contact Page', 'Adapt', 'Same form structure, different fields'],
        ['Events Page', 'Adapt', 'Same event listing pattern'],
        ['Image Optimisation', 'Full reuse', 'Sharp integration, lazy loading'],
        ['Scroll Animations', 'Full reuse', 'IntersectionObserver reveal system'],
        ['Deployment Config', 'Full reuse', 'Vercel + Railway, webhooks, auto-deploy'],
        ['Order Management API', 'Full reuse', 'Order creation, status, dashboard, CSV export'],
    ],
    col_widths=[4, 2.5, 9.5]
)

add_heading_styled('New Development Required', 2)
add_table_with_header(
    ['Component', 'Effort', 'Notes'],
    [
        ['Reset Stories Theme/Design', 'Medium', 'New palette, typography, rustic/retro editorial CSS'],
        ['Homepage (story-led)', 'Medium', 'Video hero, pathways, credentials strip'],
        ['Work With Me Page', 'Medium', '4-category layout with booking CTAs'],
        ['Blog/Stories System', 'Medium', 'New Strapi content type + Astro pages'],
        ['Podcast Page', 'Low', 'Episode listing + embedded players'],
        ['Community Features', 'High', 'User accounts, story sharing, comments, moderation'],
        ['About Page (editorial)', 'Low', 'Content-driven, mostly layout'],
        ['Calendly Integration', 'Low', 'Embed widget in booking CTAs'],
        ['Cross-Site Linking', 'Low', 'Smart banners/CTAs linking to annalouoflondon.com'],
    ],
    col_widths=[4.5, 2, 9.5]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 5. THREE-PHASE DEVELOPMENT PLAN
# ════════════════════════════════════════════════════════════

add_heading_styled('5. Three-Phase Development Plan', 1)

# Phase 1
add_heading_styled('PHASE 1 \u2014 Foundation & Core Site', 2)
add_para('Timeline: 4\u20135 weeks', bold=True, color=GOLD, size=12)
add_para('Goal: Launch a complete, polished website with all essential pages and the shop', italic=True, space_after=12)

add_heading_styled('Design & Theming', 3)
add_bullet('Brand palette implementation (Deep Purple, Gold, Lavender, Blush, Light Gold)')
add_bullet('Typography system (serif-forward editorial feel matched to Anna Lou\'s brand)')
add_bullet('Rustic/retro magazine aesthetic: textured backgrounds, film-grain overlays, warm tones')
add_bullet('Responsive design system (mobile-first)')
add_bullet('Page transition animations + scroll reveal effects')

add_heading_styled('Pages & Features', 3)
add_bullet('', bold_prefix='Homepage \u2014 ')
add_para('Story-led hero (video/image + powerful headline), three visitor pathways, testimonials section, press/credentials strip, curated Instagram section, newsletter CTA', size=10)
add_bullet('', bold_prefix='Work With Me \u2014 ')
add_para('Consolidated services page with 4 categories, Calendly embed for discovery calls, enquiry forms for corporate', size=10)
add_bullet('', bold_prefix='Shop \u2014 ')
add_para('Product listing with category filtering, product detail pages, story-connected product descriptions', size=10)
add_bullet('', bold_prefix='Cart & Checkout \u2014 ')
add_para('Full cart system with Stripe-integrated payment', size=10)
add_bullet('', bold_prefix='Events \u2014 ')
add_para('Scannable event listings with dates and ticket links', size=10)
add_bullet('', bold_prefix='About \u2014 ')
add_para('Story-led editorial page (Anna Lou of London \u2192 coaching journey), credentials, Zen Hustler philosophy', size=10)
add_bullet('', bold_prefix='Contact \u2014 ')
add_para('Clean form with enquiry routing', size=10)
add_bullet('Privacy Policy, Terms, 404 Page')

add_heading_styled('Infrastructure', 3)
add_bullet('Strapi CMS setup with all content types (Homepage, Services, Products, Categories, Events, About, Contact, Site Settings, Testimonials, Press, Orders)')
add_bullet('Strapi admin panel (white-labelled for Anna Lou Wellness)')
add_bullet('Deployment pipeline (Vercel + Railway) with webhook auto-deploy')
add_bullet('SEO infrastructure (sitemap, robots.txt, structured data)')
add_bullet('Cookie consent banner + Google Analytics')
add_bullet('Stripe webhook for order confirmation')
add_bullet('Email setup (order confirmations via Resend)')

add_heading_styled('Phase 1 Deliverables', 3)
add_bullet('Fully functional website at annalouwellness.com')
add_bullet('CMS admin panel with user manual')
add_bullet('Shop with Stripe payments')
add_bullet('All core pages live and editable via CMS')
add_bullet('Mobile-optimised across all breakpoints')
add_bullet('SEO-ready with analytics integration')

add_separator()

# Phase 2
add_heading_styled('PHASE 2 \u2014 Content Engine', 2)
add_para('Timeline: 2\u20133 weeks', bold=True, color=GOLD, size=12)
add_para('Goal: Build the Reset Stories content ecosystem that drives traffic and conversion', italic=True, space_after=12)

add_heading_styled('Blog / Reset Stories', 3)
add_bullet('Blog content type in Strapi (title, body, featured image, category, tags, author, SEO fields, related services, related products)')
add_bullet('Blog listing page with category/tag filtering')
add_bullet('Individual blog post pages with rich editorial typography layout')
add_bullet('Related content suggestions (other stories, connected services, symbolic shop items)')
add_bullet('Social sharing buttons + RSS feed')

add_heading_styled('Podcast', 3)
add_bullet('Podcast content type in Strapi (episode title, description, audio embed URL, guest info, show notes)')
add_bullet('Podcast landing page with episode grid')
add_bullet('Individual episode pages with embedded audio player')
add_bullet('Subscribe links (Apple Podcasts, Spotify, etc.)')

add_heading_styled('Lead Magnets & Email', 3)
add_bullet('Lead magnet content type in Strapi (title, description, download file)')
add_bullet('Download delivery flow (email capture \u2192 file delivery via Resend)')
add_bullet('Nervous system recalibration guide + Narcissistic abuse ebook')
add_bullet('Newsletter signup integration')

add_heading_styled('Cross-Site Linking', 3)
add_bullet('Smart contextual banners linking to annalouoflondon.com')
add_bullet('"Explore the full collection" CTAs on shop pages')
add_bullet('Shared footer branding between both sites')

add_heading_styled('Phase 2 Deliverables', 3)
add_bullet('Blog system live with initial Reset Stories content')
add_bullet('Podcast page with episode management via CMS')
add_bullet('Lead magnets downloadable via email capture')
add_bullet('Content calendar ready \u2014 Anna can publish stories, episodes, and resources independently')
add_bullet('Cross-linking ecosystem between both Anna Lou sites')

add_separator()

# Phase 3
add_heading_styled('PHASE 3 \u2014 Community & Membership', 2)
add_para('Timeline: 3\u20134 weeks', bold=True, color=GOLD, size=12)
add_para('Goal: Launch the community platform and Crystal Clear Collective membership', italic=True, space_after=12)

add_heading_styled('User Accounts & Authentication', 3)
add_bullet('User registration and login (email/password)')
add_bullet('User profiles (avatar, display name, bio, join date)')
add_bullet('Password reset flow + email verification')
add_bullet('Optional: OAuth (Google sign-in) for convenience')

add_heading_styled('Community Hub', 3)
add_bullet('"Share Your Reset Story" \u2014 authenticated users can submit their transformation experiences (title, story text, optional image)')
add_bullet('Commenting system \u2014 users can comment on shared stories and blog posts')
add_bullet('Likes/Hearts \u2014 simple engagement mechanic on stories')
add_bullet('Moderation dashboard in Strapi \u2014 approve/reject user stories, flag/remove comments')
add_bullet('Community guidelines page')
add_bullet('Notification system \u2014 email alerts for comments on your story, new community stories')

add_heading_styled('Crystal Clear Collective Membership', 3)
add_bullet('Membership tier management in Strapi (name, description, price, benefits)')
add_bullet('14-day free trial flow')
add_bullet('Stripe subscription integration (recurring billing)')
add_bullet('Members-only content gating (exclusive stories, resources, events)')
add_bullet('Member dashboard \u2014 subscription status, billing history, exclusive content access')
add_bullet('Cancellation/pause flow')

add_heading_styled('Phase 3 Deliverables', 3)
add_bullet('Full community platform with user registration, story sharing, and commenting')
add_bullet('Crystal Clear Collective membership with Stripe recurring billing')
add_bullet('Moderation tools for Anna in the CMS')
add_bullet('Member-only content areas')
add_bullet('Notification system for community engagement')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 6. DESIGN DIRECTION
# ════════════════════════════════════════════════════════════

add_heading_styled('6. Design Direction', 1)

add_heading_styled('Magazine / Editorial Aesthetic', 2)
add_para(
    'The site should feel like opening a beautifully designed wellness magazine \u2014 each page a curated spread, '
    'not a web application. Key design principles:'
)
add_bullet('Full-bleed imagery with moody, film-grain overlays')
add_bullet('Generous white space \u2014 let content breathe')
add_bullet('Pull quotes and large display typography between sections')
add_bullet('Story-first layouts \u2014 text and images interwoven editorially')
add_bullet('Textured backgrounds \u2014 subtle paper/linen textures for rustic warmth')
add_bullet('Muted, warm colour overlays on photography')
add_bullet('Handwritten or vintage-style accent fonts for headings/callouts')
add_bullet('Rounded, inviting CTAs \u2014 not sharp corporate buttons')

add_heading_styled('Rustic / Retro Elements', 2)
add_bullet('Subtle grain/noise texture overlays on images and backgrounds')
add_bullet('Warm, slightly desaturated photography treatment')
add_bullet('Vintage-inspired decorative dividers between sections')
add_bullet('Earthy, layered backgrounds (cream, parchment, soft gold)')
add_bullet('Serif-dominant typography conveying editorial authority with warmth')

add_heading_styled('Reference Points', 2)
add_bullet('', bold_prefix='thejamiesea.com \u2014 ')
add_para('Navigation simplicity, emotional-first design, moody imagery, atmospheric video backgrounds', size=10)
add_bullet('', bold_prefix='kinfolk.com \u2014 ')
add_para('Magazine editorial layout, typography-forward, generous breathing room', size=10)
add_bullet('', bold_prefix='goop.com \u2014 ')
add_para('Wellness brand with integrated editorial content + commerce', size=10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 7. INVESTMENT SUMMARY
# ════════════════════════════════════════════════════════════

add_heading_styled('7. Investment Summary', 1)

add_heading_styled('Development Timeline', 2)
add_table_with_header(
    ['Phase', 'Scope', 'Timeline'],
    [
        ['Phase 1', 'Foundation & Core Site (all pages, shop, payments)', '4\u20135 weeks'],
        ['Phase 2', 'Content Engine (blog, podcast, lead magnets, cross-linking)', '2\u20133 weeks'],
        ['Phase 3', 'Community & Membership (user accounts, story sharing, commenting, Crystal Clear Collective)', '3\u20134 weeks'],
        ['Total', 'Complete platform delivery', '9\u201312 weeks'],
    ],
    col_widths=[2, 9.5, 3]
)

add_heading_styled('Ongoing Costs (Post-Launch)', 2)
add_table_with_header(
    ['Service', 'Monthly Cost'],
    [
        ['Vercel (frontend hosting)', 'Free tier / ~\u00a30'],
        ['Railway (CMS + database)', '~\u00a35\u201310'],
        ['Stripe', '1.4% + 20p per transaction'],
        ['Resend (email)', 'Free tier up to 3k emails/month'],
        ['Domain (annual)', '~\u00a315/year'],
        ['Total estimated', '~\u00a35\u201310/month'],
    ],
    col_widths=[6, 6]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 8. WHY THIS APPROACH
# ════════════════════════════════════════════════════════════

add_heading_styled('8. Why This Approach', 1)

add_heading_styled('Why Astro + Strapi (vs WordPress / Showit)', 2)
add_para(
    'The project brief mentions a discussion about WordPress vs Showit. We recommend Astro + Strapi for the following reasons:'
)

add_table_with_header(
    ['Criteria', 'WordPress / Showit', 'Astro + Strapi (Our Approach)'],
    [
        ['Performance', 'Plugin-heavy, slower page loads', 'Static HTML on CDN \u2014 near-instant loads'],
        ['Security', 'Constant plugin vulnerabilities', 'No server-side attack surface'],
        ['Editorial Feel', 'Template-constrained', 'Fully custom, magazine-quality layouts'],
        ['Shop', 'WooCommerce (good but heavy)', 'Lightweight custom shop + Stripe'],
        ['Community', 'Plugin-dependent (BuddyPress)', 'Custom-built, exactly what\'s needed'],
        ['CMS Experience', 'WordPress admin (complex)', 'Clean Strapi admin (simple for editors)'],
        ['Hosting Cost', '\u00a320\u201350+/month', '~\u00a35\u201310/month'],
        ['Code Reuse', 'Start from scratch', '60\u201370% reuse from proven codebase'],
        ['Scalability', 'Server-dependent', 'CDN-first, globally fast'],
    ],
    col_widths=[2.5, 5, 6.5]
)

add_heading_styled('Why Three Phases', 2)
add_bullet('', bold_prefix='Phase 1 gets Anna live fast \u2014 ')
add_para('A complete, polished site replacing the current one within 4\u20135 weeks', size=10)
add_bullet('', bold_prefix='Phase 2 builds the content engine \u2014 ')
add_para('The Reset Stories narrative needs a blog and podcast to come alive; lead magnets start capturing emails', size=10)
add_bullet('', bold_prefix='Phase 3 is the growth play \u2014 ')
add_para('Community and membership create recurring revenue and sticky engagement, justified once the foundation is proven', size=10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 9. NEXT STEPS
# ════════════════════════════════════════════════════════════

add_heading_styled('9. Next Steps', 1)

steps = [
    ('Review this proposal', 'and confirm scope/priorities'),
    ('Schedule kick-off call', 'to align on design direction and content needs'),
    ('Photography brief', '\u2014 we will provide exact image specifications and quantities per section for the photographer'),
    ('Content gathering', '\u2014 testimonials, press logos, About page narrative, initial blog content'),
    ('Confirm podcast name', 'with Amina before build'),
    ('Phase 1 development begins', ''),
]
for i, (title, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    run_num = p.add_run(f'{i}. ')
    run_num.bold = True
    run_num.font.size = Pt(11)
    run_num.font.color.rgb = GOLD
    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(11)
    if desc:
        run_desc = p.add_run(f' {desc}')
        run_desc.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)

for _ in range(4):
    doc.add_paragraph()

add_separator()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared by Xceed Code')
run.italic = True
run.font.size = Pt(11)
run.font.color.rgb = MID_PURPLE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Building on a proven technology foundation to deliver a world-class wellness platform.')
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = GREY_TEXT

# ── Save ──
output_path = r'e:\Xceed\Code\Anna\Anna Lou Wellness - Website Rebuild Proposal.docx'
doc.save(output_path)
print(f'Proposal saved to: {output_path}')
